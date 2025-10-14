# ==== Standard Library ====
import calendar
import difflib
import hashlib
import html
import io
import json
import math
import logging
import os
import random
import re
import tempfile
import time
import textwrap
import urllib.parse
import urllib.parse as _urllib
from collections import Counter
from urllib.parse import urlsplit, parse_qs, urlparse, quote_plus
from datetime import date, timedelta, timezone as _timezone, UTC
from datetime import datetime
from datetime import datetime as _dt
from uuid import uuid4
from pathlib import Path
from typing import Any, Dict, Optional, Tuple, List, Iterable, MutableMapping, Sequence
from functools import lru_cache

# ==== Third-Party Packages ====
import bcrypt
import pandas as pd
import requests
import streamlit as st
import streamlit.components.v1 as components
from docx import Document
from google.cloud.firestore_v1 import FieldFilter
from firebase_admin import firestore  # Firebase
from openai import OpenAI
from src.styles import inject_global_styles
from src.lesson_language_support import gather_language_support
from src.discussion_board import (
    CLASS_DISCUSSION_LABEL,
    CLASS_DISCUSSION_LINK_TMPL,
    CLASS_DISCUSSION_ANCHOR,
    CLASS_DISCUSSION_PROMPT,
    CLASS_DISCUSSION_REMINDER,
    go_class_thread,
)
from src.topic_coach_persistence import (
    get_topic_coach_doc,
    load_topic_coach_state,
    persist_topic_coach_state,
)
from src.forum_timer import (
    _to_datetime_any,
    build_forum_reply_indicator_text,
    build_forum_timer_indicator,
)
from src.level_sync import sync_level_state

from flask import Flask
from auth import auth_bp
from src.routes.health import register_health_route
from src.group_schedules import load_group_schedules
from src.blog_feed import fetch_blog_feed
from src.blog_cards_widget import render_blog_cards
import src.schedule as _schedule
load_level_schedules = _schedule.load_level_schedules
refresh_level_schedules = getattr(_schedule, "refresh_level_schedules", lambda: None)

app = Flask(__name__)
app.register_blueprint(auth_bp)
register_health_route(app)

ICON_PATH = Path(__file__).parent / "static/icons/falowen-512.png"

st.set_page_config(
    page_title="Falowen – Your German Conversation Partner",
    page_icon=str(ICON_PATH),  # now uses official Falowen icon
    layout="wide",
)

# Load global CSS classes and variables
inject_global_styles()

st.markdown("""
<style>
html, body { overscroll-behavior-y: none; }
</style>
""", unsafe_allow_html=True)

HERR_FELIX_TYPING_HTML = (
    "<div class='typing-notice'>"
    "👨‍🏫 Herr Felix is typing"
    "<span class='typing'><span></span><span></span><span></span></span>"
    "</div>"
)

_TYPING_TRACKER_PREFIX = "__typing_meta__"
_TYPING_PING_INTERVAL = 4.0
_NEW_POST_TYPING_ID = "__new_post__"


def _safe_str(v, default: str = "") -> str:
    if v is None:
        return default
    if isinstance(v, float):
        try:
            if math.isnan(v):
                return default
        except Exception:
            pass
    s = str(v).strip()
    return "" if s.lower() in ("nan", "none") else s


def _safe_upper(v, default: str = "") -> str:
    s = _safe_str(v, default)
    return s.upper() if s else default


def _safe_lower(v, default: str = "") -> str:
    s = _safe_str(v, default)
    return s.lower() if s else default


def _typing_meta_key(draft_key: str) -> str:
    return f"{_TYPING_TRACKER_PREFIX}:{draft_key}"


def _update_typing_state(
    *,
    level: str,
    class_code: str,
    qid: str,
    draft_key: str,
    student_code: str,
    student_name: str,
    text: str,
) -> None:
    if not (level and class_code and qid and student_code):
        return

    meta_key = _typing_meta_key(draft_key)
    meta = st.session_state.get(meta_key, {"is_typing": False, "last_sent": 0.0})
    try:
        last_sent = float(meta.get("last_sent", 0.0))
    except Exception:
        last_sent = 0.0
    last_state = bool(meta.get("is_typing", False))

    is_typing = bool((text or "").strip())
    now_ts = time.time()
    should_send = False

    if is_typing != last_state:
        should_send = True
    elif is_typing and (now_ts - last_sent) >= _TYPING_PING_INTERVAL:
        should_send = True

    if should_send:
        set_typing_indicator(
            level,
            class_code,
            qid,
            student_code,
            student_name,
            is_typing=is_typing,
        )
        last_sent = now_ts

    st.session_state[meta_key] = {"is_typing": is_typing, "last_sent": last_sent}


def _clear_typing_state(
    *,
    level: str,
    class_code: str,
    qid: str,
    draft_key: str,
    student_code: str,
    student_name: str,
) -> None:
    if not (level and class_code and qid and student_code):
        return

    set_typing_indicator(
        level,
        class_code,
        qid,
        student_code,
        student_name,
        is_typing=False,
    )
    st.session_state.pop(_typing_meta_key(draft_key), None)


def _format_typing_banner(entries: List[Dict[str, Any]], current_code: str) -> str:
    names: List[str] = []
    for entry in entries:
        if entry.get("student_code") == current_code:
            continue
        name = _safe_str(entry.get("student_name", "")) or _safe_str(
            entry.get("student_code", "")
        )
        if name:
            names.append(name)
    if not names:
        return ""
    if len(names) == 1:
        return f"{names[0]} is typing…"
    if len(names) == 2:
        return f"{names[0]} and {names[1]} are typing…"
    return f"{', '.join(names[:-1])}, and {names[-1]} are typing…"


def _coerce_day(value: Any) -> Optional[int]:
    """Return ``value`` as an ``int`` day when possible."""

    if value is None:
        return None
    if isinstance(value, bool):
        return None
    if isinstance(value, (int, float)):
        try:
            return int(value)
        except (TypeError, ValueError):
            return None
    try:
        return int(str(value).strip())
    except Exception:
        return None


def _submission_block_reason(
    lesson: Optional[MutableMapping[str, Any]],
    schedule: Sequence[MutableMapping[str, Any]],
) -> str:
    """Return a message explaining why submissions are disabled for ``lesson``."""

    day = _coerce_day((lesson or {}).get("day"))
    if day is None:
        return ""

    first_day: Optional[int] = None
    for entry in schedule:
        first_day = _coerce_day((entry or {}).get("day"))
        if first_day is not None:
            break

    last_day: Optional[int] = None
    for entry in reversed(schedule):
        last_day = _coerce_day((entry or {}).get("day"))
        if last_day is not None:
            break

    if first_day is not None and day == first_day == 0:
        return "Day 0 is a tutorial day — there is nothing to submit yet."
    if last_day is not None and day == last_day:
        return "This final celebration day has no assignment to submit."
    return ""


def _timestamp_to_epoch(ts: Optional[datetime]) -> float:
    """Return ``ts`` as epoch seconds (or ``0.0`` if unavailable)."""

    if ts is None:
        return 0.0

    if isinstance(ts, (int, float)):
        try:
            return float(ts)
        except (TypeError, ValueError):
            return 0.0

    try:
        return float(ts.timestamp())  # type: ignore[call-arg]
    except Exception:
        return 0.0


def _topic_coach_state_key(base: str, student_code: str, level: str) -> str:
    student = _safe_str(student_code, "_") or "_"
    level_token = _safe_str(level, "_") or "_"
    return f"{base}:{student}:{level_token}"


def _initialise_topic_coach_session_state(
    session_state: MutableMapping[str, Any],
    *,
    student_code: str,
    level: str,
    messages: Iterable[Dict[str, Any]],
    qcount: Any,
    finalized: Any,
    focus_tips: Iterable[str] | None = None,
    identity_key: str = "_cchat_active_identity",
) -> Tuple[str, str, str, str]:
    """Return scoped Topic Coach session-state keys after initialising values."""

    student_token = _safe_str(student_code)
    level_token = _safe_str(level)
    identity = (student_token, level_token)
    previous_identity = session_state.get(identity_key)
    identity_changed = previous_identity != identity

    chat_key = _topic_coach_state_key("cchat_data_chat", student_token, level_token)
    qcount_key = _topic_coach_state_key("cchat_data_qcount", student_token, level_token)
    finalized_key = _topic_coach_state_key("cchat_data_finalized", student_token, level_token)
    focus_key = _topic_coach_state_key("cchat_data_focus", student_token, level_token)

    for legacy_key in ("cchat_data_chat", "cchat_data_qcount", "cchat_data_finalized"):
        session_state.pop(legacy_key, None)

    if identity_changed or chat_key not in session_state:
        session_state[chat_key] = list(messages or [])
    elif not session_state[chat_key] and messages:
        session_state[chat_key] = list(messages)

    try:
        qcount_value = int(qcount or 0)
    except Exception:
        qcount_value = 0
    qcount_value = max(0, qcount_value)
    if identity_changed or qcount_key not in session_state:
        session_state[qcount_key] = qcount_value

    finalized_value = bool(finalized)
    if identity_changed or finalized_key not in session_state:
        session_state[finalized_key] = finalized_value

    if identity_changed or focus_key not in session_state:
        session_state[focus_key] = list(focus_tips or [])
    elif not session_state[focus_key] and focus_tips:
        session_state[focus_key] = list(focus_tips)

    session_state[identity_key] = identity
    return chat_key, qcount_key, finalized_key, focus_key


def _extract_focus_tips_from_history(
    history: Iterable[Dict[str, Any]],
) -> List[str]:
    """Return up to three recurring correction themes from the transcript."""

    correction_lines: List[str] = []
    stop_pattern = re.compile(
        r"(?:^|\n)\s*(?:idea|💡|next question|question|focus|summary|tip|strengths|improvements?)",
        re.IGNORECASE,
    )

    for message in history:
        if not isinstance(message, dict) or message.get("role") != "assistant":
            continue
        raw = str(message.get("content") or "")
        if not raw:
            continue
        text = re.sub(r"<br\s*/?>", "\n", raw, flags=re.IGNORECASE)
        text = re.sub(r"</(?:div|p|li)>", "\n", text, flags=re.IGNORECASE)
        text = re.sub(r"<[^>]+>", " ", text)
        text = html.unescape(text)
        text = text.replace("**", "")
        text = text.replace("__", "")
        text = text.replace("*", "")
        text = re.sub(r"\r\n?", "\n", text)

        lower_text = text.lower()
        match = re.search(r"corrections?\s*:?", lower_text)
        if not match:
            continue
        segment = text[match.end() :]
        stop = stop_pattern.search(segment)
        if stop:
            segment = segment[: stop.start()]

        for line in segment.splitlines():
            cleaned = re.sub(r"^[\-•●▪▫➤▶︎\*\s]+", "", line).strip()
            if not cleaned:
                continue
            if len(cleaned) > 160:
                cleaned = cleaned[:157].rstrip() + "…"
            correction_lines.append(cleaned)

    counts: Counter[str] = Counter()
    canonical: Dict[str, str] = {}

    for line in correction_lines:
        normalised = re.sub(r"[^\wäöüÄÖÜß ]+", " ", line.lower()).strip()
        if not normalised:
            continue
        counts[normalised] += 1
        canonical.setdefault(normalised, line)

    tips = [canonical[key] for key, _ in counts.most_common(3)]
    return tips


def _resolve_class_name(
    raw_value: Any,
    *,
    level: str = "",
    default_suffix: str = "General",
) -> Tuple[str, str]:
    """Return ``(clean, resolved)`` class labels.

    ``clean`` is the sanitized class name treating "nan"/"None" as empty while
    ``resolved`` falls back to ``"{LEVEL} {default_suffix}"`` when available.
    """

    clean = _safe_str(raw_value)
    level_key = _safe_upper(level, "")
    fallback = f"{level_key} {default_suffix}".strip() if level_key else ""
    return clean, clean or fallback


def _build_missing_code_link(
    name: str = "",
    *,
    level: str = "",
    lesson_info: Optional[Dict[str, Any]] = None,
) -> str:
    display_name = _safe_str(name, "Student") or "Student"
    level_text = _safe_upper(level, "")
    lesson = lesson_info or {}

    day = _safe_str(lesson.get("day"))
    chapter = _safe_str(lesson.get("chapter"))
    topic = _safe_str(lesson.get("topic"))

    subject_parts: List[str] = []
    if level_text:
        subject_parts.append(f"Level {level_text}")
    if day:
        subject_parts.append(f"Day {day}")
    if chapter:
        subject_parts.append(f"Chapter {chapter}")

    subject = f"{display_name} - Missing student code"
    if subject_parts:
        subject = f"{subject} ({' • '.join(subject_parts)})"

    lesson_line_parts: List[str] = []
    if day:
        lesson_line_parts.append(f"Day {day}")
    if chapter:
        lesson_line_parts.append(f"Chapter {chapter}")
    if topic:
        lesson_line_parts.append(f"Topic: {topic}")

    body_lines: List[str] = [
        "Hello Learn German Ghana team,",
        "",
        "I couldn't find my student code in Falowen.",
        "",
        f"Name: {display_name}",
    ]
    if level_text:
        body_lines.append(f"Level: {level_text}")
    if lesson_line_parts:
        body_lines.append("Lesson: " + " • ".join(lesson_line_parts))
    body_lines.extend(["", "My work (paste below):", ""])

    body = "\n".join(body_lines)
    return (
        "mailto:learngermanghana@gmail.com"
        f"?subject={quote_plus(subject)}&body={quote_plus(body)}"
    )


def _show_missing_code_warning(
    name: str = "",
    *,
    level: str = "",
    lesson_info: Optional[Dict[str, Any]] = None,
) -> None:
    mailto_link = _build_missing_code_link(
        name=name,
        level=level,
        lesson_info=lesson_info,
    )
    st.warning(
        "We couldn't find your student code. Try switching to another tab and back, "
        "or logging out and back in if it doesn't load automatically. You can "
        "manually enter your student code in the field below. If it still isn't "
        f"recognized, [email us]({mailto_link}) and paste your work so the administration can assist."
    )


def render_resubmit_email_cta(
    *,
    lesson_info: Optional[Dict[str, Any]] = None,
    student_name: str = "",
    student_code: str = "",
) -> None:
    """Display the resubmission email instructions."""

    info = lesson_info or {}
    assignment_day = _safe_str(info.get("day"))
    safe_name = _safe_str(student_name)
    safe_code = _safe_str(student_code)
    if safe_code.lower() == "demo001":
        safe_code = ""

    st.write("**Need to resubmit?**")
    st.text(
        "Click the link below to email learngermanghana@gmail.com with your revised work."
    )

    resubmit_body = (
        "Paste your revised work here.\n\n"
        f"Name: {safe_name}\n"
        f"Student Code: {safe_code}\n"
        f"Assignment number: {assignment_day}"
    )

    resubmit_link = (
        "mailto:learngermanghana@gmail.com"
        "?subject=Assignment%20Resubmission"
        f"&body={_urllib.quote(resubmit_body)}"
    )

    st.markdown(
        f"[Click here to open a pre-filled resubmission email]({resubmit_link})"
    )
    st.text(
        "The email will already include your name, student code, and assignment number—just paste your revised work before sending."
    )
    st.text(
        "If the link doesn't open an email app, copy the address and send your resubmission manually."
    )


def _update_student_code_session_state(code: str) -> Dict[str, Any]:
    """Persist ``code`` into session state and return the updated row."""

    clean_code = _safe_lower(code)
    if not clean_code or clean_code == "demo001":
        return st.session_state.get("student_row") or {}

    current_row = st.session_state.get("student_row")
    try:
        updated_row = dict(current_row or {})
    except Exception:
        updated_row = {}

    updated_row["StudentCode"] = clean_code
    st.session_state["student_row"] = updated_row
    st.session_state["student_code"] = clean_code
    st.session_state["assignment_student_code_input"] = clean_code
    st.session_state["submit_student_code_input"] = clean_code

    return updated_row


def _recover_student_code(
    *,
    lesson_key: str = "",
    draft_text: Optional[str] = None,
) -> str:
    """Attempt to resolve the active student's code from session or drafts."""

    session_row = st.session_state.get("student_row") or {}
    candidates = [
        session_row.get("StudentCode"),
        st.session_state.get("student_code"),
        st.session_state.get("assignment_student_code_input"),
        st.session_state.get("submit_student_code_input"),
    ]

    for candidate in candidates:
        candidate_code = _safe_upper(candidate)
        if candidate_code and candidate_code.lower() != "demo001":
            return candidate_code

    if not lesson_key:
        return ""

    cache = st.session_state.setdefault("_student_code_lookup_cache", {})
    text_token = ""
    if draft_text:
        try:
            text_token = hashlib.sha1(str(draft_text).encode("utf-8")).hexdigest()
        except Exception:
            text_token = ""
    cache_key = f"{lesson_key}::{text_token}" if lesson_key else lesson_key
    cached = cache.get(cache_key)
    if cached:
        return _safe_upper(cached)

    recovered = recover_student_code_from_drafts(
        lesson_key,
        draft_text=draft_text if draft_text else None,
    )
    if recovered:
        cache[cache_key] = recovered
        return _safe_upper(recovered)

    return ""


def hide_sidebar() -> None:
    """Hide Streamlit's sidebar for pages where it isn't needed."""
    st.markdown(
        """
        <style>
            div[data-testid="stSidebarNav"] {display: none;}
            div[data-testid="stSidebarHeader"] {display: none;}
        </style>
        """,
        unsafe_allow_html=True,
    )


# ------------------------------------------------------------------------------
# Shared umlaut helper
# ------------------------------------------------------------------------------


def render_umlaut_pad(widget_key: str, *, context: str, disabled: bool = False) -> None:
    """Render a lightweight umlaut reminder next to a textarea."""

    markdown_fn = getattr(st, "markdown", None)
    if markdown_fn is None:
        return

    note = "_Umlaut keys (copy or paste):_"
    parts = ["**Umlaute:**", "ä", "ö", "ü", "ß"]
    separator = " · "
    buttons = f"{parts[0]} {separator.join(parts[1:])}"
    body = f"{note} {buttons}"

    if disabled:
        body = f"{body} _(copy only)_"

    markdown_fn(body)


def render_message(role: str, text: str) -> None:
    """Render a chat bubble for ``role`` (``"assistant"`` or ``"user"``)."""

    markdown_fn = getattr(st, "markdown", None)
    if markdown_fn is None:
        return

    role_key = (role or "").strip().lower()

    content = _safe_str(text)
    content_html = html.escape(content).replace("\n", "<br>")

    if role_key == "assistant":
        label_html = "<div class='bubble-wrap'><div class='lbl-a'>Herr Felix</div></div>"
        bubble_html = f"<div class='bubble-a'>{content_html}</div>"
    else:
        student_label = (
            _safe_str(
                st.session_state.get("student_display_name")
                or st.session_state.get("student_name")
                or st.session_state.get("student_code")
                or "You",
            )
            or "You"
        )
        label_html = (
            "<div class='bubble-wrap'><div class='lbl-u'>"
            f"{html.escape(student_label)}"
            "</div></div>"
        )
        bubble_html = f"<div class='bubble-u'>{content_html}</div>"

    markdown_fn(label_html, unsafe_allow_html=True)
    markdown_fn(bubble_html, unsafe_allow_html=True)



def ensure_student_row(*, stop_if_missing: bool = False) -> Dict[str, Any]:
    """Ensure ``st.session_state['student_row']`` is populated from the roster."""

    student_row = st.session_state.get("student_row", {}) or {}
    if not student_row:
        load_student_data_fn = globals().get("load_student_data")
        if load_student_data_fn is None:

            def load_student_data_fn():
                return pd.DataFrame(columns=["StudentCode"])

        student_code = (st.session_state.get("student_code", "") or "").strip().lower()
        logging.debug("Attempting roster lookup for student code '%s'", student_code)

        try:
            df_students = load_student_data_fn()
        except Exception as exc:
            logging.debug(
                "Roster fetch raised while looking for student code '%s': %s",
                student_code,
                exc,
            )
            df_students = pd.DataFrame(columns=["StudentCode"])

        if df_students is None:
            logging.debug(
                "Roster fetch returned no data while looking for student code '%s'",
                student_code,
            )
            df_students = pd.DataFrame(columns=["StudentCode"])
        else:
            try:
                row_count = len(df_students)
            except Exception:
                row_count = 0
            logging.debug(
                "Roster fetch returned %d rows while looking for student code '%s'",
                row_count,
                student_code,
            )

        if not student_code:
            logging.debug("No student code available in session for roster lookup")
        elif isinstance(df_students, pd.DataFrame) and not df_students.empty and "StudentCode" in df_students.columns:
            try:
                matches = df_students[
                    df_students["StudentCode"].astype(str).str.strip().str.lower()
                    == student_code
                ]
                match_count = int(matches.shape[0])
                logging.debug(
                    "Roster lookup for student code '%s' found=%s (matches=%d)",
                    student_code,
                    match_count > 0,
                    match_count,
                )
                if match_count > 0:
                    student_row = matches.iloc[0].to_dict()
                    st.session_state["student_row"] = student_row
            except Exception as exc:
                logging.debug(
                    "Roster lookup failed for student code '%s': %s",
                    student_code,
                    exc,
                )
        else:
            logging.debug(
                "Roster lookup for student code '%s' skipped because roster is empty or missing StudentCode column",
                student_code,
            )

    student_row = st.session_state.get("student_row", {}) or {}
    if stop_if_missing and not student_row:
        st.info("🚩 No student selected.")
        st.stop()

    return student_row


# Ensure the latest lesson schedule is loaded
if "level_schedules_initialized" not in st.session_state:
    refresh_level_schedules()
    st.session_state["level_schedules_initialized"] = True


# --- Falowen modules ---
from falowen.email_utils import send_reset_email, build_gas_reset_link
import falowen.sessions as _falowen_sessions
from falowen.sessions import (
    db,
    create_session_token,
    destroy_session_token,
    api_post,
)
from src.utils.falowen_imports import load_falowen_db
from src.contracts import (
    is_contract_expired,
)
from src.services.contracts import contract_active
from src.utils.currency import format_cedis
from src.utils.toasts import toast_ok, refresh_with_toast, toast_once
from src.firestore_utils import (
    _draft_doc_ref,
    load_chat_draft_from_db,
    load_draft_from_db,
    load_draft_meta_from_db,
    save_chat_draft_to_db,
    save_draft_to_db,
    save_ai_response,
    fetch_attendance_summary,
    load_student_profile,
    save_student_profile,
    recover_student_code_from_drafts,
    fetch_active_typists,
    set_typing_indicator,
)
from src.draft_management import (
    _draft_state_keys,
    clear_draft_after_post,
    initialize_draft_state,
    save_now,
    autosave_maybe,
    reset_local_draft_state,
    load_notes_from_db,
    save_notes_to_db,
    autosave_learning_note,
    on_cb_subtab_change,
)
from src.falowen.chat_core import (
    back_step,
    render_chat_stage,
    reset_falowen_chat_flow,
)
from src.firestore_helpers import (
    lesson_key_build,
    lock_id,
    has_existing_submission,
    acquire_lock,
    is_locked,
    resolve_current_content,
    fetch_latest,
)
from src.attendance_utils import load_attendance_records
import src.ui_components as _ui_components
from src.ui_components import (
    render_assignment_reminder,
    render_link,
    render_vocab_lookup,
)

_falowen_db = load_falowen_db()
SCHREIBEN_DAILY_LIMIT = _falowen_db.SCHREIBEN_DAILY_LIMIT
inc_sprechen_usage = _falowen_db.inc_sprechen_usage

prepare_audio_url = getattr(_ui_components, "prepare_audio_url", lambda url: url)
render_audio_player = getattr(_ui_components, "render_audio_player", lambda *a, **k: None)
from src.stats import (
    get_student_level,
    save_vocab_attempt,
    vocab_attempt_exists,
)
from src.stats_ui import render_vocab_stats, render_schreiben_stats
from src.schreiben import (
    highlight_feedback,
    get_schreiben_usage,
    inc_schreiben_usage,
    save_letter_coach_progress,
    load_letter_coach_progress,
    save_letter_coach_draft,
    load_letter_coach_draft,
    clear_letter_coach_draft,
    get_level_from_code,
)
from src.ui.auth import (
    render_signup_form,
    render_login_form,
    render_forgot_password_panel,
    render_returning_login_area,
    render_signup_request_banner,
    render_google_oauth,
    render_returning_login_form,
)
from src.ui.auth import renew_session_if_needed
from src.ui.login import render_falowen_login
from src.services.vocab import VOCAB_LISTS, AUDIO_URLS, get_audio_url
from src.schreiben import (
    update_schreiben_stats,
    get_schreiben_stats,
    save_submission,
    save_schreiben_feedback,
    load_schreiben_feedback,
    delete_schreiben_feedback,
)
from src.ui_helpers import (
    qp_get,
    qp_clear,
    seed_falowen_state_from_qp,
    highlight_terms,
    filter_matches,
)
from src.auth import (
    persist_session_client,
    reset_password_page,
)
from src.assignment_ui import (
    load_assignment_scores,
    render_results_and_resources_tab,
    get_assignment_summary,
    select_best_assignment_attempts,
)
from src.session_management import (
    bootstrap_state,
    determine_level,
    ensure_student_level,
    bootstrap_session_from_qp,
)
from src.sentence_bank import SENTENCE_BANK
from src.config import SB_SESSION_TARGET
from src.data_loading import load_student_data
from src.youtube import (
    get_playlist_ids_for_level,
    fetch_youtube_playlist_videos,
)
from src.ui_widgets import (
    render_google_brand_button_once,
    render_announcements_once,
)
from src.logout import do_logout
from src.pdf_handling import (
    extract_text_from_pdf,
    generate_notes_pdf,
    generate_single_note_pdf,
    generate_chat_pdf,
)
from src.sentence_builder import render_sentence_builder

# ------------------------------------------------------------------------------
# Google OAuth (Gmail sign-in) — single-source, no duplicate buttons
# ------------------------------------------------------------------------------
GOOGLE_CLIENT_ID     = st.secrets.get("GOOGLE_CLIENT_ID", "180240695202-3v682khdfarmq9io9mp0169skl79hr8c.apps.googleusercontent.com")
GOOGLE_CLIENT_SECRET = st.secrets.get("GOOGLE_CLIENT_SECRET", "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")
REDIRECT_URI         = st.secrets.get("GOOGLE_REDIRECT_URI", "https://www.falowen.app/")

# Mapping of CEFR levels to teacher codes that should receive admin rights.
# Extend this dictionary as new levels or teachers are added.
ADMINS_BY_LEVEL = {
    "A1": {"felixa177", "felixa1"},
    "A2": {"felixa2"},
    "B1": {"felixb1"},
}

# Study tips shown on the dashboard for each CEFR level.
LEVEL_TIPS = {
    "A1": "Focus on everyday phrases and practice listening with simple dialogues.",
    "A2": "Build vocabulary around daily routines and start speaking in longer sentences.",
    "B1": "Read short articles to boost comprehension and keep a diary in German.",
    "B2": "Engage with podcasts or news to refine your listening and expand vocabulary.",
    "C1": "Discuss complex topics in German and review grammar nuances regularly.",
}

# Exam preparation advice for each CEFR level.
EXAM_ADVICE = {
    "A1": "Focus on listening comprehension to build a strong foundation.",
    "A2": "Review core grammar and practice everyday conversations.",
    "B1": "Work on grammar accuracy and write short essays.",
    "B2": "Engage with longer authentic texts and audio.",
    "C1": "Hone precision in complex discussions and essays.",
}





def inject_notice_css():
    from src.ui.login import inject_notice_css as _inject_css
    _inject_css()

# Legacy hero rendering moved to src.login_ui

# ------------------------------------------------------------------------------
# Sign up / Login / Forgot password
# ------------------------------------------------------------------------------


def calc_blog_height(num_posts: int) -> int:
    """Return the container height needed for the desktop blog card grid.

    The blog cards render inside a CSS grid defined as
    ``repeat(auto-fill, minmax(240px, 1fr))`` with a ``16px`` gap and a maximum
    width of ``1120px``.  On wide viewports this produces four columns, so each
    row requires space for four cards plus the gaps between them.  Each card is
    approximately ``312px`` tall.  The function computes the minimum height
    necessary to show all posts without leaving excessive blank space.

    Parameters
    ----------
    num_posts: int
        Number of blog posts to display.

    Returns
    -------
    int
        Height in pixels for the blog card container.
    """

    CARD_HEIGHT = 312
    ROW_GAP = 16
    CARD_MIN_WIDTH = 240
    GRID_MAX_WIDTH = 1120

    # Desktop layout shows at most four columns: floor((max_width + gap) /
    # (min_width + gap)) = floor((1120 + 16) / (240 + 16)) = 4.
    CARDS_PER_ROW = max(
        1, math.floor((GRID_MAX_WIDTH + ROW_GAP) / (CARD_MIN_WIDTH + ROW_GAP))
    )

    if num_posts <= 0:
        return 0

    rows = math.ceil(num_posts / CARDS_PER_ROW)
    return CARD_HEIGHT * rows + ROW_GAP * (rows - 1)


def login_page():
    if st.session_state.get("logged_in"):
        try:
            render_logged_in_topbar()
        except Exception:
            pass
        try:
            ensure_student_level()
        except Exception:
            pass
        return

    hide_fn = globals().get("hide_sidebar")
    if callable(hide_fn):
        hide_fn()

    try:
        renew_session_if_needed()
    except Exception:
        pass

    # 1) Get Google auth URL (also completes flow if ?code=...)
    auth_url = render_google_oauth(return_url=True) or ""

    # 2) Branded hero (Google button suppressed inside the template)
    render_falowen_login(auth_url, show_google_in_hero=False)
    st.divider()

    # 3) Returning user section (Google CTA below the form)
    login_success = render_returning_login_area()
    render_google_brand_button_once(auth_url, center=True)

    # Guard: only schedule the post-login rerun once, and clear URL params first
    def _run_once(key: str) -> bool:
        if st.session_state.get(key):
            return False
        st.session_state[key] = True
        return True

    if login_success and _run_once("post_login_rerun"):
        try:
            # Remove any lingering query params like ?code=... from OAuth/deeplinks
            for k in list(st.query_params.keys()):
                st.query_params[k] = ""
        except Exception:
            pass
        st.session_state["need_rerun"] = True

    # 4) Explanation banner + tabs (keep your existing content below)
    render_signup_request_banner()
    tab2, tab3 = st.tabs(["🧾 Sign Up (Approved)", "📝 Request Access"])
    with tab2:
        render_signup_form()
    with tab3:
        st.markdown(
            """
            <div class="page-wrap" style="text-align:center; margin-top:8px;">
              <a href="https://docs.google.com/forms/d/e/1FAIpQLSenGQa9RnK9IgHbAn1I9rSbWfxnztEUcSjV0H-VFLT-jkoZHA/viewform?usp=header" 
                 target="_blank" rel="noopener">
                <button style="background:#1f2d7a; color:white; padding:10px 20px; border:none; border-radius:8px; cursor:pointer;">
                  📝 Open Request Access Form
                </button>
              </a>
              <div style="color:#64748b; font-size:.95rem; margin-top:6px;">
                We’ll email you once your account is ready.
              </div>
            </div>
            """,
            unsafe_allow_html=True,
        )


    # (Optional) help/links/steps/footer blocks can follow...


    # Help + quick contacts
    st.markdown("""
    <div class="page-wrap">
      <div class="help-contact-box" aria-label="Help and contact options" style="text-align:center;">
        <b>❓ Need help or access?</b><br>
        <a href="https://api.whatsapp.com/send?phone=233205706589" target="_blank" rel="noopener">📱 WhatsApp us</a>
        &nbsp;|&nbsp;
        <a href="mailto:learngermanghana@gmail.com" target="_blank" rel="noopener">✉️ Email</a>
      </div>
    </div>
    """, unsafe_allow_html=True)

    # Steps (1-2-3)
    st.markdown("---")
    LOGIN_IMG_URL      = "https://i.imgur.com/pFQ5BIn.png"
    COURSEBOOK_IMG_URL = "https://i.imgur.com/pqXoqSC.png"
    RESULTS_IMG_URL    = "https://i.imgur.com/uiIPKUT.png"
    c1, c2, c3 = st.columns(3)
    with c1:
        st.markdown(f"""
        <img src="{LOGIN_IMG_URL}" alt="Login screenshot"
             style="width:100%; height:220px; object-fit:cover; border-radius:12px; pointer-events:none; user-select:none;">
        <div style="height:8px;"></div>
        <h3 style="margin:0 0 4px 0;">1️⃣ Sign in</h3>
        <p style="margin:0;">Use your <b>student code or email</b> — or the <b>Google</b> button above.</p>
        """, unsafe_allow_html=True)
    with c2:
        st.markdown(f"""
        <img src="{COURSEBOOK_IMG_URL}" alt="Course Book screenshot"
             style="width:100%; height:220px; object-fit:cover; border-radius:12px; pointer-events:none; user-select:none;">
        <div style="height:8px;"></div>
        <h3 style="margin:0 0 4px 0;">2️⃣ Learn & submit</h3>
        <p style="margin:0;">Watch lessons, practice vocab, and <b>submit assignments</b> in the Course Book.</p>
        """, unsafe_allow_html=True)
    with c3:
        st.markdown(f"""
        <img src="{RESULTS_IMG_URL}" alt="Results screenshot"
             style="width:100%; height:220px; object-fit:cover; border-radius:12px; pointer-events:none; user-select:none;">
        <div style="height:8px;"></div>
        <h3 style="margin:0 0 4px 0;">3️⃣ Get results</h3>
        <p style="margin:0;">You’ll get an <b>email when marked</b>. Check <b>Results & Resources</b> for feedback.</p>
        """, unsafe_allow_html=True)

    # Blog posts / announcements
    blog_posts = fetch_blog_feed(limit=4)[:4]
    if blog_posts:
        st.markdown("---")
        st.markdown(
            '<h2 style="text-align:center;">Falowen Blog</h2>',
            unsafe_allow_html=True,
        )
        height = globals().get("calc_blog_height", lambda n: 380)(len(blog_posts))
        render_blog_cards(blog_posts, height=height)
        st.markdown(
            '<div style="text-align:center;margin:4px 0 0;">'
            '<a href="https://blog.falowen.app/" target="_blank" rel="noopener">Read more</a>'
            '</div>',
            unsafe_allow_html=True,
        )

    # Footer links
    st.markdown("""
    <div class="page-wrap" style="text-align:center; margin:12px 0;">
      <a href="https://www.learngermanghana.com/tutors"           target="_blank" rel="noopener">👩‍🏫 Tutors</a>
      &nbsp;|&nbsp;
      <a href="https://www.learngermanghana.com/upcoming-classes" target="_blank" rel="noopener">🗓️ Upcoming Classes</a>
      &nbsp;|&nbsp;
      <a href="https://register.falowen.app/#privacy-policy"      target="_blank" rel="noopener">🔒 Privacy</a>
      &nbsp;|&nbsp;
      <a href="https://register.falowen.app/#terms-of-service"    target="_blank" rel="noopener">📜 Terms</a>
      &nbsp;|&nbsp;
      <a href="https://www.learngermanghana.com/contact-us"       target="_blank" rel="noopener">✉️ Contact</a>
      &nbsp;|&nbsp;
      <a href="https://register.falowen.app"                      target="_blank" rel="noopener">📝 Register</a>
      &nbsp;|&nbsp;
      <a href="https://register.falowen.app/#about-us"            target="_blank" rel="noopener">ℹ️ About Us</a>
      &nbsp;|&nbsp;
      <a href="https://script.google.com/macros/s/AKfycbwXrfiuKl65Va_B2Nr4dFnyLRW5z6wT5kAbCj6cNl1JxdOzWVKT_ZMwdh2pN_dbdFoy/exec" target="_blank" rel="noopener">🗑️ Delete Account</a>
      &nbsp;|&nbsp;
      <a href="https://blog.falowen.app/"                         target="_blank" rel="noopener">📰 Blog</a>
    </div>
    """, unsafe_allow_html=True)

    from datetime import UTC, datetime as _dt_now
    st.markdown(f"""
    <div class="page-wrap" style="text-align:center;color:#64748b; margin-bottom:16px;">
      © {_dt_now.now(UTC).year} Learn Language Education Academy • Accra, Ghana<br>
      Need help? <a href="mailto:learngermanghana@gmail.com">Email</a> •
      <a href="https://api.whatsapp.com/send?phone=233205706589" target="_blank" rel="noopener">WhatsApp</a>
    </div>
    """, unsafe_allow_html=True)


def render_logged_in_topbar():
    name  = st.session_state.get("student_name", "")
    level = st.session_state.get("student_level", "—")
    code  = st.session_state.get("student_code", "—")

    st.markdown(
        """
        <style>
          .dash-topwrap{
            background:#f5f9ff;
            border:1px solid rgba(30,64,175,.12);
            border-radius:14px;
            padding:14px 16px;
            margin:4px 0 10px 0;
            box-shadow:0 6px 14px rgba(2,6,23,.06);
          }
          .dash-title{ font-size:1.55rem; font-weight:900; color:#19213a; margin:0 0 4px 0; }
          .dash-sub{ color:#475569; font-size:.95rem; }
          div[data-testid="stButton"] > button[kind="primary"]{
            background:#1f2d7a; border:1px solid #1b2a6e; border-radius:10px; font-weight:700;
            box-shadow:0 4px 10px rgba(31,45,122,.18);
          }
          div[data-testid="stButton"] > button[kind="primary"]:hover{ filter:brightness(1.05); }
        </style>
        """,
        unsafe_allow_html=True
    )
    top = st.container()
    with top:
        c1, c2 = st.columns([1, 0.18])
        with c1:
            st.markdown(
                f"""
                <div class="dash-topwrap">
                  <div class="dash-title">👋 Welcome, {name}</div>
                  <div class="dash-sub">Level: {level} · Code: {code}</div>
                </div>
                """,
                unsafe_allow_html=True
            )
        with c2:
            st.button(
                "Log out",
                key="logout_global",
                type="primary",
                width="stretch",
                on_click=do_logout,
            )

    level_key = (level or "").strip().upper()
    tip = LEVEL_TIPS.get(level_key, "Keep practicing and immerse yourself daily.")
    st.info(tip)


# ------------------------------------------------------------------------------
# Level-aware welcome video (YouTube) used in the sidebar (IDs can be added later)
# ------------------------------------------------------------------------------
def dashboard_page():
    """Render the dashboard for logged-in users."""
    if not st.session_state.get("logged_in"):
        login_page()
        return
    try:
        render_logged_in_topbar()
    except Exception:
        pass
    try:
        ensure_student_level()
    except Exception:
        pass


def render_level_welcome_video(level: str | None):
    level = (level or "").strip().upper() or "A1"
    YT_WELCOME = {"A1":"", "A2":"", "B1":"", "B2":"", "C1":"", "C2":""}  # fill IDs later
    vid = YT_WELCOME.get(level) or ""
    if not vid:
        st.info(f"No welcome video added yet for {level}. Check back soon!")
        return
    components.html(
        f"""
        <div style="position:relative;padding-bottom:56.25%;height:0;overflow:hidden;border-radius:12px;
                    box-shadow:0 4px 12px rgba(0,0,0,.08);">
          <iframe
            src="https://www.youtube.com/embed/{vid}"
            title="Welcome • {level}"
            frameborder="0"
            allow="accelerometer; autoplay; clipboard-write; encrypted-media; gyroscope; picture-in-picture; web-share"
            allowfullscreen
            style="position:absolute;top:0;left:0;width:100%;height:100%;border:0;border-radius:12px;">
          </iframe>
        </div>
        """, height=320, scrolling=False
    )


# ------------------------------------------------------------------------------
# Sidebar (publish-ready)
# ------------------------------------------------------------------------------
def render_sidebar_published():
    def _qp_set_safe(**kwargs):
        if "_qp_set" in globals():
            try: _qp_set(**kwargs); return
            except Exception: pass
        try:
            for k, v in kwargs.items():
                st.query_params[k] = "" if v is None else str(v)
        except Exception:
            pass

    def _go(tab_name: str):
        st.session_state["nav_sel"] = tab_name
        st.session_state["main_tab_select"] = tab_name
        _qp_set_safe(tab=tab_name)
        if tab_name != "Chat • Grammar • Exams":
            st.session_state.pop("_chat_focus_tab", None)
        st.session_state["need_rerun"] = True

    def _go_chat_main():
        st.session_state["_chat_focus_tab"] = None
        _go("Chat • Grammar • Exams")

    def _go_zoom_class():
        st.session_state["nav_sel"] = "My Course"
        st.session_state["main_tab_select"] = "My Course"
        st.session_state["coursebook_subtab"] = "🧑‍🏫 Classroom"
        st.session_state["classroom_page"] = "Join on Zoom"
        _qp_set_safe(tab="My Course")
        st.session_state.pop("_chat_focus_tab", None)
        st.session_state["need_rerun"] = True

    def _go_post_qna():
        st.session_state["nav_sel"] = "My Course"
        st.session_state["main_tab_select"] = "My Course"
        st.session_state["coursebook_subtab"] = "🧑‍🏫 Classroom"
        st.session_state["classroom_page"] = "Class Notes & Q&A"
        _qp_set_safe(tab="My Course")
        st.session_state.pop("_chat_focus_tab", None)
        st.session_state["need_rerun"] = True

    def _go_course_submit():
        st.session_state["nav_sel"] = "My Course"
        st.session_state["main_tab_select"] = "My Course"
        st.session_state["coursebook_subtab"] = "📘 Course Book"
        st.session_state["coursebook_page"] = "Submit"
        st.session_state["coursebook_prev_page"] = "Submit"
        _qp_set_safe(tab="My Course")
        st.session_state.pop("_chat_focus_tab", None)
        st.session_state["need_rerun"] = True
    if st.session_state.get("logged_in", False):
        st.sidebar.markdown("## Quick access")
        st.sidebar.button("🏠 Dashboard",                width="stretch", on_click=_go, args=("Dashboard",))
        st.sidebar.button("📈 My Course",                width="stretch", on_click=_go, args=("My Course",))
        st.sidebar.button("📊 Results & Resources",      width="stretch", on_click=_go, args=("My Results and Resources",))
        st.sidebar.button("🗣️ Chat • Grammar • Exams", width="stretch", on_click=_go_chat_main)
        st.sidebar.button("✅ Submit Assignment",       width="stretch", on_click=_go_course_submit)
        st.sidebar.button("✍️ Schreiben Trainer",        width="stretch", on_click=_go, args=("Schreiben Trainer",))
        st.sidebar.button("🎥 Join on Zoom",             width="stretch", on_click=_go_zoom_class)
        st.sidebar.button("❓ Class Notes & Q&A",         width="stretch", on_click=_go_post_qna)
        st.sidebar.divider()

        st.sidebar.markdown("## Our Socials")
        st.sidebar.markdown(
            """
- 📸 [Instagram](https://www.instagram.com/lleaghana/)
- ▶️ [YouTube](https://www.youtube.com/@LLEAGhana)
- 🎵 [TikTok](https://www.tiktok.com/@lleaghana)
- 💼 [LinkedIn](https://www.linkedin.com/in/lleaghana/)
            """
        )
        st.sidebar.divider()

        st.sidebar.markdown("## How-to & tips")
        first_time_quick_guide = not st.session_state.get("_seen_quick_guide", False)
        if first_time_quick_guide:
            st.toast("👋 New here? Peek at the Quick guide in the sidebar to get started!")
        with st.sidebar.expander("📚 Quick guide", expanded=first_time_quick_guide):
            st.markdown(
                """
                - **Submit work:** My Course → Submit → **Confirm & Submit** (locks after submission).
                - **Check feedback:** **Results & Resources** shows marks, comments, downloads.
                - **Practice speaking:** **Tools → Sprechen** for instant pronunciation feedback.
                - **Build vocab:** **Schreiben Trainer → Vocab Trainer** for daily words & review cycles.
                - **Track progress:** **Dashboard** shows streaks, next lesson, and missed items.
                """
            )
        if first_time_quick_guide:
            st.session_state["_seen_quick_guide"] = True

        with st.sidebar.expander("🧭 Dashboard tabs, explained", expanded=False):
            st.markdown(
                """
                - **Dashboard:** Overview (streak, next lesson, missed, leaderboard, new posts).
                - **My Course:** Lessons, materials, and submission flow.
                - **Results & Resources:** Marks, feedback, downloadable resources.
                - **Chat • Grammar • Exams:** Guided conversation practice plus instant pronunciation feedback.
                - **Schreiben Trainer:** Structured writing with iterative feedback and 📚 vocab practice tools.
                """
            )

        with st.sidebar.expander("🔔 Telegram notifications", expanded=False):
            st.markdown(
                """
- Open the Falowen bot link or search for `@falowenbot`, tap **Start**, then type your student code (e.g. `kwame202`)
- To deactivate: send `/stop`
                """
            )

        st.sidebar.divider()

        st.sidebar.markdown("## Support")
        st.sidebar.markdown(
            """
- 📱 [WhatsApp](https://api.whatsapp.com/send?phone=233205706589)
- ✉️ [Email](mailto:learngermanghana@gmail.com)
- 🐞 [Report an issue](mailto:learngermanghana@gmail.com?subject=Falowen%20Bug%20Report)
            """
        )

        st.sidebar.markdown("## Resources")
        st.sidebar.markdown(
            """
- 👩‍🏫 [Tutors](https://www.learngermanghana.com/tutors)
- 🗓️ [Upcoming Classes](https://www.learngermanghana.com/upcoming-classes)
- 📰 [Blog](https://blog.falowen.app)
- ✉️ [About Us](https://register.falowen.app/#about-us)
            """
        )



# ------------------------------------------------------------------------------
# OpenAI (used elsewhere in app)
# ------------------------------------------------------------------------------
OPENAI_API_KEY = st.secrets.get("OPENAI_API_KEY") or os.getenv("OPENAI_API_KEY")
if not OPENAI_API_KEY:
    st.error("Missing OpenAI API key. Please add OPENAI_API_KEY in Streamlit secrets.")
    raise RuntimeError("Missing OpenAI API key")
os.environ["OPENAI_API_KEY"] = OPENAI_API_KEY
client = OpenAI(api_key=OPENAI_API_KEY)


def apply_profile_ai_correction(about_key: str) -> None:
    """Use OpenAI to correct and enhance the user's profile biography."""
    current_text = st.session_state.get(about_key, "")
    if not current_text.strip():
        return
    if not OPENAI_API_KEY:
        st.error("Missing OpenAI API key.")
        return
    try:
        resp = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are a helpful assistant that corrects and enhances a student's biography. "
                        "Return only the improved biography."
                    ),
                },
                {"role": "user", "content": current_text},
            ],
            temperature=0,
            max_tokens=300,
        )
        ai_text = (resp.choices[0].message.content or "").strip()
        if ai_text:
            st.session_state[about_key] = ai_text
    except Exception as e:
        logging.exception("Profile AI correction error")
        st.error(f"AI correction failed: {e}")


def apply_status_ai_correction(text: str) -> Tuple[str, str]:
    """Return an AI-improved version of *text* and a brief explanation."""
    if not text.strip():
        return text, ""
    try:
        resp = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are a helpful assistant that improves a student's forum post. "
                        "Return a JSON object with keys 'improved' for the corrected post "
                        "and 'explanation' for a short explanation of the changes."
                    ),
                },
                {"role": "user", "content": text},
            ],
            temperature=0,
            max_tokens=400,
        )
        raw = resp.choices[0].message.content or ""
        try:
            data = json.loads(raw)
            improved = (data.get("improved") or "").strip()
            explanation = (data.get("explanation") or "").strip()
        except Exception:
            improved = raw.strip()
            explanation = ""
        return improved, explanation
    except Exception as e:
        logging.exception("Status AI correction error")
        st.error(f"AI correction failed: {e}")
        return text, ""


def apply_note_ai_correction(text: str) -> Tuple[str, str]:
    """Return an AI-improved version of a learning note and an explanation."""
    if not text.strip():
        return text, ""
    try:
        resp = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are a helpful assistant that improves a student's learning note. "
                        "Return a JSON object with keys 'improved' for the corrected note "
                        "and 'explanation' for a short explanation of the changes."
                    ),
                },
                {"role": "user", "content": text},
            ],
            temperature=0,
            max_tokens=400,
        )
        raw = resp.choices[0].message.content or ""
        try:
            data = json.loads(raw)
            improved = (data.get("improved") or "").strip()
            explanation = (data.get("explanation") or "").strip()
        except Exception:
            improved = raw.strip()
            explanation = ""
        return improved, explanation
    except Exception as e:
        logging.exception("Learning note AI correction error")
        st.error(f"AI correction failed: {e}")
        return text, ""


def diff_with_markers(original: str, corrected: str) -> str:
    """Generate HTML diff using <mark> tags for additions and deletions."""
    diff_lines = difflib.unified_diff(
        original.splitlines(),
        corrected.splitlines(),
        lineterm="",
    )
    html_lines = ["<pre>"]
    for line in diff_lines:
        if line.startswith(("---", "+++", "@@")):
            continue
        if line.startswith("+"):
            html_lines.append(
                f"<mark style='background-color:#d4fcbc'>+ {html.escape(line[1:])}</mark>"
            )
        elif line.startswith("-"):
            html_lines.append(
                f"<mark style='background-color:#ffbdbd'>- {html.escape(line[1:])}</mark>"
            )
        else:
            html_lines.append(html.escape(line))
    html_lines.append("</pre>")
    return "\n".join(html_lines)


# ------------------------------------------------------------------------------
# Seed state from query params / restore session / reset-link path / go to login
# ------------------------------------------------------------------------------
bootstrap_state()
seed_falowen_state_from_qp()
bootstrap_session_from_qp()

# If visiting with password-reset token
if not st.session_state.get("logged_in", False):
    tok = st.query_params.get("token")
    if isinstance(tok, list):
        tok = tok[0] if tok else None
    if tok:
        reset_password_page(tok)
        st.stop()

# Gate
if not st.session_state.get("logged_in", False):
    login_page()
    if not st.session_state.get("logged_in", False):
        st.stop()

# ==================== LOGGED IN ====================
# Ensure the roster row is available for downstream tabs
ensure_student_row()

# Show header immediately after login on every page
render_logged_in_topbar()

# Theme bits (chips etc.)
inject_notice_css()

# Sidebar (no logout; logout lives in the header)
render_sidebar_published()

# Falowen blog updates (render once)
new_posts = fetch_blog_feed()

st.markdown("---")
st.markdown("**You’re logged in.** Continue to your lessons and tools from the navigation.")



# =========================================================
# ============== Data loaders & helpers ===================
# =========================================================
@st.cache_data(ttl=43200)
def _load_full_vocab_sheet_cached():
    SHEET_ID = "1I1yAnqzSh3DPjwWRh9cdRSfzNSPsi7o4r5Taj9Y36NU"
    csv_url = f"https://docs.google.com/spreadsheets/d/{SHEET_ID}/export?format=csv&gid=0"
    try:
        resp = requests.get(csv_url, timeout=8)
        resp.raise_for_status()
        df = pd.read_csv(io.StringIO(resp.text), dtype=str)
    except requests.RequestException as e:
        st.error(f"Could not load vocab sheet: {e}")
        return pd.DataFrame(columns=["level", "german", "english", "example"])
    except Exception:
        st.error("Could not load vocab sheet.")
        return pd.DataFrame(columns=["level", "german", "english", "example"])
    df.columns = df.columns.str.strip().str.lower()

    def _match(colnames, *cands):
        s = set(colnames)
        for c in cands:
            if c in s: return c
        for c in colnames:
            if any(c.startswith(x) for x in cands): return c
        return None

    col_level   = _match(df.columns, "level")
    col_german  = _match(df.columns, "german", "de", "word", "wort")
    col_english = _match(df.columns, "english", "en", "meaning", "translation")
    col_example = _match(df.columns, "example", "sentence", "usage")
    if not (col_level and col_german and col_english):
        return pd.DataFrame(columns=["level", "german", "english", "example"])

    rename = {col_level:"level", col_german:"german", col_english:"english"}
    if col_example: rename[col_example] = "example"
    df = df.rename(columns=rename)
    if "example" not in df.columns: df["example"] = ""
    for c in ["level","german","english","example"]:
        df[c] = df[c].astype(str).str.strip()
    df = df[df["level"].notna() & (df["level"] != "")]
    df["level"] = df["level"].str.upper()
    return df[["level","german","english","example"]]

def load_full_vocab_sheet():
    """Return full vocab sheet DataFrame from session state or cache."""
    if "full_vocab_df" not in st.session_state:
        st.session_state["full_vocab_df"] = _load_full_vocab_sheet_cached()
    return st.session_state["full_vocab_df"]

def get_vocab_of_the_day(df: pd.DataFrame, level: str):
    if df is None or df.empty: return None
    if not {"level","german","english","example"}.issubset(df.columns): return None
    lvl = (level or "").upper().strip()
    subset = df[df["level"] == lvl]
    if subset.empty: return None
    idx = date.today().toordinal() % len(subset)
    row = subset.reset_index(drop=True).iloc[idx]
    return {"german": row.get("german",""), "english": row.get("english",""), "example": row.get("example","")}


def render_lesson_language_support(info: MutableMapping[str, Any] | None, level_key: str) -> None:
    """Show quick vocab and grammar reminders for the current lesson."""

    info = info or {}
    grammar_topic = _safe_str(info.get("grammar_topic"))
    vocab_df = load_full_vocab_sheet()
    suggestions = gather_language_support(info, level_key, vocab_df, VOCAB_LISTS)

    if not grammar_topic and not suggestions:
        return

    st.markdown("###### 💡 Quick Language Support")
    if grammar_topic:
        st.caption(f"Grammar focus: {grammar_topic}")

    if suggestions:
        for entry in suggestions:
            german = _safe_str(entry.get("german"))
            english = _safe_str(entry.get("english"))
            example = _safe_str(entry.get("example"))
            bullet = f"**{german}**" if german else ""
            if english:
                bullet = f"{bullet} — {english}" if bullet else english
            if bullet:
                line = f"- {bullet}"
            elif example:
                line = f"- _{example}_"
            else:
                continue
            if example and bullet:
                line += f"  \n  _{example}_"
            st.markdown(line)
    else:
        st.caption("No quick vocabulary matches found for this lesson.")


@st.cache_data(ttl=3600)
def _load_reviews_cached():
    SHEET_ID = "137HANmV9jmMWJEdcA1klqGiP8nYihkDugcIbA-2V1Wc"
    url = f"https://docs.google.com/spreadsheets/d/{SHEET_ID}/gviz/tq?tqx=out:csv&sheet=Sheet1"
    df = pd.read_csv(url)
    df.columns = df.columns.str.strip().str.lower()
    return df

def load_reviews():
    """Return reviews DataFrame cached and stored in session state."""
    if "reviews_df" not in st.session_state:
        st.session_state["reviews_df"] = _load_reviews_cached()
    return st.session_state["reviews_df"]

CONTRACT_DATE_FORMATS = (
    "%Y-%m-%d",
    "%Y-%m-%d %H:%M:%S",
    "%Y-%m-%dT%H:%M:%S",
    "%m/%d/%Y",
    "%d.%m.%y",
    "%d.%m.%Y",
    "%d/%m/%Y",
    "%d-%m-%Y",
)


def _parse_contract_date_value(date_str: Any):
    """Return a naive ``datetime`` for *date_str* or ``None`` if parsing fails."""

    if isinstance(date_str, datetime):
        dt = date_str
        if dt.tzinfo is not None:
            return dt.astimezone(UTC).replace(tzinfo=None)
        return dt

    if hasattr(date_str, "to_pydatetime"):
        try:
            converted = date_str.to_pydatetime()
        except Exception:
            converted = None
        if converted is not None:
            return _parse_contract_date_value(converted)

    text = "" if date_str is None else str(date_str).strip()
    if not text or text.lower() in ("nan", "none"):
        return None

    iso_candidate = text.replace("Z", "+00:00")
    try:
        parsed = datetime.fromisoformat(iso_candidate)
    except ValueError:
        parsed = None
    if parsed is not None:
        if parsed.tzinfo is not None:
            parsed = parsed.astimezone(UTC).replace(tzinfo=None)
        return parsed

    trimmed = text[:-1].strip() if text.endswith("Z") else text
    for fmt in CONTRACT_DATE_FORMATS:
        try:
            return datetime.strptime(trimmed, fmt)
        except ValueError:
            continue

    for sep in (" ", "T"):
        if sep in trimmed:
            base = trimmed.split(sep, 1)[0].strip()
            if base and base != trimmed:
                return _parse_contract_date_value(base)

    return None


def parse_contract_start(date_str: Any):

    """Parse a contract start date across common and ISO timestamp formats."""

    return _parse_contract_date_value(date_str)


def _compute_finish_date_estimates(start_str: Any, total_lessons: Any, parse_start_fn):
    """Return projected completion dates keyed by weekly study frequency."""

    try:
        total = int(total_lessons)
    except (TypeError, ValueError):
        return None

    if total <= 0:
        return None

    parse_fn = parse_start_fn if callable(parse_start_fn) else parse_contract_start
    try:
        parsed = parse_fn(start_str)
    except Exception:
        parsed = None

    if not parsed:
        return None

    if isinstance(parsed, datetime):
        start_date = parsed.date()
    elif isinstance(parsed, date):
        start_date = parsed
    elif hasattr(parsed, "date"):
        start_date = parsed.date()
    else:
        return None

    weeks_three = (total + 2) // 3
    weeks_two = (total + 1) // 2
    weeks_one = total

    return {
        3: start_date + timedelta(weeks=weeks_three),
        2: start_date + timedelta(weeks=weeks_two),
        1: start_date + timedelta(weeks=weeks_one),
    }


def _dict_tts_bytes_de(text: str) -> Optional[bytes]:
    """Return MP3 bytes for German *text* using gTTS.

    On failure, log and return ``None`` instead of raising to avoid crashing the app.
    """
    if not text:
        return None
    try:
        from gtts import gTTS

        buf = io.BytesIO()
        gTTS(text=text, lang="de").write_to_fp(buf)
        return buf.getvalue()
    except Exception as exc:  # pragma: no cover - best effort
        logging.warning("gTTS synthesis failed: %s", exc)
        return None

# ------------------------------- Footer -------------------------------
FOOTER_LINKS = {
    "👩‍🏫 Tutors": "https://www.learngermanghana.com/tutors",
    "🗓️ Upcoming Classes": "https://www.learngermanghana.com/upcoming-classes",
    "🔒 Privacy": "https://register.falowen.app/#privacy-policy",
    "📜 Terms": "https://register.falowen.app/#terms-of-service",
    "✉️ Contact": "https://www.learngermanghana.com/contact-us",
    "📝 Register": "https://register.falowen.app",
    "ℹ️ About Us": "https://register.falowen.app/#about-us",
    "🗑️ Delete Account": "https://script.google.com/macros/s/AKfycbwXrfiuKl65Va_B2Nr4dFnyLRW5z6wT5kAbCj6cNl1JxdOzWVKT_ZMwdh2pN_dbdFoy/exec",
    "📰 Blog": "https://blog.falowen.app/",
}

def render_app_footer(links: dict):
    st.markdown(
        """
        <style>
          .app-footer{ margin-top:18px; padding:16px 14px; border-top:1px solid rgba(148,163,184,.35); color:#475569; }
          .app-footer a{ text-decoration:none; font-weight:700; }
          .app-footer .row{ display:flex; flex-wrap:wrap; gap:14px; }
          @media (max-width:640px){ .app-footer{ padding:14px 10px; } }
        </style>
        """,
        unsafe_allow_html=True
    )
    parts = [f'<a href="{href}" target="_blank">{label}</a>' for label, href in links.items()]
    st.markdown(
        f"""
        <div class="app-footer">
          <div class="row">
            {" | ".join(parts)}
          </div>
          <div style="margin-top:6px;font-size:.9rem;">© 2025 Learn Language Education Academy • Accra, Ghana – Need help? Email • WhatsApp</div>
        </div>
        """,
        unsafe_allow_html=True
    )




# =========================================================
# ===================== NAV & HELPERS =====================
# =========================================================

# --- Query-param helpers (single API; no experimental mix) ---
if "_qp_get_first" not in globals():
    def _qp_get_first(key: str, default: str = "") -> str:
        """Return first value from st.query_params (new API-safe)."""
        try:
            val = st.query_params.get(key, default)
            if isinstance(val, list):
                return (val[0] if val else default)
            return str(val)
        except Exception:
            return default

if "_qp_set" not in globals():
    def _qp_set(**kwargs):
        """Set URL query params using only the production API."""
        try:
            for k, v in kwargs.items():
                st.query_params[k] = "" if v is None else str(v)
        except Exception:
            # If browser doesn't allow URL changes, just skip
            pass

if "build_course_day_link" not in globals():
    def build_course_day_link(day: int | str, tab: str = "My Course") -> str:
        """Return a link to a specific Course Book day."""
        try:
            day_val = int(day)
        except Exception:
            day_val = day
        try:
            tab_q = _urllib.quote(tab)
        except Exception:
            tab_q = tab
        return f"?tab={tab_q}&day={day_val}"

# --- Nav dropdown (mobile-friendly, simple text) ---
def render_dropdown_nav():
    tabs = [
        "Dashboard",
        "My Course",
        "My Results and Resources",
        "Chat • Grammar • Exams",
        "Schreiben Trainer",
    ]
    icons = {
        "Dashboard": "🏠",
        "My Course": "📈",
        "My Results and Resources": "📊",
        "Chat • Grammar • Exams": "🗣️",
        "Schreiben Trainer": "✍️",
    }

    # Sticky banner
    st.markdown(
        """
        <div class="nav-sticky">
          <div style="padding:8px 14px;background:#ecfeff;border:1px solid #67e8f9;border-radius:12px;
                      margin:0;display:flex;align-items:center;gap:10px;justify-content:space-between;">
            <div style="font-weight:800;color:#0f172a;font-size:1.05rem;">🧭 Main Menu</div>
            <div style="color:#0c4a6e;font-size:0.95rem;">Use the selector <b>below</b> to switch sections</div>
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    # Default from URL OR session
    default = _qp_get_first("tab", st.session_state.get("main_tab_select", "Dashboard"))
    if default == "Vocab Trainer":
        default = "Schreiben Trainer"
        st.session_state["main_tab_select"] = default
        st.session_state["nav_sel"] = default
        st.session_state.pop("_chat_focus_tab", None)
        st.session_state["schreiben_pending_subtab"] = "Vocab Trainer"
        student_code_for_nav = st.session_state.get("student_code")
        if student_code_for_nav:
            st.session_state[f"schreiben_sub_tab_{student_code_for_nav}"] = "Vocab Trainer"
    if default not in tabs:
        default = "Dashboard"

    def _fmt(x: str) -> str:
        return f"{icons.get(x,'•')}  {x}"

    def _on_nav_change() -> None:
        sel_val = st.session_state["nav_dd"]
        st.session_state["main_tab_select"] = sel_val
        st.session_state["nav_sel"] = sel_val
        if sel_val != "Chat • Grammar • Exams":
            st.session_state.pop("_chat_focus_tab", None)
        _qp_set(tab=sel_val)

    sel = st.selectbox(
        "🧭 Main menu (tap ▾)",
        tabs,
        index=tabs.index(default),
        key="nav_dd",
        format_func=_fmt,
        help="This is the main selector. Tap ▾ to view all sections.",
        on_change=_on_nav_change,
    )



    # “You’re here” chip
    st.markdown(
        f"""
        <div style="margin-top:6px;">
          <span style="background:#e0f2fe;border:1px solid #7dd3fc;color:#075985;
                       padding:4px 10px;border-radius:999px;font-size:0.92rem;">
            You’re viewing: {icons.get(sel,'•')} <b>{sel}</b>
          </span>
        </div>
        """,
        unsafe_allow_html=True,
    )
    return sel

# --- Initialize nav (MUST be before any "if tab == ..." checks) ---
inject_notice_css()
try:
    if "nav_sel" not in st.session_state:
        st.session_state["nav_sel"] = _qp_get_first("tab", "Dashboard")
        st.session_state["main_tab_select"] = st.session_state["nav_sel"]
    tab = render_dropdown_nav()
except Exception as e:
    st.warning(f"Navigation init issue: {e}. Falling back to Dashboard.")
    tab = "Dashboard"
render_announcements_once(new_posts, tab == "Dashboard")


def _go_attendance() -> None:
    """Jump to the Attendance page inside the Classroom tab."""

    st.session_state["nav_sel"] = "My Course"
    st.session_state["main_tab_select"] = "My Course"
    st.session_state["coursebook_subtab"] = "🧑‍🏫 Classroom"
    st.session_state["cb_prev_subtab"] = "🧑‍🏫 Classroom"
    st.session_state["classroom_page"] = "Attendance"
    st.session_state["classroom_prev_page"] = "Attendance"
    _qp_set(tab="My Course")
    st.session_state.pop("_chat_focus_tab", None)
    st.session_state["need_rerun"] = True


def _go_next_assignment(day_value: Any) -> None:
    """Jump straight to the Assignment tab for ``day_value``."""

    st.session_state["nav_sel"] = "My Course"
    st.session_state["main_tab_select"] = "My Course"
    st.session_state["coursebook_subtab"] = "📘 Course Book"
    st.session_state["cb_prev_subtab"] = "📘 Course Book"
    st.session_state["coursebook_page"] = "Assignment"
    st.session_state["coursebook_prev_page"] = "Assignment"
    params = {"tab": "My Course"}
    if day_value not in (None, "", "?"):
        try:
            params["day"] = int(day_value)
        except Exception:
            params["day"] = day_value
    _qp_set(**params)
    st.session_state.pop("_chat_focus_tab", None)
    st.session_state["need_rerun"] = True


# =========================================================
# ===================== Dashboard =========================
# =========================================================
_ASSIGNMENT_DATE_COLUMNS = (
    "date",
    "submission_date",
    "submissiondate",
    "submitted_on",
    "submittedon",
    "submitted_at",
    "submittedat",
    "submitted",
    "timestamp",
    "created_at",
    "createdat",
    "created",
    "completed_at",
    "completedat",
)

_ASSIGNMENT_DATE_REGEX_PATTERNS = (
    re.compile(r"(?P<year>\d{4})[-/](?P<month>\d{1,2})[-/](?P<day>\d{1,2})"),
    re.compile(r"(?P<day>\d{1,2})[\./-](?P<month>\d{1,2})[\./-](?P<year>\d{4})"),
)


def _normalize_assignment_submission_dates(
    df: pd.DataFrame,
    *,
    candidate_columns: Iterable[str] = _ASSIGNMENT_DATE_COLUMNS,
) -> pd.Series:
    """Return normalized ``datetime.date`` objects for assignment submissions."""

    if not isinstance(df, pd.DataFrame):
        return pd.Series(dtype="object")

    if df.empty:
        return pd.Series([pd.NaT] * len(df), index=df.index, dtype="object")

    column_lookup = {str(col).strip().lower(): col for col in df.columns}
    source_column: Optional[str] = None
    for candidate in candidate_columns:
        if candidate in column_lookup:
            source_column = column_lookup[candidate]
            break

    if not source_column:
        return pd.Series([pd.NaT] * len(df), index=df.index, dtype="object")

    source = df[source_column]

    def _coerce(value: Any) -> Any:
        parsed_dt = _parse_contract_date_value(value)
        if parsed_dt is not None:
            return parsed_dt.date()

        try:
            ts = pd.to_datetime(value, errors="coerce")
        except Exception:
            ts = pd.NaT

        if pd.notna(ts):
            if isinstance(ts, pd.Timestamp):
                if ts.tzinfo is not None:
                    try:
                        ts = ts.tz_convert(None)
                    except (TypeError, ValueError):
                        try:
                            ts = ts.tz_localize(None)
                        except (TypeError, ValueError):
                            pass
                return ts.date()
            try:
                return ts.date()
            except Exception:
                pass

        text = "" if value is None else str(value).strip()
        if not text or text.lower() in ("nan", "none"):
            return pd.NaT

        for pattern in _ASSIGNMENT_DATE_REGEX_PATTERNS:
            match = pattern.search(text)
            if not match:
                continue
            parts = match.groupdict()
            try:
                year = int(parts["year"])
                month = int(parts["month"])
                day = int(parts["day"])
                return date(year, month, day)
            except (KeyError, ValueError):
                continue

        return pd.NaT

    normalized = source.apply(_coerce)
    normalized.index = source.index
    return normalized


if tab == "Dashboard":
    # ---------- Helpers ----------
    def safe_get(row, key, default=""):
        try: return row.get(key, default)
        except Exception: pass
        try: return getattr(row, key, default)
        except Exception: pass
        try: return row[key]
        except Exception: return default

    # Fallback parsers if globals not present
    def _fallback_parse_date(s):
        return _parse_contract_date_value(s)

    def _fallback_add_months(dt, n):
        y = dt.year + (dt.month - 1 + n) // 12
        m = (dt.month - 1 + n) % 12 + 1
        d = min(dt.day, calendar.monthrange(y, m)[1])
        return dt.replace(year=y, month=m, day=d)

    parse_contract_start_fn = globals().get("parse_contract_start", _fallback_parse_date)
    parse_contract_end_fn   = globals().get("parse_contract_end",   _fallback_parse_date)
    add_months_fn           = globals().get("add_months",           _fallback_add_months)

    # Global styles for chips & mini-cards
    inject_notice_css()

    # ---------- Ensure we have a student row ----------
    student_row = ensure_student_row(stop_if_missing=True)

    st.markdown("<div style='height:6px'></div>", unsafe_allow_html=True)

    st.divider()
    # ---------- 3) Motivation mini-cards (streak / vocab / leaderboard) ----------
    _student_code_raw = (st.session_state.get("student_code", "") or "").strip()
    _student_code = _student_code_raw.lower()
    _df_assign = load_assignment_scores()
    _df_assign["date"] = _normalize_assignment_submission_dates(_df_assign)
    _mask_student = _df_assign["studentcode"].str.lower().str.strip() == _student_code

    _dates = sorted(_df_assign[_mask_student]["date"].dropna().unique(), reverse=True)
    _streak = 1 if _dates else 0
    for i in range(1, len(_dates)):
        if (_dates[i - 1] - _dates[i]).days == 1:
            _streak += 1
        else:
            break

    _monday = date.today() - timedelta(days=date.today().weekday())
    _weekly_goal = 3
    _submitted_this_week = _df_assign[_mask_student & (_df_assign["date"] >= _monday)].shape[0]
    _goal_left = max(0, _weekly_goal - _submitted_this_week)

    _level = (safe_get(student_row, "Level", "A1") or "A1").upper().strip()
    _vocab_df = load_full_vocab_sheet()
    _vocab_item = get_vocab_of_the_day(_vocab_df, _level)

    _df_assign['level'] = _df_assign['level'].astype(str).str.upper().str.strip()
    _df_assign['score'] = pd.to_numeric(_df_assign['score'], errors='coerce')
    _df_assign_best = select_best_assignment_attempts(_df_assign)
    _min_assignments = 3
    _df_level = (
        _df_assign_best[_df_assign_best['level'] == _level]
        .groupby(['studentcode', 'name'], as_index=False)
        .agg(total_score=('score', 'sum'), completed=('assignment', 'nunique'))
    )
    _df_level = _df_level[_df_level['completed'] >= _min_assignments]
    _df_level = _df_level.sort_values(['total_score', 'completed'], ascending=[False, False]).reset_index(drop=True)
    _df_level['Rank'] = _df_level.index + 1
    _your_row = _df_level[
        _df_level['studentcode'].astype(str).str.strip().str.lower() == _student_code
    ]
    _total_students = len(_df_level)

    _streak_line = (
        f"<span class='pill pill-green'>{_streak} day{'s' if _streak != 1 else ''} streak</span>"
        if _streak > 0 else
        "<span class='pill pill-amber'>Start your streak today</span>"
    )
    _goal_line = (
        f"Submitted {_submitted_this_week}/{_weekly_goal} this week"
        + (f" — {_goal_left} to go" if _goal_left else " — goal met 🎉")
    )

    if _vocab_item:
        _vocab_chip = f"<span class='pill pill-purple'>{_vocab_item.get('german','')}</span>"
        _vocab_sub = f"{_vocab_item.get('english','')} · Level {_level}"
    else:
        _vocab_chip = "<span class='pill pill-amber'>No vocab available</span>"
        _vocab_sub = f"Level {_level}"

    if not _your_row.empty:
        _rank = int(_your_row.iloc[0]["Rank"])
        _total_score = int(_your_row.iloc[0]["total_score"])
        _rank_text = f"Rank #{_rank} of {_total_students} — {_total_score} pts"
        _lead_chip = "<span class='pill pill-purple'>On the board</span>"
    else:
        _rank_text = "Complete 3+ assignments to be ranked"
        _lead_chip = "<span class='pill pill-amber'>Not ranked yet</span>"

    _summary = get_assignment_summary(_student_code, _level, _df_assign)

    _missed_raw = _summary.get("missed", [])
    if isinstance(_missed_raw, (list, tuple, set)):
        _missed_list = [str(item).strip() for item in _missed_raw if str(item).strip()]
    elif isinstance(_missed_raw, str):
        _missed_list = [_missed_raw.strip()] if _missed_raw.strip() else []
    else:
        _missed_list = []

    _failed_raw = _summary.get("failed", [])
    if isinstance(_failed_raw, (list, tuple, set)):
        _failed_list = [str(item).strip() for item in _failed_raw if str(item).strip()]
    elif isinstance(_failed_raw, str):
        _failed_list = [_failed_raw.strip()] if _failed_raw.strip() else []
    else:
        _failed_list = []

    _next_lesson = _summary.get("next")

    if _missed_list:
        _missed_chip = f"<span class='pill pill-amber'>{len(_missed_list)} missed</span>"
        _missed_preview = ", ".join(_missed_list[:2]) + ("…" if len(_missed_list) > 2 else "")
    else:
        _missed_chip = "<span class='pill pill-green'>None</span>"
        _missed_preview = "You're on track"

    if _failed_list:
        _next_chip = "<span class='pill pill-amber'>Rework failed assignment</span>"
        if len(_failed_list) == 1:
            _next_sub = _failed_list[0]
        else:
            _next_sub = f"{_failed_list[0]} (+{len(_failed_list) - 1} more)"
    elif _next_lesson:
        _next_title = (
            f"Day {_next_lesson.get('day','?')}: {_next_lesson.get('chapter','?')} – {_next_lesson.get('topic','')}"
        )
        _next_chip = f"<span class='pill pill-purple'>{_next_title}</span>"
        _next_sub = _next_lesson.get("goal", "")
    elif _missed_list:
        _next_chip = "<span class='pill pill-amber'>Finish missed work</span>"
        _next_sub = "Complete skipped assignments first"
    else:
        _next_chip = "<span class='pill pill-green'>All caught up</span>"
        _next_sub = ""
    _class_name_clean, _class_name_lookup = _resolve_class_name(
        safe_get(student_row, "ClassName", ""),
        level=_level,
    )
    _att_sessions, _att_hours = (0, 0.0)
    if _class_name_lookup and _student_code_raw:
        _att_sessions, _att_hours = fetch_attendance_summary(
            _student_code_raw,
            _class_name_lookup,
        )
    _attendance_chip = (
        f"<span class='pill pill-purple'>{_att_sessions} sessions • {_att_hours:.1f}h</span>"
    )



    st.markdown(
        f"""
        <div class="minirow">
          <div class="minicard">
            <h4>🏅 Assignment Streak</h4>
            <div>{_streak_line}</div>
            <div class="sub">{_goal_line}</div>
          </div>
          <div class="minicard">
            <h4>🗣️ Vocab of the Day</h4>
            <div>{_vocab_chip}</div>
            <div class="sub">{_vocab_sub}</div>
          </div>
          <div class="minicard">
            <h4>🏆 Leaderboard</h4>
            <div>{_lead_chip}</div>
            <div class="sub">{_rank_text}</div>
          </div>
          <div class="minicard">
            <h4>📚 Missed Assignments</h4>
            <div>{_missed_chip}</div>
            <div class="sub">{_missed_preview}</div>
          </div>
          <div class="minicard">
            <h4>⏭️ Next Assignment</h4>
            <div>{_next_chip}</div>
            <div class="sub">{_next_sub}</div>
          </div>
          <div class="minicard">
            <h4>🕛 Attendance</h4>
            <div>{_attendance_chip}</div>
            <div class="sub"></div>
          </div>
        </div>
        """,
        unsafe_allow_html=True
    )
    if _next_lesson:
        raw_day = _next_lesson.get("day", "?")
        day_text = str(raw_day).strip()
        if not day_text or day_text.lower() in {"none", "nan"}:
            button_label = "Start your next assignment"
        else:
            button_label = f"Start Day {day_text} assignment"
        st.button(
            button_label,
            type="primary",
            width="stretch",
            key="btn_dash_next_assignment",
            on_click=_go_next_assignment,
            args=(_next_lesson.get("day"),),
        )
    st.button("View attendance", on_click=_go_attendance)
    st.divider()

    # ---------- Student header (compact) + details (expander) ----------
    name = safe_get(student_row, "Name")
    level = safe_get(student_row, "Level", "")
    code  = safe_get(student_row, "StudentCode", "")
    _class_raw = safe_get(student_row, "ClassName", "")
    _, class_name = _resolve_class_name(_class_raw, level=level)
    try:
        bal_val = float(str(safe_get(student_row, "Balance", 0)).replace(",", "").strip() or 0)
    except Exception:
        bal_val = 0.0

    st.markdown(
        f"<div style='display:flex;flex-wrap:wrap;gap:10px;align-items:center;"
        f"padding:8px 10px;border:1px solid rgba(148,163,184,.35);border-radius:10px;"
        f"background:#ffffff;'>"
        f"<b>👤 {name}</b>"
        f"<span style='background:#eef4ff;color:#2541b2;padding:2px 8px;border-radius:999px;'>Level: {level}</span>"
        f"<span style='background:#f3e8ff;color:#6b21a8;padding:2px 8px;border-radius:999px;'>Class: {class_name}</span>"
        f"<span style='background:#f1f5f9;color:#334155;padding:2px 8px;border-radius:999px;'>Code: <code>{code}</code></span>"
        + (
            f"<span style='background:#fff7ed;color:#7c2d12;padding:2px 8px;border-radius:999px;'>Balance: {format_cedis(bal_val)}</span>"
            if bal_val > 0
            else f"<span style='background:#ecfdf5;color:#065f46;padding:2px 8px;border-radius:999px;'>Balance: {format_cedis(0)}</span>"
        )
        + "</div>",
        unsafe_allow_html=True
    )

    with st.expander("👤 Student details", expanded=False):
        info_html = f"""
        <div style='
            background:#f8fbff;
            border:1.6px solid #cfe3ff;
            border-radius:12px;
            padding:12px 14px;
            margin-top:8px;
            box-shadow:0 2px 8px rgba(44,106,221,0.04);
            font-size:1.04em;
            color:#17325e;
            font-family:"Segoe UI","Arial",sans-serif;
            letter-spacing:.01em;'>
            <div style="font-weight:700;font-size:1.12em;margin-bottom:6px;">
                👤 {name}
            </div>
            <div style="font-size:1em; margin-bottom:4px;">
                <b>Level:</b> {safe_get(student_row, 'Level', '')} &nbsp;|&nbsp; 
                <b>Code:</b> <code>{safe_get(student_row, 'StudentCode', '')}</code> &nbsp;|&nbsp;
                <b>Status:</b> {safe_get(student_row, 'Status', '')}
            </div>
            <div style="font-size:1em; margin-bottom:4px;">
                <b>Email:</b> {safe_get(student_row, 'Email', '')} &nbsp;|&nbsp;
                <b>Phone:</b> {safe_get(student_row, 'Phone', '')} &nbsp;|&nbsp;
                <b>Location:</b> {safe_get(student_row, 'Location', '')}
            </div>
            <div style="font-size:1em;">
                <b>Contract:</b> {safe_get(student_row, 'ContractStart', '')} ➔ {safe_get(student_row, 'ContractEnd', '')} &nbsp;|&nbsp;
                <b>Enroll Date:</b> {safe_get(student_row, 'EnrollDate', '')}
            </div>
        </div>
        """
        st.markdown(info_html, unsafe_allow_html=True)

    # ---------- Payments & Renewal (policy-aligned, all inside one expander) ----------
    from datetime import datetime as _dt
    import calendar as _cal

    _read_money = globals().get("_read_money")
    if _read_money is None:
        def _read_money(x):
            try:
                s = str(x).replace(",", "").strip()
                return float(s) if s not in ("", "nan", "None") else 0.0
            except Exception:
                return 0.0

    def _fallback_parse_date(s):
        for f in ("%Y-%m-%d", "%m/%d/%Y", "%d.%m.%y", "%d/%m/%Y", "%d-%m-%Y"):
            try:
                return _dt.strptime(str(s).strip(), f)
            except Exception:
                pass
        return None

    def _fallback_add_months(dt, n):
        y = dt.year + (dt.month - 1 + n) // 12
        m = (dt.month - 1 + n) % 12 + 1
        d = min(dt.day, _cal.monthrange(y, m)[1])
        return dt.replace(year=y, month=m, day=d)

    _parse_start = (
        globals().get("parse_contract_start_fn")
        or globals().get("parse_contract_start")
        or _fallback_parse_date
    )
    _parse_end = (
        globals().get("parse_contract_end_fn")
        or globals().get("parse_contract_end")
        or _fallback_parse_date
    )
    _add_months = (
        globals().get("add_months_fn")
        or globals().get("add_months")
        or _fallback_add_months
    )

    _today = _dt.today().date()

    _cs = None
    for _k in ["ContractStart", "StartDate", "ContractBegin", "Start", "Begin"]:
        _s = str(safe_get(student_row, _k, "") or "").strip()
        if _s:
            _cs = _parse_start(_s)
            break
    _first_due_dt = _add_months(_cs, 1) if _cs else None
    _first_due = _first_due_dt.date() if _first_due_dt and hasattr(_first_due_dt, "date") else _first_due_dt

    _balance = _read_money(safe_get(student_row, "Balance", 0))

    _exp_title = "💳 Payments (info)"
    _severity = "info"
    if _balance > 0 and _first_due:
        if _today > _first_due:
            _days_over = (_today - _first_due).days
            _exp_title = f"💳 Payments • overdue {_days_over}d"
            _severity = "error"
            _msg = (
                f"💸 **Overdue by {_days_over} day{'s' if _days_over != 1 else ''}.** "
                f"Amount due: **{format_cedis(_balance)}**. First due: {_first_due:%d %b %Y}."
            )
        elif _today == _first_due:
            _exp_title = "💳 Payments • due today"
            _severity = "warning"
            _msg = f"⏳ **Payment due today** ({_first_due:%d %b %Y}). Amount due: **{format_cedis(_balance)}**."
        else:
            _exp_title = "💳 Payments (info)"
            _severity = "info"
            _days_left = (_first_due - _today).days
            _msg = (
                f"No payment expected yet. Your first payment date is **{_first_due:%d %b %Y}** "
                f"(in {_days_left} day{'s' if _days_left != 1 else ''}). Current balance: **{format_cedis(_balance)}**."
            )
    elif _balance > 0 and not _first_due:
        _exp_title = "💳 Payments • schedule unknown"
        _severity = "info"
        _msg = (
            "ℹ️ You have a positive balance, but I couldn’t read your contract start date "
            "to compute the first payment date. Please contact the office."
        )
    else:
        _exp_title = "💳 Payments (info)"
        _severity = "info"
        if _first_due:
            _msg = (
                "No outstanding balance. You’re not expected to pay anything now. "
                f"Your first payment date (if applicable) is **{_first_due:%d %b %Y}**."
            )
        else:
            _msg = (
                "No outstanding balance. You’re not expected to pay anything now. "
                "We’ll compute your first payment date after your contract start is on file."
            )

    with st.expander(_exp_title, expanded=False):
        if _severity == "error":
            st.error(_msg)
        elif _severity == "warning":
            st.warning(_msg)
        else:
            st.info(_msg)

        _cs_str = _cs.strftime("%d %b %Y") if _cs else "—"
        _fd_str = _first_due.strftime("%d %b %Y") if _first_due else "—"
        st.markdown(
            f"""
            **Details**
            - Contract start: **{_cs_str}**
            - First payment due (start + 1 month): **{_fd_str}**
            - Current balance: **{format_cedis(_balance)}**
            """
        )

        EXT_FEE = 1000
        _ce = _parse_end(safe_get(student_row, "ContractEnd", ""))
        _ce_date = _ce.date() if hasattr(_ce, "date") else _ce
        if _ce_date:
            _days_left = (_ce_date - _today).days
            if _days_left < 0:
                st.error(
                    f"⚠️ Your contract ended on **{_ce_date:%d %b %Y}**. "
                    f"If you need more time, extension costs **{format_cedis(EXT_FEE)}/month**."
                )
            elif _days_left <= 14:
                st.warning(
                    f"⏰ Your contract ends in **{_days_left} day{'s' if _days_left != 1 else ''}** "
                    f"(**{_ce_date:%d %b %Y}**). Extension costs **{format_cedis(EXT_FEE)}/month**."
                )

    # ---------- Always-visible Contract Alert ----------
    from datetime import datetime as _dt

    def _fallback_parse_date(_s):
        for _f in ("%Y-%m-%d", "%m/%d/%Y", "%d.%m.%y", "%d/%m/%Y", "%d-%m-%Y"):
            try:
                return _dt.strptime(str(_s).strip(), _f)
            except Exception:
                pass
        return None

    _parse_end = (
        globals().get("parse_contract_end_fn")
        or globals().get("parse_contract_end")
        or _fallback_parse_date
    )

    _today = _dt.today().date()
    _ce_raw = _parse_end(safe_get(student_row, "ContractEnd", ""))
    _ce_date = _ce_raw.date() if hasattr(_ce_raw, "date") else _ce_raw

    st.markdown("""
    <style>
      .contract-alert { border-radius:12px; padding:12px 14px; margin:8px 0 10px 0; font-weight:600; }
      .ca-warn { background:#fff7ed; color:#7c2d12; border:1px solid #fed7aa; }
      .ca-err  { background:#fef2f2; color:#991b1b; border:1px solid #fecaca; }
      .ca-text { font-size:1rem; line-height:1.55; }
      .ca-cta  { margin-top:6px; font-size:.95rem; }
      @media (max-width:640px){
        .contract-alert{ padding:10px 12px; }
        .ca-text{ font-size:1.02rem; }
      }
    </style>
    """, unsafe_allow_html=True)

    if _ce_date:
        _days_left = (_ce_date - _today).days
        _student_code = str(safe_get(student_row, "StudentCode", "") or "").strip().lower()
        _alert_key = f"hide_contract_alert:{_student_code}:{_ce_date.isoformat()}:{_today.isoformat()}"
        _ext_fee = 1000

        if not st.session_state.get(_alert_key, False):
            if _days_left < 0:
                _msg = (
                    f"⚠️ <b>Your contract ended on {_ce_date:%d %b %Y}.</b> "
                    f"To continue, extension costs <b>{format_cedis(_ext_fee)}/month</b>."
                )
                _cls = "ca-err"
            elif _days_left <= 14:
                _msg = (
                    f"⏰ <b>Your contract ends in {_days_left} day{'s' if _days_left != 1 else ''} "
                    f"({_ce_date:%d %b %Y}).</b> Extension costs <b>{format_cedis(_ext_fee)}/month</b>."
                )
                _cls = "ca-warn"
            else:
                _msg = ""
                _cls = ""

            if _msg:
                st.markdown(
                    f"<div class='contract-alert {_cls}'><div class='ca-text'>{_msg}</div></div>",
                    unsafe_allow_html=True
                )
                if st.button("Got it — hide this notice for today", key=f"btn_contract_alert_{_student_code}"):
                    st.session_state[_alert_key] = True
                    refresh_with_toast()

    # ---------- Class schedules ----------
    with st.expander("🗓️ Class Schedule & Upcoming Sessions", expanded=False):
        if not st.session_state.get("student_level"):
            ensure_student_level()
        GROUP_SCHEDULES = load_group_schedules()

        from datetime import datetime as _dt_local, timedelta as _td_local
        student_level_val = st.session_state.get("student_level", "")
        class_name_clean, class_name_lookup = _resolve_class_name(
            safe_get(student_row, "ClassName", ""),
            level=student_level_val,
        )
        class_name_display = class_name_clean or class_name_lookup
        class_schedule = GROUP_SCHEDULES.get(class_name_lookup)
        week_days = [
            "Monday",
            "Tuesday",
            "Wednesday",
            "Thursday",
            "Friday",
            "Saturday",
            "Sunday",
        ]

        if not class_name_lookup or not class_schedule:
            st.info("🚩 Your class is not set yet. Please contact your teacher or the office.")
        else:
            days = class_schedule.get("days", [])
            time_str = class_schedule.get("time", "")
            start_dt = class_schedule.get("start_date", "")
            end_dt = class_schedule.get("end_date", "")
            doc_url = class_schedule.get("doc_url", "")

            today = _dt_local.today().date()
            start_date_obj = None
            end_date_obj = None
            try:
                if start_dt:
                    start_date_obj = _dt_local.strptime(start_dt, "%Y-%m-%d").date()
            except Exception:
                start_date_obj = None
            try:
                if end_dt:
                    end_date_obj = _dt_local.strptime(end_dt, "%Y-%m-%d").date()
            except Exception:
                end_date_obj = None

            before_start = bool(start_date_obj and today < start_date_obj)
            after_end = bool(end_date_obj and today > end_date_obj)
            day_indices = [week_days.index(d) for d in days if d in week_days] if isinstance(days, list) else []

            def get_next_sessions(from_date, weekday_indices, limit=3, end_date=None):
                results = []
                if not weekday_indices:
                    return results
                check_date = from_date
                while len(results) < limit:
                    if end_date and check_date > end_date:
                        break
                    if check_date.weekday() in weekday_indices:
                        results.append(check_date)
                    check_date += _td_local(days=1)
                return results

            if before_start and start_date_obj:
                upcoming_sessions = get_next_sessions(start_date_obj, day_indices, limit=3, end_date=end_date_obj)
            elif after_end:
                upcoming_sessions = []
            else:
                upcoming_sessions = get_next_sessions(today, day_indices, limit=3, end_date=end_date_obj)

            if after_end:
                end_str = end_date_obj.strftime('%d %b %Y') if end_date_obj else end_dt
                st.error(
                    f"❌ Your class ({class_name_display}) ended on {end_str}. Please contact the office for next steps."
                )
            else:
                if upcoming_sessions:
                    items = []
                    for session_date in upcoming_sessions:
                        weekday_name = week_days[session_date.weekday()]
                        display_date = session_date.strftime("%d %b")
                        items.append(
                            f"<li style='margin-bottom:6px;'><b>{weekday_name}</b> "
                            f"<span style='color:#1976d2;'>{display_date}</span> "
                            f"<span style='color:#333;'>{time_str}</span></li>"
                        )
                    session_items_html = "<ul style='padding-left:16px; margin:9px 0 0 0;'>" + "".join(items) + "</ul>"
                else:
                    session_items_html = "<span style='color:#c62828;'>No upcoming sessions in the visible window.</span>"

                if before_start and start_date_obj:
                    days_until = (start_date_obj - today).days
                    label = f"Starts in {days_until} day{'s' if days_until != 1 else ''} (on {start_date_obj.strftime('%d %b %Y')})"
                    bar_html = f"""
        <div style="margin-top:8px; font-size:0.85em;">
          <div style="margin-bottom:4px;">{label}</div>
          <div style="background:#ddd; border-radius:6px; overflow:hidden; height:12px; width:100%;">
            <div style="width:3%; background:#1976d2; height:100%;"></div>
          </div>
        </div>"""
                elif start_date_obj and end_date_obj:
                    total_days = (end_date_obj - start_date_obj).days + 1
                    elapsed = max(0, (today - start_date_obj).days + 1) if today >= start_date_obj else 0
                    remaining = max(0, (end_date_obj - today).days)
                    percent = int((elapsed / total_days) * 100) if total_days > 0 else 100
                    percent = min(100, max(0, percent))
                    label = f"{remaining} day{'s' if remaining != 1 else ''} remaining in course"
                    bar_html = f"""
        <div style="margin-top:8px; font-size:0.85em;">
          <div style="margin-bottom:4px;">{label}</div>
          <div style="background:#ddd; border-radius:6px; overflow:hidden; height:12px; width:100%;">
            <div style="width:{percent}%; background: linear-gradient(90deg,#1976d2,#4da6ff); height:100%;"></div>
          </div>
          <div style="margin-top:2px; font-size:0.75em;">
            Progress: {percent}% (started {elapsed} of {total_days} days)
          </div>
        </div>"""
                else:
                    bar_html = f"""
        <div style="margin-top:8px; font-size:0.85em;">
          <b>Course period:</b> {start_dt or '[not set]'} to {end_dt or '[not set]'}
        </div>"""

                period_str = f"{start_dt or '[not set]'} to {end_dt or '[not set]'}"
                st.markdown(
                    f"""
        <div style='border:2px solid #17617a; border-radius:14px;
                    padding:13px 11px; margin-bottom:13px;
                    background:#eaf6fb; font-size:1.15em;
                    line-height:1.65; color:#232323;'>
          <b style="font-size:1.09em;">🗓️ Your Next Classes ({class_name}):</b><br>
          {session_items_html}
          {bar_html}
          <div style="font-size:0.98em; margin-top:6px;">
            <b>Course period:</b> {period_str}
          </div>
          {f'<a href="{doc_url}" target="_blank" '
            f'style="font-size:1em;color:#17617a;text-decoration:underline;margin-top:6px;display:inline-block;">📄 View/download full class schedule</a>'
            if doc_url else ''}
        </div>""",
                    unsafe_allow_html=True,
                )

    # ---------- Goethe exam ----------
    GOETHE_EXAM_DATES = {
        "A1": (date(2025, 10, 13), 2850, None),
        "A2": (date(2025, 10, 14), 2400, None),
        "B1": (date(2025, 10, 15), 2750, 880),
        "B2": (date(2025, 10, 16), 2500, 840),
        "C1": (date(2025, 10, 17), 2450, 700),
    }
    level = (safe_get(student_row, "Level", "") or "").upper().replace(" ", "")
    exam_info = GOETHE_EXAM_DATES.get(level)
    days_to_exam: Optional[int] = None
    fee_text = ""
    if exam_info:
        exam_date, fee, module_fee = exam_info
        days_to_exam = (exam_date - date.today()).days
        fee_text = f"**Fee:** {format_cedis(fee)}"
        if module_fee:
            fee_text += f" &nbsp; | &nbsp; **Per Module:** {format_cedis(module_fee)}"

    expander_title = (
        f"⏳ Goethe Exam: {days_to_exam} days left" if days_to_exam is not None else "⏳ Goethe Exam"
    )
    st.caption(expander_title)
    with st.expander(expander_title, expanded=False):
        if exam_info:
            register_link = "[Register online here](https://www.goethe.de/ins/gh/en/spr/prf.html)"
            if days_to_exam is not None and days_to_exam > 0:
                st.info(
                    "\n".join(
                        [
                            f"Your {level} exam is in {days_to_exam} days ({exam_date:%d %b %Y}).",
                            fee_text,
                            register_link,
                        ]
                    )
                )
            elif days_to_exam == 0:
                st.success(
                    "\n".join(
                        [
                            f"📝 Registration is today ({exam_date:%d %b %Y}).",
                            fee_text,
                            register_link,
                        ]
                    )
                )
            else:
                st.error(
                    "\n".join(
                        [
                            f"❌ Your {level} exam was on {exam_date:%d %b %Y}, {abs(days_to_exam)} days ago.",
                            fee_text,
                        ]
                    )
                )

            st.caption(
                EXAM_ADVICE.get(
                    level,
                    "No exam advice available for your level."
                )
            )
        else:
            st.warning("No exam date configured for your level.")
    st.divider()

    # ---------- Footer ----------
    render_app_footer(FOOTER_LINKS)



def render_section(day_info: dict, key: str, title: str, icon: str) -> None:
    """Render a lesson section (supports list or single dict)."""
    content = day_info.get(key)
    if not content:
        return
    items = content if isinstance(content, list) else [content]
    st.markdown(f"#### {icon} {title}")
    for idx, part in enumerate(items):
        if len(items) > 1:
            st.markdown(f"###### {icon} Part {idx+1} of {len(items)}: Chapter {part.get('chapter','')}")
        if part.get('video'):
            st.video(part['video'])
        if part.get('grammarbook_link'):
            render_link("📘 Grammar Book (Notes)", part['grammarbook_link'])
            st.markdown(
                '<em>Further notice:</em> 📘 contains notes; 📒 is your workbook assignment.',
                unsafe_allow_html=True
            )
        if part.get('workbook_link'):
            render_link("📒 Workbook (Assignment)", part['workbook_link'])
            render_assignment_reminder()
        extras = part.get('extra_resources')
        if extras:
            for ex in (extras if isinstance(extras, list) else [extras]):
                render_link("🔗 Extra", ex)

# -------------------------
# Slack helpers (optional)
# -------------------------
SLACK_DEBUG = (os.getenv("SLACK_DEBUG", "0") == "1")

def _slack_url() -> str:
    """Resolve Slack webhook URL (ENV first, then st.secrets)."""
    url = (os.getenv("SLACK_WEBHOOK_URL") or "").strip()
    if not url:
        try:
            url = (st.secrets.get("slack", {}).get("webhook_url", "") if hasattr(st, "secrets") else "").strip()
        except Exception:
            url = ""
    return url

def get_slack_webhook() -> str:
    """Back-compat alias to _slack_url()."""
    return _slack_url()

def notify_slack(text: str) -> Tuple[bool, str]:
    """
    Post a plain text message to the Slack webhook.
    Returns (ok, info). If SLACK_DEBUG=1, more verbose info is printed in logs.
    """
    url = _slack_url()
    if not url:
        return False, "missing_webhook"
    try:
        resp = api_post(url, json={"text": text}, timeout=6)
        ok = 200 <= resp.status_code < 300
        return ok, f"status={resp.status_code}"
    except Exception as e:
        return False, str(e)

def notify_slack_submission(
    webhook_url: str,
    *,
    student_name: str,
    student_code: str,
    level: str,
    day: int,
    chapter: str,
    receipt: str,
    preview: str
) -> None:
    """Send a compact submission notification to Slack (best-effort)."""
    if not webhook_url:
        return
    text = (
        f"*New submission* • {student_name} ({student_code})\n"
        f"*Level:* {level}  •  *Day:* {day}\n"
        f"*Chapter:* {chapter}\n"
        f"*Ref:* `{receipt}`\n"
        f"*Preview:* {preview[:180]}{'…' if len(preview) > 180 else ''}"
    )
    try:
        api_post(webhook_url, json={"text": text}, timeout=6)
    except Exception:
        pass  # never block the student


def has_telegram_subscription(student_code: str) -> bool:
    """Return True if the student has enabled Telegram notifications."""
    _db = globals().get("db")
    if _db is None:
        return False
    try:
        snap = _db.collection("telegram_subscriptions").document(student_code).get()
        return snap.exists and bool((snap.to_dict() or {}).get("chat_id"))
    except Exception:
        return False

# -------------------------
# Firestore helpers (uses your existing `db` and `from firebase_admin import firestore`)
# -------------------------

# -------------------------
# Misc existing helper preserved
# -------------------------
def post_message(
    level: str,
    class_name: str,
    code: str,
    name: str,
    content: str,
    reply_to: Optional[str] = None,
) -> None:
    """Post a message to the class board for a specific class."""
    posts_ref = (
        db.collection("class_board")
        .document(level)
        .collection("classes")
        .document(class_name)
        .collection("posts")
    )
    posts_ref.add(
        {
            "student_code": code,
            "student_name": name,
            "content": content.strip(),
            "created_at": _dt.now(_timezone.utc),
            "reply_to": reply_to,
        }
    )

RESOURCE_LABELS = {
    'video': '🎥 Video',
    'grammarbook_link': '📘 Grammar',
    'workbook_link': '📒 Workbook',
    'extra_resources': '🔗 Extra'
}


def _next_available_lesson_day(
    schedule: Sequence[MutableMapping[str, Any]] | Sequence[dict],
    current_index: int,
) -> Optional[int]:
    """Return the next non-zero ``day`` value following ``current_index``."""

    for entry in schedule[current_index + 1 :]:
        day_val = _coerce_day((entry or {}).get("day"))
        if day_val not in (None, 0):
            return day_val
    for entry in schedule:
        day_val = _coerce_day((entry or {}).get("day"))
        if day_val not in (None, 0):
            return day_val
    return None


def render_day_zero_onboarding(
    info: MutableMapping[str, Any] | dict,
    schedule: Sequence[MutableMapping[str, Any]] | Sequence[dict],
    idx: int,
) -> None:
    """Render a dedicated onboarding layout for the tutorial day."""

    st.caption("Review the tutorial if you need a refresher.")

    goal_text = _safe_str(info.get("goal"))
    intro_text = _safe_str(info.get("instruction"))
    cards = info.get("onboarding_cards") or []

    with st.expander("📘 Day 0 Tutorial – optional", expanded=False):
        st.markdown(
            """
            <style>
            .day0-card{background:#f8fafc;border:1px solid #e2e8f0;border-radius:16px;padding:18px 16px;
                       box-shadow:0 8px 24px rgba(15,23,42,.08);min-height:260px;display:flex;flex-direction:column;gap:12px;}
            .day0-card__title{font-size:1.1rem;font-weight:600;color:#0f172a;}
            .day0-card__helper{color:#475569;font-size:.95rem;line-height:1.4;}
            .day0-card ol{margin:0;padding-left:1.2rem;color:#1e293b;font-size:.95rem;line-height:1.45;}
            .day0-card__cta{margin-top:auto;}
            .day0-card__cta button{width:100%;border-radius:999px;}
            @media (max-width:768px){.day0-card{min-height:auto;}}
            </style>
            """,
            unsafe_allow_html=True,
        )

        if goal_text:
            st.markdown(f"#### 🚀 {goal_text}")
        else:
            st.markdown("#### 🚀 Welcome to Falowen")
        if intro_text:
            st.caption(intro_text)

        if not isinstance(cards, Sequence) or isinstance(cards, (str, bytes)) or not cards:
            info_lines: List[str] = []
            if goal_text:
                info_lines.append(f"🎯 **Goal:** {goal_text}")
            if intro_text:
                info_lines.append(f"📝 **Instruction:** {intro_text}")
            if info_lines:
                st.info("\n\n".join(info_lines))
            return

        target_day = _next_available_lesson_day(schedule, idx) or 1

        def _set_coursebook_page(page: str) -> None:
            st.session_state["coursebook_page"] = page
            st.session_state["coursebook_prev_page"] = page

    def _jump_course(page: str) -> None:
        _go_next_assignment(target_day)
        _set_coursebook_page(page)

    def _jump_classroom() -> None:
        st.session_state["nav_sel"] = "My Course"
        st.session_state["main_tab_select"] = "My Course"
        st.session_state["coursebook_subtab"] = "🧑‍🏫 Classroom"
        st.session_state["cb_prev_subtab"] = "🧑‍🏫 Classroom"
        _qp_set(tab="My Course")
        st.session_state["need_rerun"] = True

    def _go_dashboard() -> None:
        st.session_state["nav_sel"] = "Dashboard"
        st.session_state["main_tab_select"] = "Dashboard"
        _qp_set(tab="Dashboard")
        st.session_state["need_rerun"] = True

    action_map: Dict[str, Any] = {
        "overview": lambda: _jump_course("Overview"),
        "assignment": lambda: _jump_course("Assignment"),
        "submit": lambda: _jump_course("Submit"),
        "classroom": _jump_classroom,
        "attendance": _go_attendance,
        "dashboard": _go_dashboard,
    }

    chunk = 3
    for start in range(0, len(cards), chunk):
        cols = st.columns(len(cards[start : start + chunk]))
        for col, card in zip(cols, cards[start : start + chunk]):
            with col:
                title_raw = _safe_str(card.get("title"), "")
                helper_raw = _safe_str(card.get("helper"), "")
                title = html.escape(title_raw)
                helper = html.escape(helper_raw)
                steps = card.get("steps") or []
                if not isinstance(steps, Sequence):
                    steps = []
                steps_html = "".join(
                    f"<li>{html.escape(_safe_str(step, ''))}</li>" for step in steps if _safe_str(step, "")
                )

                st.markdown("<div class='day0-card'>", unsafe_allow_html=True)
                if title:
                    st.markdown(
                        f"<div class='day0-card__title'>{title}</div>",
                        unsafe_allow_html=True,
                    )
                if helper:
                    st.markdown(
                        f"<div class='day0-card__helper'>{helper}</div>",
                        unsafe_allow_html=True,
                    )
                if steps_html:
                    st.markdown(
                        f"<ol>{steps_html}</ol>",
                        unsafe_allow_html=True,
                    )

                cta_label = _safe_str(card.get("cta_label"))
                action_key = _safe_lower(card.get("action"))
                action_fn = action_map.get(action_key)
                if cta_label and callable(action_fn):
                    st.markdown("<div class='day0-card__cta'>", unsafe_allow_html=True)
                    st.button(
                        cta_label,
                        key=f"day0_card_btn_{start}_{title_raw}",
                        on_click=action_fn,
                        use_container_width=True,
                    )
                    st.markdown("</div>", unsafe_allow_html=True)
                st.markdown("</div>", unsafe_allow_html=True)


# ---- Firestore Helpers ----
if tab == "My Course":
    # === HANDLE ALL SWITCHING *BEFORE* ANY WIDGET ===
    # Jump flags set by buttons elsewhere
    if st.session_state.get("__go_classroom"):
        st.session_state["coursebook_subtab"] = "🧑‍🏫 Classroom"
        del st.session_state["__go_classroom"]
        refresh_with_toast()

    if st.session_state.get("__go_notes"):
        st.session_state["coursebook_subtab"] = "📒 Learning Notes"
        del st.session_state["__go_notes"]
        refresh_with_toast()

    # Backward-compat: older code may still set this
    if st.session_state.get("switch_to_notes"):
        st.session_state["coursebook_subtab"] = "📒 Learning Notes"
        del st.session_state["switch_to_notes"]
        refresh_with_toast()

    # First run default
    if "coursebook_subtab" not in st.session_state:
        st.session_state["coursebook_subtab"] = "📘 Course Book"
    if "cb_prev_subtab" not in st.session_state:
        st.session_state["cb_prev_subtab"] = st.session_state["coursebook_subtab"]

    # Header (render once)
    st.markdown(
        '''
        <div style="
            padding: 16px;
            background: #007bff;
            color: #ffffff;
            border-radius: 8px;
            text-align: center;
            margin-bottom: 16px;
            box-shadow: 0 4px 6px rgba(0,0,0,0.1);
        ">
            <span style="font-size:1.8rem; font-weight:600;">📈 My Course</span>
        </div>
        ''',
        unsafe_allow_html=True
    )
    st.divider()

    # Subtabs (1: Classroom, 2: Course Book, 3: Learning Notes)
    def on_cb_subtab_change() -> None:
        prev = st.session_state.get("cb_prev_subtab")
        curr = st.session_state.get("coursebook_subtab")
        if prev == "📒 Learning Notes":
            code = st.session_state.get("student_code", "") or ""
            if not code:
                st.error("Student code is required.")
            else:
                notes_key = f"notes_{code}"
                notes = st.session_state.get(notes_key)
                if notes is not None:
                    save_notes_to_db(code, notes)
        elif prev == "🧑‍🏫 Classroom":
            code = (
                st.session_state.get("student_code")
                or (st.session_state.get("student_row") or {}).get("StudentCode", "")
            )
            if code:
                if str(st.session_state.get("q_text", "")).strip():
                    save_now("q_text", code)
                for k in [key for key in st.session_state.keys() if key.startswith("q_reply_box_")]:
                    if str(st.session_state.get(k, "")).strip():
                        save_now(k, code)
        elif prev == "📘 Course Book":
            draft_key = st.session_state.get("coursebook_draft_key")
            code = (
                st.session_state.get("student_code")
                or (st.session_state.get("student_row") or {}).get("StudentCode", "")
            )
            if draft_key and code:
                last_val_key, *_ = _draft_state_keys(draft_key)
                if st.session_state.get(draft_key, "") != st.session_state.get(last_val_key, ""):
                    save_now(draft_key, code)
        st.session_state["cb_prev_subtab"] = curr


    # Subtabs (1: Classroom, 2: Course Book, 3: Learning Notes)
    cb_subtab = st.radio(
        "Select section:",
        ["🧑‍🏫 Classroom", "📘 Course Book", "📒 Learning Notes"],
        horizontal=True,
        key="coursebook_subtab",
        on_change=on_cb_subtab_change,
    )

    # ---------- DB (Firestore) bootstrap ----------
    def _get_db():
        global db
        existing = (
            db
            or getattr(_falowen_sessions, "db", None)
            or getattr(_falowen_sessions, "_db_client", None)
        )
        if existing is not None:
            _falowen_sessions.db = existing
            if hasattr(_falowen_sessions, "_db_client"):
                _falowen_sessions._db_client = existing
            db = existing
            return existing
        # Try Firebase Admin SDK first (firestore.client())
        client = None
        try:
            import firebase_admin
            from firebase_admin import firestore as fbfs

            if not firebase_admin._apps:
                firebase_admin.initialize_app()
            client = fbfs.client()
        except Exception:
            client = None
        if client is None:
            try:
                from google.cloud import firestore as gcf

                client = gcf.Client()
            except Exception:
                st.error(
                    "Firestore client isn't configured. Provide Firebase Admin creds or set GOOGLE_APPLICATION_CREDENTIALS.",
                    icon="🛑",
                )
                raise
        _falowen_sessions.db = client
        if hasattr(_falowen_sessions, "_db_client"):
            _falowen_sessions._db_client = client
        db = client
        return client

    db = _get_db()


    # === COURSE BOOK SUBTAB (mini-tabs inside) ===
    if cb_subtab == "📘 Course Book":
        from datetime import date, timedelta  # needed inside this branch

        st.markdown(
            '''
            <div style="
                padding: 16px;
                background: #007bff;
                color: #ffffff;
                border-radius: 8px;
                text-align: center;
                margin-bottom: 16px;
                box-shadow: 0 4px 6px rgba(0,0,0,0.1);
            ">
                <span style="font-size:1.8rem; font-weight:600;">📘 Course Book</span>
            </div>
            ''',
            unsafe_allow_html=True
        )
        st.divider()

        # ---- Load schedule (normalized) ----
        if not st.session_state.get("student_level"):
            ensure_student_level()  
        student_level = st.session_state.get("student_level", "A1")
        level_key = (student_level or "A1").strip().upper()
        schedules = load_level_schedules()
        schedule = schedules.get(level_key, schedules.get("A1", []))
        if not schedule:
            st.warning(f"No lessons found for level **{level_key}**.")
            # Removed st.stop() so downstream sections (e.g., class board) can still render

        # ---- Search ----
        query = st.text_input("🔍 Search for topic, chapter, grammar, day, or anything…")
        search_terms = [q for q in query.strip().lower().split() if q] if query else []

        if search_terms:
            matches = [(i, d) for i, d in enumerate(schedule) if filter_matches(d, search_terms)]
            if not matches:
                st.warning("No matching lessons. Try simpler terms or check spelling.")
                # Removed st.stop() so downstream sections (e.g., class board) can still render

            labels = []
            for _, d in matches:
                title = highlight_terms(f"Day {d['day']}: {d['topic']}", search_terms)
                grammar = highlight_terms(d.get("grammar_topic", ""), search_terms)
                labels.append(f"{title}  {'<span style=\"color:#007bff\">['+grammar+']</span>' if grammar else ''}")

            st.markdown("<span style='font-weight:700; font-size:1rem;'>Lessons:</span>", unsafe_allow_html=True)
            sel = st.selectbox(
                "Lesson",
                list(range(len(matches))),
                format_func=lambda i: labels[i],
                key="course_search_sel",
                label_visibility="collapsed",
            )
            idx = matches[sel][0]
        else:
            st.markdown("<span style='font-weight:700; font-size:1rem;'>Choose your lesson/day:</span>", unsafe_allow_html=True)
            qp_day = _qp_get_first("day", "")
            default_idx = 0
            if schedule:
                try:
                    day_val = int(qp_day)
                    min_day = schedule[0]["day"]
                    max_day = schedule[-1]["day"]
                    day_val = max(min_day, min(day_val, max_day))
                    default_idx = next((i for i, d in enumerate(schedule) if d["day"] == day_val), 0)
                except Exception:
                    pass
            idx = st.selectbox(
                "Lesson selection",
                list(range(len(schedule))),
                index=default_idx,
                format_func=lambda i: f"Day {schedule[i]['day']} - {schedule[i]['topic']} (Chapter {schedule[i].get('chapter', '?')})",
                label_visibility="collapsed",
            )

        st.divider()

        # ---- Progress ----
        total = len(schedule)
        done = idx + 1
        pct = int(done / total * 100) if total else 0
        st.progress(pct)
        st.markdown(f"**You’ve loaded {done} / {total} lessons ({pct}%)**")
        st.divider()

        # ---- Lesson info ----
        info = schedule[idx]
        lesson_key = lesson_key_build(
            student_level,
            info.get("day", 0),
            info.get("chapter", ""),
        )
        draft_key = f"draft_{lesson_key}"
        st.session_state["coursebook_draft_key"] = draft_key
        chapter = info.get("chapter")
        title_txt = f"Day {info['day']}: {info['topic']}"
        st.markdown(
            f"### {highlight_terms(title_txt, search_terms)} (Chapter {chapter or '?'})",
            unsafe_allow_html=True,
        )
        if info.get("grammar_topic"):
            st.markdown(f"**🔤 Grammar Focus:** {highlight_terms(info['grammar_topic'], search_terms)}", unsafe_allow_html=True)

        if _coerce_day(info.get("day")) == 0:
            render_day_zero_onboarding(info, schedule, idx)
        elif info.get("goal") or info.get("instruction"):
            st.info(
                f"🎯 **Goal:** {info.get('goal','')}\n\n"
                f"📝 **Instruction:** {info.get('instruction','')}"
            )

        # ---- Class discussion count & link ----
        student_row = st.session_state.get("student_row") or {}
        _class_name_clean, class_name_lookup = _resolve_class_name(
            student_row.get("ClassName", ""),
            level=student_level,
        )
        class_name = class_name_lookup

        if class_name and chapter:
            board_base = (
                db.collection("class_board")
                .document(student_level)
                .collection("classes")
                .document(class_name)
                .collection("posts")
            )
            post_count = sum(
                1
                for _ in board_base.where(
                    filter=FieldFilter("chapter", "==", chapter)
                ).stream()
            )
            link_key = CLASS_DISCUSSION_LINK_TMPL.format(chapter=chapter)
            count_txt = f" ({post_count})" if post_count else ""
            st.info(
                f"📣 {CLASS_DISCUSSION_PROMPT} "
                f"{CLASS_DISCUSSION_LABEL}{count_txt}. "
                f"{CLASS_DISCUSSION_REMINDER}"
            )

            def _launch_class_thread(chap: str) -> None:
                current_row = st.session_state.get("student_row") or {}
                if isinstance(current_row, dict):
                    updated_row = dict(current_row)
                else:
                    try:
                        updated_row = dict(current_row)
                    except Exception:
                        updated_row = {}
                updated_row["ClassName"] = class_name_lookup
                st.session_state["student_row"] = updated_row
                go_class_thread(chap)

            st.button(
                CLASS_DISCUSSION_LABEL,
                key=link_key,
                on_click=_launch_class_thread,
                args=(chapter,),
            )
            if post_count == 0:
                st.caption("No posts yet. Clicking will show the full board.")
        elif not class_name:
            st.error(
                "This class discussion board is unavailable. Select another "
                "classroom tab and return, or log out and back in to refresh "
                "your roster."
            )
        else:
            st.warning("Missing chapter for discussion board.")

        st.divider()

        # ---------- mini-tabs inside Course Book ----------
        if "coursebook_page" not in st.session_state:
            st.session_state["coursebook_page"] = "Overview"
        if "coursebook_prev_page" not in st.session_state:
            st.session_state["coursebook_prev_page"] = st.session_state["coursebook_page"]
        def on_coursebook_page_change() -> None:
            prev = st.session_state.get("coursebook_prev_page")
            curr = st.session_state.get("coursebook_page")
            if prev in {"Assignment", "Submit"}:
                draft_key = st.session_state.get("coursebook_draft_key")
                code = (
                    st.session_state.get("student_code")
                    or (st.session_state.get("student_row") or {}).get("StudentCode", "")
                )
                if draft_key and code:
                    last_val_key, *_ = _draft_state_keys(draft_key)
                    if st.session_state.get(draft_key, "") != st.session_state.get(last_val_key, ""):
                        save_now(draft_key, code)
            st.session_state["coursebook_prev_page"] = curr

        student_row = st.session_state.get("student_row", {})
        
        coursebook_section = st.radio(
            "Section",
            ["Overview", "Assignment", "Submit"],
            key="coursebook_page",
            on_change=on_coursebook_page_change,
        )

        # OVERVIEW
        if coursebook_section == "Overview":
        
            with st.expander("📚 Course Book & Study Recommendations", expanded=True):
                LEVEL_TIME = {"A1": 15, "A2": 25, "B1": 30, "B2": 40, "C1": 45}
                rec_time = LEVEL_TIME.get(level_key, 20)
                st.info(f"⏱️ **Recommended:** Invest about {rec_time} minutes to complete this lesson fully.")

                student_row = st.session_state.get("student_row", {})
                start_str   = student_row.get("ContractStart", "")
                parse_start = (
                    globals().get("parse_contract_start_fn")
                    or globals().get("parse_contract_start")
                )
                estimates = _compute_finish_date_estimates(start_str, total, parse_start)

                if estimates:
                    end_three = estimates[3]
                    end_two   = estimates[2]
                    end_one   = estimates[1]
                    _, content = st.columns([3, 7])
                    with content:
                        st.success(f"If you complete **three sessions per week**, you will finish by **{end_three.strftime('%A, %d %B %Y')}**.")
                        st.info(f"If you complete **two sessions per week**, you will finish by **{end_two.strftime('%A, %d %B %Y')}**.")
                        st.warning(f"If you complete **one session per week**, you will finish by **{end_one.strftime('%A, %d %B %Y')}**.")
                else:
                    _, content = st.columns([3, 7])
                    with content:
                        st.warning(
                            "❓ We couldn't load your finish date estimates. Try scrolling up, "
                            "switching to a different tab and back, or logging out and logging "
                            "in again. If the problem persists, please contact administration."
                        )

        # ASSIGNMENT (activities + resources; tolerant across A1–C1)
        elif coursebook_section == "Assignment":

            draft_text = st.session_state.get(draft_key, "")
            recovered_code = _recover_student_code(
                lesson_key=lesson_key,
                draft_text=draft_text if draft_text else None,
            )
            student_row = (
                _update_student_code_session_state(recovered_code)
                if recovered_code
                else st.session_state.get("student_row")
                or {}
            )

            if not recovered_code:
                manual_key = "assignment_student_code_input"
                st.warning(
                    "We couldn't detect your student code automatically. "
                    "Enter it below so your work is linked to your account."
                )
                manual_value = st.text_input(
                    "Student code",
                    key=manual_key,
                    placeholder="e.g. KWAME123",
                    help="This appears on your student ID card or welcome email.",
                )
                manual_code = _safe_lower(manual_value)
                if manual_code and manual_code.lower() != "demo001":
                    student_row = _update_student_code_session_state(manual_code)
                    st.caption("Student code saved. You're all set to continue.")

            # ---------- helpers ----------
            def _as_list(x):
                if not x: return []
                return x if isinstance(x, list) else [x]

            def _is_url(u: str) -> bool:
                try:
                    p = urlparse(str(u))
                    return p.scheme in ("http", "https") and bool(p.netloc)
                except Exception:
                    return False

            def _dedup(seq):
                out, seen = [], set()
                for s in seq:
                    if s and s not in seen:
                        seen.add(s); out.append(s)
                return out

            def _canon_video(u: str) -> str:
                """Stable id for a video url (YouTube => yt:ID, else normalized url)."""
                if not u:
                    return ""
                try:
                    p = urlsplit(u)
                    host = (p.netloc or "").lower().replace("www.", "")
                    if "youtube.com" in host:
                        q = parse_qs(p.query or "")
                        vid = (q.get("v", [""])[0] or "").strip()
                        return f"yt:{vid}" if vid else u.strip().lower()
                    if "youtu.be" in host:
                        vid = (p.path or "/").strip("/").split("/")[0]
                        return f"yt:{vid}" if vid else u.strip().lower()
                    return u.strip().lower()
                except Exception:
                    return str(u).strip().lower()

            def pick_sections(day_info: dict):
                """Find any section keys present for this lesson across levels."""
                candidates = [
                    ("lesen_hören",        "Lesen & Hören",        "📚"),
                    ("lesen_hoeren",       "Lesen & Hören",        "📚"),
                    ("lesenhoeren",        "Lesen & Hören",        "📚"),
                    ("lesen",              "Lesen",                "📖"),
                    ("hören",              "Hören",                "🎧"),
                    ("hoeren",             "Hören",                "🎧"),
                    ("schreiben_sprechen", "Schreiben & Sprechen", "📝"),
                    ("sprechen_schreiben", "Schreiben & Sprechen", "📝"),
                    ("sprechen",           "Sprechen",             "🗣️"),
                    ("schreiben",          "Schreiben",            "✍️"),
                ]
                found = []
                for key, title, icon in candidates:
                    if day_info.get(key):
                        found.append((key, title, icon))
                return found

            def render_section_any(day_info, key, title, icon, seen_videos: set):
                content = day_info.get(key)
                if not content:
                    return

                if isinstance(content, dict):
                    items = [content]
                elif isinstance(content, list):
                    items = [c for c in content if isinstance(c, dict)]
                    if len(items) != len(content):
                        logging.warning("Expected dict elements in '%s' list, skipping non-dict entries", key)
                else:
                    logging.warning("Expected dict or list for '%s', got %s", key, type(content).__name__)
                    return

                st.markdown(f"#### {icon} {title}")
                for idx_part, part in enumerate(items):
                    chapter = part.get('chapter', '')
                    video = part.get('video')
                    youtube_link = part.get('youtube_link')
                    grammarbook_link = part.get('grammarbook_link')
                    workbook_link = part.get('workbook_link')
                    extras = part.get('extra_resources')

                    if len(items) > 1:
                        st.markdown(
                            f"###### {icon} Part {idx_part+1} of {len(items)}: Chapter {chapter}"
                        )
                    else:
                        st.markdown(f"###### {icon} Chapter {chapter}")
                    # videos (embed once)
                    for maybe_vid in [video, youtube_link]:
                        if _is_url(maybe_vid):
                            cid = _canon_video(maybe_vid)
                            if cid not in seen_videos:
                                st.markdown(
                                    f"[🎬 Lecture Video on YouTube]({maybe_vid})"
                                )
                                seen_videos.add(cid)
                    # links/resources inline
                    if grammarbook_link:
                        st.markdown(f"- [📘 Grammar Book (Notes)]({grammarbook_link})")
                        st.markdown(
                            "<em>Reminder:</em> 📘 gives you the grammar from today's lecture; 📒 is the assignment you must complete.",
                            unsafe_allow_html=True,
                        )
                    if workbook_link:
                        st.markdown(f"- [📒 Workbook (Assignment)]({workbook_link})")
                        with st.expander("📖 Dictionary"):
                            render_vocab_lookup(
                                f"{key}-{idx_part}",
                                f"Day {day_info.get('day')} Chapter {chapter}",
                            )
                        render_assignment_reminder()
                    if extras:
                        for ex in _as_list(extras):
                            st.markdown(f"- [🔗 Extra]({ex})")

            # ---------- YOUR WORK (tolerant across levels; embeds each video at most once) ----------
            st.markdown("### 🧪 Your Work")
            seen_videos = set()
            sections = pick_sections(info)

            if sections:
                for key, title, icon in sections:
                    render_section_any(info, key, title, icon, seen_videos)
            else:
                # Fallback: show top-level resources even if there are no section keys
                showed = False
                if info.get("video"):
                    cid = _canon_video(info["video"])
                    if cid not in seen_videos:
                        st.markdown(
                            f"[🎬 Lecture Video on YouTube]({info['video']})"
                        )
                        seen_videos.add(cid)
                    showed = True
                if info.get("grammarbook_link"):
                    st.markdown(f"- [📘 Grammar Book (Notes)]({info['grammarbook_link']})")
                    showed = True
                if info.get("workbook_link"):
                    st.markdown(f"- [📒 Workbook (Assignment)]({info['workbook_link']})")
                    with st.expander("📖 Dictionary"):
                        render_vocab_lookup(
                            f"fallback-{info.get('day', '')}",
                            f"Day {info.get('day')} Chapter {info.get('chapter', '')}",
                        )
                    render_assignment_reminder()
                    showed = True
                for ex in _as_list(info.get("extra_resources")):
                    st.markdown(f"- [🔗 Extra]({ex})")
                    showed = True

                if not showed:
                    st.info(
                        "No activity sections or links found for this lesson. Check the lesson data for A2/B1 key names."
                    )

            # --- quick access to translators and language support ---
            translator_col, support_col = st.columns([1, 1.6])
            with translator_col:
                st.markdown(
                    "[🌐 DeepL Translator](https://www.deepl.com/translator) &nbsp; | &nbsp; "
                    "[🌐 Google Translate](https://translate.google.com)",
                    unsafe_allow_html=True,
                )
            with support_col:
                render_lesson_language_support(info, level_key)

            # ---------- Build a clean downloadable bundle of links (no on-page repetition) ----------
            st.divider()
            st.markdown("### 📎 Lesson Links — Download")

            # Collect links (top-level + nested)
            resources = {"Grammar Notes": [], "Workbook": [], "Videos": [], "Extras": []}

            def _add(kind, val):
                for v in _as_list(val):
                    if _is_url(v):
                        resources[kind].append(v)

            # top-level
            _add("Videos", info.get("video"))
            _add("Grammar Notes", info.get("grammarbook_link"))
            _add("Workbook", info.get("workbook_link"))
            _add("Extras", info.get("extra_resources"))

            # nested: include whatever sections exist for this lesson
            for section_key, _, _ in sections or []:
                for part in _as_list(info.get(section_key)):
                    if not isinstance(part, dict):
                        continue
                    _add("Videos", [part.get("video"), part.get("youtube_link")])
                    _add("Grammar Notes", part.get("grammarbook_link"))
                    _add("Workbook", part.get("workbook_link"))
                    _add("Extras", part.get("extra_resources"))

            # dedupe + remove videos already embedded above
            for k in list(resources.keys()):
                resources[k] = _dedup(resources[k])

            # If nothing remains after filtering, don't show anything
            if not any(resources.values()):
                st.caption("All lesson links are already shown above. No extra links to download.")
            else:
                # Prepare TXT bundle
                lesson_header = f"Level: {level_key} | Day: {info.get('day','?')} | Chapter: {info.get('chapter','?')} | Topic: {info.get('topic','')}"
                parts_txt = [lesson_header, "-" * len(lesson_header)]
                for title, key_name in [("📘 Grammar Notes", "Grammar Notes"),
                                        ("📒 Workbook", "Workbook"),
                                        ("🎥 Videos", "Videos"),
                                        ("🔗 Extras", "Extras")]:
                    if resources[key_name]:
                        parts_txt.append(title)
                        parts_txt.extend([f"- {u}" for u in resources[key_name]])
                        parts_txt.append("")
                bundle_txt = "\n".join(parts_txt).strip() + "\n"

                temp_path = st.session_state.get("links_temp_path")
                if not temp_path or not os.path.exists(temp_path):
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".txt") as tmp:
                        tmp.write(bundle_txt.encode("utf-8"))
                        temp_path = tmp.name
                    st.session_state["links_temp_path"] = temp_path

                cdl1, cdl2 = st.columns([1, 1])
                with cdl1:
                    file_obj = open(temp_path, "rb")
                    clicked = st.download_button(
                        "⬇️ Download lesson links (TXT)",
                        data=file_obj,
                        file_name=f"lesson_links_{level_key}_day{info.get('day','')}.txt",
                        mime="text/plain",
                        key="dl_links_txt",
                    )
                    file_obj.close()
                    if clicked:
                        try:
                            os.remove(temp_path)
                        finally:
                            st.session_state.pop("links_temp_path", None)

                

            with st.expander("📚 Study Resources"):
                if _is_url(info.get("video")):
                    st.video(info["video"])
                elif info.get("video"):
                    st.markdown(
                        f"[🎬 Lecture Video on YouTube]({info['video']})"
                    )
                    
                if _is_url(info.get("grammarbook_link")):
                    render_link("📘 Grammar Book (Notes)", info["grammarbook_link"])

                render_link("📗 Dictionary", "https://dict.leo.org/german-english")


            st.markdown("#### 🎬 Video of the Day for Your Level")
            playlist_ids = get_playlist_ids_for_level(level_key)
            fetch_videos = fetch_youtube_playlist_videos
            playlist_id = random.choice(playlist_ids) if playlist_ids else None

            reflection_prompts = [
                "📝 After watching, jot down two new words or phrases you heard.",
                "🗣️ Pause the video and repeat a key sentence aloud to practice pronunciation.",
                "💬 Summarize the main idea of the video in one or two simple sentences.",
            ]

            if playlist_id:
                if st.button("🔄 Refresh videos", key=f"refresh_vod_{level_key}"):
                    st.cache_data.clear()
                    st.session_state["need_rerun"] = True
                st.caption(
                    "Click 'Refresh videos' to clear cached playlist data and reload from YouTube if results look out of date."
                )
                try:
                    video_list = fetch_videos(playlist_id)
                except Exception:
                    video_list = []
                if video_list:
                    today_idx = date.today().toordinal() % len(video_list)
                    video = video_list[today_idx]
                    st.markdown(f"**{video['title']}**")
                    st.video(video['url'])
                    description = video.get("description")
                    if description:
                        st.caption(description)
                    else:
                        st.caption(random.choice(reflection_prompts))
                else:
                    st.info("No videos found for your level’s playlist. Check back soon!")
            else:
                st.info("No playlist found for your level yet. Stay tuned!")
            st.markdown("**The End**")


        # SUBMIT
        elif coursebook_section == "Submit":
            submission_disabled_reason = _submission_block_reason(info, schedule)
            submission_disabled = bool(submission_disabled_reason)

            st.markdown("### ✅ Submit Your Assignment")
            st.markdown(
                f"""
                <div style="box-sizing:border-box;padding:14px 16px;border-radius:10px;
                            background:#f0f9ff;border:1px solid #bae6fd;margin:6px 0 12px 0;">
                  <div style="font-size:1.05rem;">
                    📌 <b>You're on:</b> Level <b>{student_level}</b> • Day <b>{info['day']}</b> • Chapter <b>{info['chapter']}</b>
                  </div>
                  <div style="color:#0369a1;margin-top:4px;">
                    Make sure this matches the assignment your tutor set. If not, change the lesson from the dropdown above.
                  </div>
                </div>
                """,
                unsafe_allow_html=True
            )

            draft_text = st.session_state.get(draft_key, "")
            recovered_code = _recover_student_code(
                lesson_key=lesson_key,
                draft_text=draft_text if draft_text else None,
            )
            student_row = (
                _update_student_code_session_state(recovered_code)
                if recovered_code
                else st.session_state.get("student_row")
                or {}
            )
            code = _safe_str(student_row.get("StudentCode"), "demo001")
            name_default = _safe_str(student_row.get("Name"))
            missing_code = (not code) or (code.lower() == "demo001")

            locked_key = f"{lesson_key}_locked"
            needs_resubmit_key = f"{lesson_key}__needs_resubmit"

            code_input_key = "submit_student_code_input"
            if missing_code:
                _show_missing_code_warning(
                    name=name_default,
                    level=student_level,
                    lesson_info=info,
                )
                manual_value = st.text_input(
                    "Enter your student code to continue",
                    key=code_input_key,
                    placeholder="e.g. KWAME123",
                    help="This appears on your student ID card or welcome email.",
                )
                manual_code = _safe_lower(manual_value)
                if manual_code and manual_code.lower() != "demo001":
                    student_row = _update_student_code_session_state(manual_code)
                    code = manual_code
                    missing_code = False
                    st.success("Student code saved. You can now submit your work.")

                resubmit_should_show = (
                    st.session_state.get(locked_key, False)
                    or bool(st.session_state.get(f"{lesson_key}__receipt"))
                    or (st.session_state.get(needs_resubmit_key) is not None)
                )
                if resubmit_should_show:
                    stored_row = st.session_state.get("student_row") or {}
                    resubmit_name = name_default or _safe_str(stored_row.get("Name"))
                    candidate_code = (
                        st.session_state.get("student_code")
                        or stored_row.get("StudentCode")
                        or manual_code
                    )
                    render_resubmit_email_cta(
                        lesson_info=info,
                        student_name=resubmit_name,
                        student_code=candidate_code,
                    )
            else:
                st.session_state[code_input_key] = code

            if not missing_code:
                if submission_disabled:
                    st.info(submission_disabled_reason)

                st.session_state["student_code"] = code
                chapter_name = f"{info['chapter']} – {info.get('topic', '')}"

                name = st.text_input("Name", value=student_row.get('Name', ''))
                email = st.text_input("Email", value=student_row.get('Email', ''))

                db_locked = is_locked(student_level, code, lesson_key)
                success_notice_key = f"{lesson_key}__submit_success_notice"
                if db_locked:
                    st.session_state[locked_key] = True
                locked = db_locked or st.session_state.get(locked_key, False)
                locked_ui = locked or submission_disabled
                submit_in_progress_key = f"{lesson_key}_submit_in_progress"

                # ---------- save previous lesson on switch + force hydrate for this one ----------
                prev_active_key = st.session_state.get("__active_draft_key")
                if prev_active_key and prev_active_key != draft_key:
                    try:
                        prev_text = st.session_state.get(prev_active_key, "")
                        save_draft_to_db(code, prev_active_key, prev_text)
                    except Exception:
                        pass  # never block UI
                    # ensure the newly selected lesson re-hydrates from cloud
                    st.session_state.pop(f"{draft_key}__hydrated_v2", None)
                st.session_state["__active_draft_key"] = draft_key

                # ---------- Decide what to show (guarded hydration) ----------
                pending_key      = f"{draft_key}__pending_reload"
                pending_text_key = f"{draft_key}__reload_text"
                pending_ts_key   = f"{draft_key}__reload_ts"
                hydrated_key     = f"{draft_key}__hydrated_v2"  # only hydrate once per lesson

                last_val_key, last_ts_key, saved_flag_key, saved_at_key = _draft_state_keys(draft_key)

                # 1) If a forced reload was requested, apply it BEFORE widget creation
                if st.session_state.get(pending_key):
                    cloud_text = st.session_state.pop(pending_text_key, "")
                    cloud_ts   = st.session_state.pop(pending_ts_key, None)
                    st.session_state[pending_key] = False

                    st.session_state[draft_key]      = cloud_text or ""
                    st.session_state[last_val_key]   = st.session_state[draft_key]
                    st.session_state[last_ts_key]    = time.time()
                    st.session_state[saved_flag_key] = True
                    st.session_state[saved_at_key]   = (cloud_ts or datetime.now(_timezone.utc))
                    st.session_state[hydrated_key]   = True

                    try:
                        when = (cloud_ts.strftime('%Y-%m-%d %H:%M') + " UTC") if cloud_ts else "now"
                    except Exception:
                        when = "now"
                    st.info(f"Reloaded cloud draft (saved {when}).")

                else:
                    # 2) If a SUBMISSION exists, always enforce it (locked) on every run
                    latest = fetch_latest(student_level, code, lesson_key)
                    if latest and (latest.get("answer", "") is not None):
                        sub_txt = latest.get("answer", "") or ""
                        sub_ts  = latest.get("updated_at")

                        st.session_state[draft_key]      = sub_txt
                        st.session_state[last_val_key]   = sub_txt
                        st.session_state[last_ts_key]    = time.time()
                        st.session_state[saved_flag_key] = True
                        st.session_state[saved_at_key]   = (sub_ts or datetime.now(_timezone.utc))
                        st.session_state[locked_key]     = True
                        st.session_state[hydrated_key]   = True
                        locked = True  # enforce read-only

                        when = f"{sub_ts.strftime('%Y-%m-%d %H:%M')} UTC" if sub_ts else ""
                        st.success(f"Showing your submitted answer. {('Updated ' + when) if when else ''}")

                    else:
                        # 3) No submission → hydrate ONCE from cloud; after that, never clobber local typing
                        if not st.session_state.get(hydrated_key, False):
                            cloud_text, cloud_ts = load_draft_meta_from_db(code, draft_key)
                            if cloud_text is not None:
                                st.session_state[draft_key]      = cloud_text or ""
                                st.session_state[last_val_key]   = st.session_state[draft_key]
                                st.session_state[last_ts_key]    = time.time()
                                st.session_state[saved_flag_key] = True
                                st.session_state[saved_at_key]   = (cloud_ts or datetime.now(_timezone.utc))
                            else:
                                st.session_state.setdefault(draft_key, "")
                                st.session_state.setdefault(last_val_key, "")
                                st.session_state.setdefault(last_ts_key, time.time())
                                st.session_state.setdefault(saved_flag_key, False)
                                st.session_state.setdefault(saved_at_key, None)

                            st.session_state[hydrated_key] = True

                            if cloud_text:
                                when = f"{cloud_ts.strftime('%Y-%m-%d %H:%M')} UTC" if cloud_ts else ""
                                st.info(f"💾 Restored your saved draft. {('Last saved ' + when) if when else ''}")
                            elif not submission_disabled:
                                st.caption("Start typing your answer.")
                        else:
                            # If 'hydrated' but local is empty, pull cloud once
                            if not st.session_state.get(draft_key, "") and not locked_ui:
                                ctext, cts = load_draft_meta_from_db(code, draft_key)
                                if ctext:
                                    st.session_state[draft_key]      = ctext
                                    st.session_state[last_val_key]   = ctext
                                    st.session_state[last_ts_key]    = time.time()
                                    st.session_state[saved_flag_key] = True
                                    st.session_state[saved_at_key]   = (cts or datetime.now(_timezone.utc))

                st.subheader("✍️ Your Answer")

                locked_warning_message = None
                needs_resubmit = False
                if locked:
                    locked_warning_message = (
                        "This box is locked because you have already submitted your work."
                    )
                    needs_resubmit = st.session_state.get(needs_resubmit_key)
                    if needs_resubmit is None:
                        answer_text = st.session_state.get(draft_key, "").strip()
                        MIN_WORDS = 20
                        needs_resubmit = len(answer_text.split()) < MIN_WORDS
                    needs_resubmit = bool(needs_resubmit)
                    st.session_state[needs_resubmit_key] = needs_resubmit

                    render_resubmit_email_cta(
                        lesson_info=info,
                        student_name=(name or name_default),
                        student_code=code,
                    )
                show_resubmit_hint = locked and bool(needs_resubmit)

                # ---------- Editor (save on blur + debounce) ----------
                st.text_area(
                    "Type all your answers here",
                    height=500,
                    key=draft_key,              # value already hydrated in st.session_state[draft_key]
                    on_change=save_now,         # guaranteed save on blur/change
                    args=(draft_key, code),
                    disabled=locked_ui,
                    help="Autosaves on blur and in the background while you type."
                )
                render_umlaut_pad(draft_key, context=f"coursebook_{lesson_key}", disabled=locked_ui)

                # Debounced autosave (safe so empty first-render won't wipe a non-empty cloud draft)
                current_text = st.session_state.get(draft_key, "")
                last_val = st.session_state.get(last_val_key, "")
                if not locked_ui and (current_text.strip() or not last_val.strip()):
                    autosave_maybe(code, draft_key, current_text, min_secs=2.0, min_delta=12, locked=locked_ui)

                # ---------- Manual save + last saved time + safe reload ----------
                csave1, csave2, csave3 = st.columns([1, 1, 1])

                with csave1:
                    if st.button("💾 Save Draft now", disabled=locked_ui):
                        save_draft_to_db(code, draft_key, current_text)
                        st.session_state[last_val_key]   = current_text
                        st.session_state[last_ts_key]    = time.time()
                        st.session_state[saved_flag_key] = True
                        st.session_state[saved_at_key]   = datetime.now(_timezone.utc)
                        st.success("Draft saved.")

                with csave2:
                    ts = st.session_state.get(saved_at_key)
                    if ts:
                        st.caption("Last saved: " + ts.strftime("%Y-%m-%d %H:%M") + " UTC")
                    else:
                        st.caption("No local save yet")

                with csave3:
                    # Current draft text
                    draft_txt = st.session_state.get(draft_key, "") or ""

                    # Last-saved timestamp (for header)
                    _, _, _, saved_at_key = _draft_state_keys(draft_key)
                    ts = st.session_state.get(saved_at_key)
                    when = (
                        ts.astimezone(_timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
                        if ts else datetime.now(_timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
                    )

                    # Strip any previous backup header the student may have pasted back
                    def _strip_old_header(txt: str) -> str:
                        if not txt:
                            return ""
                        # Remove ONE leading “Falowen — Draft Backup … ======” block if present
                        pattern = r"(?s)\AFalowen\s+—\s+Draft\s+Backup.*?\n[-=]{8,}\n\n"
                        return re.sub(pattern, "", txt, count=1)

                    clean_body = (_strip_old_header(draft_txt).rstrip() + "\n")

                    # Build a simple, single header
                    header_lines = [
                        "Falowen — Draft Backup",
                        f"Level: {student_level}  •  Day: {info['day']}  •  Chapter: {info.get('chapter','')}",
                        f"Student: {name}  •  Code: {code}",
                        f"Saved (UTC): {when}",
                        "=" * 56,
                        ""  # blank line before body
                    ]
                    header = "\n".join(header_lines)

                    # Safe filename
                    safe_chapter = re.sub(r"[^A-Za-z0-9_.-]+", "_", str(info.get("chapter", "")))
                    fname = f"falowen_draft_{student_level}_day{info['day']}_{safe_chapter}.txt"

                    st.download_button(
                        "⬇️ Download draft (TXT)",
                        data=(header + clean_body).encode("utf-8"),
                        file_name=fname,
                        mime="text/plain",
                        help="Save a clean backup of your current draft"
                    )

                if show_resubmit_hint:
                    st.info(
                        "🔒 This box is locked because you already submitted. Scroll up to use the resubmit email link if you need to send an update."
                    )

                with st.expander("📌 How to Submit", expanded=False):
                    st.markdown(f"""
                        1) Check you’re on the correct page: **Level {student_level} • Day {info['day']} • Chapter {info['chapter']}**.  
                        2) Tick the two confirmations below.  
                        3) Click **Confirm & Submit**.  
                        4) Your box will lock (read-only).  
                        _You’ll get an **email** when it’s marked. See **Results & Resources** for scores & feedback._
                    """)

                col1, col2 = st.columns([1, 1.2])
                success_notice = st.session_state.pop(success_notice_key, None)
                with col1:
                    st.markdown("#### 🧾 Finalize")
                    if success_notice:
                        success_msg = success_notice.get("message")
                        caption_msg = success_notice.get("caption")
                        if success_msg:
                            st.success(success_msg)
                        if caption_msg:
                            st.caption(caption_msg)
                    if locked_warning_message:
                        st.warning(locked_warning_message)
                    confirm_final = st.checkbox(
                        f"I confirm this is my complete work for Level {student_level} • Day {info['day']} • Chapter {info['chapter']}.",
                        key=f"confirm_final_{lesson_key}",
                        disabled=locked_ui
                    )
                    confirm_lock = st.checkbox(
                        "I understand it will be locked after I submit.",
                        key=f"confirm_lock_{lesson_key}",
                        disabled=locked_ui
                    )
                    can_submit = (
                        confirm_final
                        and confirm_lock
                        and (not locked_ui)
                    )

                with col2:
                    st.markdown("#### 🔍 Live Preview")
                    preview_text = st.session_state.get(draft_key, "") or ""
                    word_count = len(preview_text.split()) if preview_text else 0
                    if preview_text.strip():
                        st.caption(f"Word count: {word_count}")
                        preview_box = (
                            "<div style=\"max-height:360px;overflow:auto;padding:12px;border:1px solid #e2e8f0;"
                            "border-radius:8px;background:#f8fafc;white-space:pre-wrap;line-height:1.5;\">"
                            f"{html.escape(preview_text)}"
                            "</div>"
                        )
                        st.markdown(preview_box, unsafe_allow_html=True)
                    else:
                        st.info("Draft preview will appear here once you start typing.")
                        st.caption("Word count: 0")

                submit_in_progress = st.session_state.get(submit_in_progress_key, False)

                with col1:
                    if st.button(
                        "✅ Confirm & Submit",
                        type="primary",
                        disabled=(not can_submit) or submit_in_progress,
                    ):
                        st.session_state[submit_in_progress_key] = True
                        
                        try:

                            # 1) Try to acquire the lock first
                            got_lock = acquire_lock(student_level, code, lesson_key)

                            # If lock exists already, check whether a submission exists; if yes, reflect lock and rerun.
                            if not got_lock:
                                if has_existing_submission(student_level, code, lesson_key):
                                    st.session_state[locked_key] = True
                                    st.warning("You have already submitted this assignment. It is locked.")
                                    refresh_with_toast()
                                else:
                                    st.info("Found an old lock without a submission — recovering and submitting now…")

                            posts_ref = db.collection("submissions").document(student_level).collection("posts")

                            # 2) Pre-create doc (avoids add() tuple-order mismatch)
                            doc_ref = posts_ref.document()  # auto-ID now available
                            short_ref = f"{doc_ref.id[:8].upper()}-{info['day']}"

                            payload = {
                                "student_code": code,
                                "student_name": name or "Student",
                                "student_email": email,
                                "level": student_level,
                                "day": info["day"],
                                "chapter": chapter_name,
                                "lesson_key": lesson_key,
                                "answer": (st.session_state.get(draft_key, "") or "").strip(),
                                "status": "submitted",
                                "receipt": short_ref,  # persist receipt immediately
                                "created_at": firestore.SERVER_TIMESTAMP,
                                "updated_at": firestore.SERVER_TIMESTAMP,
                                "version": 1,
                            }

                            saved_ok = False

                            # Archive the draft so it won't rehydrate again (drafts_v2)
                            try:

                                doc_ref.set(payload)  # write the submission
                                saved_ok = True
                                st.caption(f"Saved to: `{doc_ref.path}`")  # optional debug
                            except Exception as e:
                                st.error(f"Could not save submission: {e}")

                            if saved_ok:
                                # 3) Success: lock UI, remember receipt, archive draft, notify, rerun
                                st.session_state[locked_key] = True
                                st.session_state[f"{lesson_key}__receipt"] = short_ref

                                success_msg = (
                                    f"Well done, {name or 'Student'}! Remember the pass mark is 60, "
                                    "and if you score below that you must revisit this Submit page to try again."
                                )
                                caption_msg = (
                                    f"Receipt: `{short_ref}` • Marks will arrive by email and via "
                                    "Telegram from @falowenbot. See **Results & Resources** for scores & feedback."
                                )
                                st.success(success_msg)
                                st.caption(caption_msg)
                                st.session_state[success_notice_key] = {
                                    "message": success_msg,
                                    "caption": caption_msg,
                                }
                                row = st.session_state.get("student_row") or {}
                                tg_subscribed = bool(
                                    row.get("TelegramChatID")
                                    or row.get("telegram_chat_id")
                                    or row.get("Telegram")
                                    or row.get("telegram")
                                )
                                if not tg_subscribed:
                                    try:
                                        tg_subscribed = has_telegram_subscription(code)
                                    except Exception:
                                        tg_subscribed = False
                                if tg_subscribed:
                                    st.info("You'll also receive a Telegram notification when your score is posted.")
                                else:
                                    with st.expander("🔔 Subscribe to Telegram notifications", expanded=False):
                                        st.markdown(
                                            "\n".join(
                                                [
                                                    "1. Search for **@falowenbot** on Telegram and open the chat.",
                                                    "2. Tap **Start**, then follow the prompts to connect your account so you can receive your marks.",
                                                    "3. To deactivate: send `/stop`",
                                                ]
                                            )
                                        )
                                answer_text = st.session_state.get(draft_key, "").strip()
                                MIN_WORDS = 20

                                st.session_state[f"{lesson_key}__needs_resubmit"] = (
                                    len(answer_text.split()) < MIN_WORDS
                                )


                                # Archive the draft so it won't rehydrate again (drafts_v2)
                                try:
                                    _draft_doc_ref(student_level, lesson_key, code).set(
                                        {"status": "submitted", "archived_at": firestore.SERVER_TIMESTAMP}, merge=True
                                    )
                                except Exception:
                                    pass

                                # Notify Slack (best-effort)
                                webhook = get_slack_webhook()
                                if webhook:
                                    notify_slack_submission(
                                        webhook_url=webhook,
                                        student_name=name or "Student",
                                        student_code=code,
                                        level=student_level,
                                        day=info["day"],
                                        chapter=chapter_name,
                                        receipt=short_ref,
                                        preview=st.session_state.get(draft_key, "")
                                    )

                                # Rerun so hydration path immediately shows locked view
                                refresh_with_toast(
                                    "Submission saved! Remember you need at least 60 points to pass."
                                )
                            else:
                                # 4) Failure: remove the lock doc so student can retry cleanly
                                try:
                                    db.collection("submission_locks").document(lock_id(student_level, code, lesson_key)).delete()
                                except Exception:
                                    pass
                                st.warning("Submission not saved. Please fix the issue and try again.")
                        finally:
                            st.session_state[submit_in_progress_key] = False
                            st.markdown("**The End**")



    if cb_subtab == "🧑‍🏫 Classroom":
        # --- Classroom banner (top of subtab) ---
        st.markdown(
            '''
            <div style="
                padding: 16px;
                background: #0ea5e9;
                color: #ffffff;
                border-radius: 8px;
                text-align: center;
                margin-bottom: 16px;
                box-shadow: 0 4px 6px rgba(0,0,0,0.1);
            ">
                <span style="font-size:1.8rem; font-weight:600;">🧑‍🏫 Classroom</span>
            </div>
            ''',
            unsafe_allow_html=True
        )
        st.divider()

        # ---------- Shared helpers & imports used across tabs ----------
        import math
        import os
        import io
        import re
        import json
        import hashlib
        import pandas as pd
        import requests
        from uuid import uuid4
        from datetime import datetime as _dt, timedelta as _td
        import urllib.parse as _urllib
        try:
            import streamlit.components.v1 as components
        except Exception:
            components = None

        student_row   = st.session_state.get("student_row") or {}
        student_code  = _safe_str(student_row.get("StudentCode"), "demo001")
        student_name  = _safe_str(student_row.get("Name"), "Student")
        student_level = _safe_upper(student_row.get("Level"), "A1")
        if student_code == "demo001":
            _show_missing_code_warning(name=student_name, level=student_level)
        class_name    = _safe_str(student_row.get("ClassName")) or f"{student_level} General"

        ADMINS = set()
        try:
            ADMINS = set(st.secrets["roles"]["admins"])
        except Exception:
            pass
        ADMINS |= ADMINS_BY_LEVEL.get(student_level, set())
        IS_ADMIN = student_code in ADMINS

        # ---------- slack helper (use global notify_slack if present; else env/secrets) ----------
        def _notify_slack(*parts: str):
            text = "".join(parts)
            try:
                fn = globals().get("notify_slack")
                if callable(fn):
                    try:
                        fn(text)
                        return
                    except Exception:
                        pass
                url = (os.getenv("SLACK_WEBHOOK_URL") or
                       (st.secrets.get("slack", {}).get("webhook_url", "") if hasattr(st, "secrets") else "")).strip()
                if url:
                    try:
                        requests.post(url, json={"text": text}, timeout=6)
                    except Exception:
                        pass
            except Exception:
                pass

        def _ukey(base: str) -> str:
            # unique widget key per class (prevents duplicate-key crashes)
            seed = f"{base}|{class_name}"
            return f"{base}_{hashlib.md5(seed.encode()).hexdigest()[:8]}"


        # ---------- MINI-TABS INSIDE 'CLASSROOM' (radio style) ----------
        if "classroom_page" not in st.session_state:
            st.session_state["classroom_page"] = "Calendar"
        if "classroom_prev_page" not in st.session_state:
            st.session_state["classroom_prev_page"] = st.session_state["classroom_page"]

        def on_classroom_page_change() -> None:
            prev = st.session_state.get("classroom_prev_page")
            curr = st.session_state.get("classroom_page")
            st.session_state["classroom_prev_page"] = curr

        classroom_section = st.radio(
            "Classroom section",
            [
                "Calendar",
                "Join on Zoom",
                "Members & Profile",
                "Class Notes & Q&A",
                "Attendance",
            ],
            horizontal=True,
            key="classroom_page",
            on_change=on_classroom_page_change,
        )

                # ===================== CALENDAR =====================
        if classroom_section == "Calendar":
            # Banner
            st.markdown(
                '''
                <div style="
                    padding: 12px;
                    background: #0ea5e9;
                    color: #ffffff;
                    border-radius: 8px;
                    text-align: center;
                    margin-bottom: 12px;
                    box-shadow: 0 2px 6px rgba(0,0,0,0.08);
                    font-weight: 600;
                ">
                    <span style="font-size:1.2rem;">📅 Calendar</span>
                    <div style="font-weight:500; font-size:0.98rem; margin-top:2px;">
                        Download the full course schedule or add reminders to your phone.
                    </div>
                </div>
                ''',
                unsafe_allow_html=True
            )
            st.divider()

            # Try dateutil if available; fall back gracefully.
            try:
                from dateutil import parser as _dateparse
            except Exception:
                _dateparse = None

            # -------- group schedule config (global/secrets/firestore/fallback) --------
            def _load_group_schedules():
                if not st.session_state.get("student_level"):
                    ensure_student_level()
                # 1) global
                cfg = globals().get("GROUP_SCHEDULES")
                if isinstance(cfg, dict) and cfg:
                    return cfg
                # 2) session_state
                cfg = st.session_state.get("GROUP_SCHEDULES")
                if isinstance(cfg, dict) and cfg:
                    globals()["GROUP_SCHEDULES"] = cfg
                    return cfg
                # 3) secrets
                try:
                    raw = st.secrets.get("group_schedules", None)
                    if raw:
                        cfg = json.loads(raw) if isinstance(raw, str) else raw
                        if isinstance(cfg, dict) and cfg:
                            st.session_state["GROUP_SCHEDULES"] = cfg
                            globals()["GROUP_SCHEDULES"] = cfg
                            return cfg
                except Exception:
                    pass
                # 4) Firestore (optional)
                try:
                    doc = db.collection("config").document("group_schedules").get()
                    if doc and getattr(doc, "exists", False):
                        data = doc.to_dict() or {}
                        cfg = data.get("data", data)
                        if isinstance(cfg, dict) and cfg:
                            st.session_state["GROUP_SCHEDULES"] = cfg
                            globals()["GROUP_SCHEDULES"] = cfg
                            return cfg
                except Exception:
                    pass
                    
                # 5) Shared fallback from module
                cfg = load_group_schedules()
                st.session_state["GROUP_SCHEDULES"] = cfg
                globals()["GROUP_SCHEDULES"] = cfg
                return cfg

            def _gdrive_direct_download(url: str) -> Optional[bytes]:
                if not url:
                    return None
                m = re.search(r"/file/d/([A-Za-z0-9_-]{20,})/", url) or re.search(r"[?&]id=([A-Za-z0-9_-]{20,})", url)
                file_id = m.group(1) if m else None
                if not file_id:
                    return None
                dl = f"https://drive.google.com/uc?export=download&id={file_id}"
                try:
                    r = requests.get(dl, timeout=15)
                    if r.status_code == 200 and r.content:
                        if b"uc-download-link" in r.content[:4000] and b"confirm" in r.content[:4000]:
                            return None
                        return r.content
                except Exception:
                    pass
                return None

            def _extract_text_from_pdf(pdf_bytes: bytes) -> str:
                from src.pdf_handling import extract_text_from_pdf
                return extract_text_from_pdf(pdf_bytes)

            _DATE_PATTERNS = [
                r"\b(20\d{2}-\d{2}-\d{2})\b",
                r"\b(\d{1,2}/\d{1,2}/20\d{2})\b",
                r"\b(\d{1,2}\s+(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*\s+20\d{2})\b",
                r"\b((Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*\s+\d{1,2},\s*20\d{2})\b",
            ]

            def _parse_any_date(raw: str):
                if _dateparse:
                    for dayfirst in (False, True):
                        try:
                            return _dateparse.parse(raw, dayfirst=dayfirst, fuzzy=True).date()
                        except Exception:
                            pass
                for fmt in ("%Y-%m-%d", "%d %b %Y", "%b %d, %Y", "%m/%d/%Y", "%d/%m/%Y"):
                    try:
                        return _dt.strptime(raw, fmt).date()
                    except Exception:
                        pass
                return None

            def _find_dates_in_text(txt: str):
                found = []
                if not txt:
                    return found
                for pat in _DATE_PATTERNS:
                    for m in re.finditer(pat, txt, flags=re.IGNORECASE):
                        d = _parse_any_date(m.group(1))
                        if d:
                            found.append(d)
                uniq = []
                seen = set()
                for d in sorted(found):
                    if d not in seen:
                        seen.add(d)
                        uniq.append(d)
                return uniq

            def infer_start_end_from_doc(doc_url: str):
                pdf_bytes = _gdrive_direct_download(doc_url)
                if not pdf_bytes:
                    return None, None
                text = _extract_text_from_pdf(pdf_bytes)
                dates = _find_dates_in_text(text)
                if len(dates) >= 2:
                    return dates[0], dates[-1]
                if len(dates) == 1:
                    return dates[0], None
                return None, None

            GROUP_SCHEDULES = _load_group_schedules()

            class_cfg   = GROUP_SCHEDULES.get(class_name, {})
            days        = class_cfg.get("days", [])
            time_str    = class_cfg.get("time", "")
            start_str   = class_cfg.get("start_date", "")
            end_str     = class_cfg.get("end_date", "")
            doc_url     = class_cfg.get("doc_url", "")

            start_date_obj = None
            end_date_obj   = None
            try:
                if start_str:
                    start_date_obj = _dt.strptime(start_str, "%Y-%m-%d").date()
            except Exception:
                pass
            try:
                if end_str:
                    end_date_obj = _dt.strptime(end_str, "%Y-%m-%d").date()
            except Exception:
                pass

            _inferred_start = _inferred_end = False
            if (not start_date_obj or not end_date_obj) and doc_url:
                s, e = infer_start_end_from_doc(doc_url)
                if s and not start_date_obj:
                    start_date_obj = s; _inferred_start = True
                if e and not end_date_obj:
                    end_date_obj = e; _inferred_end = True

            if not (start_date_obj and end_date_obj and isinstance(time_str, str) and time_str.strip() and days):
                st.warning("This class doesn’t have a full calendar setup yet. Please contact the office.", icon="⚠️")
            else:
                _note_bits = []
                if _inferred_start or _inferred_end:
                    _note_bits.append("dates inferred from the schedule document")
                _note = f" ({', '.join(_note_bits)})" if _note_bits else ""
                st.info(
                    f"**Course period:** {start_date_obj.strftime('%d %b %Y')} → {end_date_obj.strftime('%d %b %Y')}{_note}",
                    icon="📅",
                )

                _WKD_ORDER = ["MO","TU","WE","TH","FR","SA","SU"]
                _FULL_TO_CODE = {
                    "monday":"MO","tuesday":"TU","wednesday":"WE","thursday":"TH","friday":"FR","saturday":"SA","sunday":"SU",
                    "mon":"MO","tue":"TU","tues":"TU","wed":"WE","thu":"TH","thur":"TH","thurs":"TH","fri":"FR","sat":"SA","sun":"SU"
                }
                DEFAULT_AMPM = "pm"

                def _normalize_time_groups(s: str) -> str:
                    s = (s or "").strip()
                    s = s.replace("–", "-").replace("—", "-")
                    s = re.sub(
                        r"(?i)\b(mon|tue|tues|wed|thu|thur|thurs|fri|sat|sun|monday|tuesday|wednesday|thursday|friday|saturday|sunday)\s*(\d)",
                        r"\1: \2",
                        s,
                    )
                    return s

                def _to_24h(h, m, ampm):
                    h = int(h); m = int(m); ap = (ampm or "").lower()
                    if ap == "pm" and h != 12: h += 12
                    if ap == "am" and h == 12: h = 0
                    return h, m

                def _parse_time_component_relaxed(s, default_ampm=DEFAULT_AMPM):
                    s = (s or "").strip().lower()
                    m = re.match(r"^(\d{1,2})(?::(\d{2}))?\s*(am|pm)?$", s)
                    if not m: return None
                    hh = int(m.group(1)); mm = int(m.group(2) or 0); ap = m.group(3)
                    if ap:
                        return _to_24h(hh, mm, ap)
                    if 0 <= hh <= 23:
                        if hh <= 12 and default_ampm in ("am","pm"):
                            return _to_24h(hh, mm, default_ampm)
                        return (hh, mm)
                    return None

                def _parse_time_range_relaxed(rng, default_ampm=DEFAULT_AMPM):
                    rng = (rng or "").strip().lower().replace("–","-").replace("—","-")
                    parts = [p.strip() for p in rng.split("-", 1)]
                    if len(parts) != 2: return None
                    a = _parse_time_component_relaxed(parts[0], default_ampm=default_ampm)
                    if not a: return None
                    ap_hint = re.search(r"(am|pm)\s*$", parts[0])
                    second_default = ap_hint.group(1) if ap_hint else default_ampm
                    b = _parse_time_component_relaxed(parts[1], default_ampm=second_default)
                    return (a, b) if b else None

                def _expand_day_token(tok):
                    tok = (tok or "").strip().lower().replace("–","-").replace("—","-")
                    if "-" in tok:
                        a, b = [t.strip() for t in tok.split("-", 1)]
                        a_code = _FULL_TO_CODE.get(a, ""); b_code = _FULL_TO_CODE.get(b, "")
                        if a_code and b_code:
                            ai = _WKD_ORDER.index(a_code); bi = _WKD_ORDER.index(b_code)
                            return _WKD_ORDER[ai:bi+1] if ai <= bi else _WKD_ORDER[ai:] + _WKD_ORDER[:bi+1]
                        return []
                    c = _FULL_TO_CODE.get(tok, "")
                    return [c] if c else []

                def _parse_time_blocks(time_str, days_list):
                    s = _normalize_time_groups(time_str)
                    blocks = []
                    if ":" in s:
                        groups = [g.strip() for g in s.split(",") if g.strip()]
                        for g in groups:
                            if ":" not in g:
                                continue
                            left, right = [x.strip() for x in g.split(":", 1)]
                            day_tokens = re.split(r"/", left)
                            codes = []
                            for tok in day_tokens:
                                codes.extend(_expand_day_token(tok))
                            tr = _parse_time_range_relaxed(right)
                            if codes and tr:
                                (sh, sm), (eh, em) = tr
                                blocks.append({
                                    "byday": sorted(set(codes), key=_WKD_ORDER.index),
                                    "start": (sh, sm), "end": (eh, em)
                                })
                        return blocks
                    tr = _parse_time_range_relaxed(s)
                    if not tr:
                        return []
                    (sh, sm), (eh, em) = tr
                    codes = []
                    for d in (days_list or []):
                        c = _FULL_TO_CODE.get(str(d).lower().strip(), "")
                        if c: codes.append(c)
                    codes = sorted(set(codes), key=_WKD_ORDER.index) or _WKD_ORDER[:]
                    return [{"byday": codes, "start": (sh, sm), "end": (eh, em)}]

                def _next_on_or_after(d, weekday_index):
                    delta = (weekday_index - d.weekday()) % 7
                    return d + _td(days=delta)

                _blocks = _parse_time_blocks(time_str, days)
                if not _blocks and (days and str(time_str or "").strip()):
                    tr_fallback = _parse_time_range_relaxed(str(time_str))
                    if tr_fallback:
                        (sh, sm), (eh, em) = tr_fallback
                        codes = []
                        for d in (days or []):
                            c = _FULL_TO_CODE.get(str(d).lower().strip(), "")
                            if c: codes.append(c)
                        if codes:
                            codes = sorted(set(codes), key=_WKD_ORDER.index)
                            _blocks = [{"byday": codes, "start": (sh, sm), "end": (eh, em)}]

                # === Next class countdown ======================
                def _compute_next_class_instance(now_utc: _dt):
                    if not _blocks:
                        return None, None, ""
                    _wmap = {"MO":0,"TU":1,"WE":2,"TH":3,"FR":4,"SA":5,"SU":6}
                    best = None
                    cur = max(start_date_obj, now_utc.date())
                    while cur <= end_date_obj:
                        widx = cur.weekday()
                        for blk in _blocks:
                            if any(_wmap[c] == widx for c in blk["byday"]):
                                sh, sm = blk["start"]; eh, em = blk["end"]
                                sdt = _dt(cur.year, cur.month, cur.day, sh, sm, tzinfo=_timezone.utc)   # Ghana == UTC
                                edt = _dt(cur.year, cur.month, cur.day, eh, em, tzinfo=_timezone.utc)
                                if edt <= now_utc:
                                    continue
                                def _fmt_ampm(h, m):
                                    ap = "AM" if h < 12 else "PM"
                                    hh = h if 1 <= h <= 12 else (12 if h % 12 == 0 else h % 12)
                                    return f"{hh}:{m:02d}{ap}"
                                label = f"{cur.strftime('%a %d %b')} • {_fmt_ampm(sh, sm)}–{_fmt_ampm(eh, em)}"
                                cand = (sdt, edt, label)
                                if (best is None) or (sdt < best[0]):
                                    best = cand
                        cur += _td(days=1)
                    return best if best else (None, None, "")

                def _human_delta_ms(ms: int) -> str:
                    s = max(0, ms // 1000)
                    d, r = divmod(s, 86400)
                    h, r = divmod(r, 3600)
                    m, _ = divmod(r, 60)
                    parts = []
                    if d: parts.append(f"{d}d")
                    if h: parts.append(f"{h}h")
                    if (d == 0) and (m or not parts):
                        parts.append(f"{m}m")
                    return " ".join(parts) if parts else "0m"

                _now = _dt.now(_timezone.utc)
                nxt_start, nxt_end, nxt_label = _compute_next_class_instance(_now)
                if nxt_start and nxt_end:
                    start_ms = int(nxt_start.timestamp() * 1000)
                    now_ms   = int(_now.timestamp() * 1000)
                    time_left_label = _human_delta_ms(start_ms - now_ms) if now_ms < start_ms else "now"
                    st.info(f"**Next class:** {nxt_label}  •  **Starts in:** {time_left_label}", icon="⏰")
                    if components:
                        components.html(
                            f"""
                            <div id="nextCount" style="margin:6px 0 2px;color:#0f172a;font-weight:600;"></div>
                            <script>
                              (function(){{
                                const startMs = {start_ms};
                                const el = document.getElementById('nextCount');
                                function tick(){{
                                  const now = Date.now();
                                  if (now >= startMs) {{
                                    el.textContent = "Class is LIVE or started.";
                                  }} else {{
                                    const diff = startMs - now;
                                    const s = Math.floor(diff/1000);
                                    const d = Math.floor(s/86400);
                                    const h = Math.floor((s%86400)/3600);
                                    const m = Math.floor((s%3600)/60);
                                    const sec = s % 60;
                                    let txt = "Starts in: ";
                                    if (d) txt += d + "d ";
                                    if (h) txt += h + "h ";
                                    if (d || h) {{
                                      txt += m + "m";
                                    }} else {{
                                      txt += m + "m " + sec + "s";
                                    }}
                                    el.textContent = txt;
                                  }}
                                  setTimeout(tick, 1000);
                                }}
                                tick();
                              }})();
                            </script>
                            """,
                            height=28,
                        )

                # ================= ICS BUILD (full course) =================
                ZOOM = {
                    "link": (st.secrets.get("zoom", {}).get("link", "") if hasattr(st, "secrets") else "") or "https://zoom.us",
                    "meeting_id": (st.secrets.get("zoom", {}).get("meeting_id", "") if hasattr(st, "secrets") else "") or "",
                    "passcode": (st.secrets.get("zoom", {}).get("passcode", "") if hasattr(st, "secrets") else "") or "",
                }
                _zl = (ZOOM or {}).get("link", "")
                _zid = (ZOOM or {}).get("meeting_id", "")
                _zpw = (ZOOM or {}).get("passcode", "")
                _details = f"Zoom link: {_zl}\\nMeeting ID: {_zid}\\nPasscode: {_zpw}"
                _dtstamp = _dt.now(_timezone.utc).strftime("%Y%m%dT%H%M%SZ")
                _until = _dt(end_date_obj.year, end_date_obj.month, end_date_obj.day, 23, 59, 59, tzinfo=_timezone.utc).strftime("%Y%m%dT%H%M%SZ")
                _summary = f"{class_name} — Live German Class"

                USE_TZID = False
                TZID = "Africa/Accra"

                _ics_lines = [
                    "BEGIN:VCALENDAR","VERSION:2.0","PRODID:-//Falowen//Course Scheduler//EN",
                    "CALSCALE:GREGORIAN","METHOD:PUBLISH",
                ]

                if not _blocks:
                    _start_dt = _dt(start_date_obj.year, start_date_obj.month, start_date_obj.day, 18, 0)
                    _end_dt   = _dt(start_date_obj.year, start_date_obj.month, start_date_obj.day, 19, 0)
                    if USE_TZID:
                        dtfmt = "%Y%m%dT%H%M%S"
                        dtstart_line = f"DTSTART;TZID={TZID}:{_start_dt.strftime(dtfmt)}"
                        dtend_line   = f"DTEND;TZID={TZID}:{_end_dt.strftime(dtfmt)}"
                    else:
                        dtstart_line = f"DTSTART:{_start_dt.strftime('%Y%m%dT%H%M%SZ')}"
                        dtend_line   = f"DTEND:{_end_dt.strftime('%Y%m%dT%H%M%SZ')}"
                    _ics_lines += [
                        "BEGIN:VEVENT",
                        f"UID:{uuid4()}@falowen",
                        f"DTSTAMP:{_dtstamp}",
                        dtstart_line,
                        dtend_line,
                        f"SUMMARY:{_summary}",
                        f"DESCRIPTION:{_details}",
                        f"URL:{_zl}",
                        "LOCATION:Zoom",
                        "BEGIN:VALARM","ACTION:DISPLAY","DESCRIPTION:Class starts soon","TRIGGER:-PT15M","END:VALARM",
                        "END:VEVENT",
                    ]
                else:
                    for blk in _blocks:
                        byday_codes = blk["byday"]
                        sh, sm = blk["start"]; eh, em = blk["end"]
                        _wmap = {"MO":0,"TU":1,"WE":2,"TH":3,"FR":4,"SA":5,"SU":6}
                        first_dates = []
                        for code in byday_codes:
                            widx = _wmap[code]
                            first_dates.append(_next_on_or_after(start_date_obj, widx))
                        first_date = min(first_dates)
                        dt_start = _dt(first_date.year, first_date.month, first_date.day, sh, sm)
                        dt_end   = _dt(first_date.year, first_date.month, first_date.day, eh, em)

                        if USE_TZID:
                            dtfmt = "%Y%m%dT%H%M%S"
                            dtstart_line = f"DTSTART;TZID={TZID}:{dt_start.strftime(dtfmt)}"
                            dtend_line   = f"DTEND;TZID={TZID}:{dt_end.strftime(dtfmt)}"
                        else:
                            dtstart_line = f"DTSTART:{dt_start.strftime('%Y%m%dT%H%M%SZ')}"
                            dtend_line   = f"DTEND:{dt_end.strftime('%Y%m%dT%H%M%SZ')}"

                        _ics_lines += [
                            "BEGIN:VEVENT",
                            f"UID:{uuid4()}@falowen",
                            f"DTSTAMP:{_dtstamp}",
                            dtstart_line,
                            dtend_line,
                            f"RRULE:FREQ=WEEKLY;BYDAY={','.join(byday_codes)};UNTIL={_until}",
                            f"SUMMARY:{_summary}",
                            f"DESCRIPTION:{_details}",
                            f"URL:{_zl}",
                            "LOCATION:Zoom",
                            "BEGIN:VALARM","ACTION:DISPLAY","DESCRIPTION:Class starts soon","TRIGGER:-PT15M","END:VALARM",
                            "END:VEVENT",
                        ]

                _ics_lines.append("END:VCALENDAR")
                _course_ics = "\n".join(_ics_lines)

                c1, c2 = st.columns([1, 1])
                with c1:
                    st.download_button(
                        "⬇️ Download full course (.ics)",
                        data=_course_ics,
                        file_name=f"{class_name.replace(' ', '_')}_course.ics",
                        mime="text/calendar",
                        key=_ukey("dl_course_ics"),
                    )
                with c2:
                    st.caption("Calendar created. Use the download button to import the full course.")

                # --- Quick Android repeat links ---
                _gcal_repeat_links = []
                try:
                    if _blocks:
                        _wmap = {"MO":0,"TU":1,"WE":2,"TH":3,"FR":4,"SA":5,"SU":6}
                        _code_to_pretty = {"MO":"Mon","TU":"Tue","WE":"Wed","TH":"Thu","FR":"Fri","SA":"Sat","SU":"Sun"}

                        def _fmt_time(h, m):
                            ap = "AM" if h < 12 else "PM"
                            hh = h if 1 <= h <= 12 else (12 if h % 12 == 0 else h % 12)
                            return f"{hh}:{m:02d}{ap}"

                        for blk in _blocks:
                            byday_codes = blk["byday"]
                            sh, sm = blk["start"]; eh, em = blk["end"]

                            first_dates = []
                            for code in byday_codes:
                                widx = _wmap[code]
                                first_dates.append(_next_on_or_after(start_date_obj, widx))
                            first_date = min(first_dates)

                            _start_dt = _dt(first_date.year, first_date.month, first_date.day, sh, sm)
                            _end_dt   = _dt(first_date.year, first_date.month, first_date.day, eh, em)
                            _start_str = _start_dt.strftime("%Y%m%dT%H%M%SZ")
                            _end_str   = _end_dt.strftime("%Y%m%dT%H%M%SZ")

                            _until = _dt(end_date_obj.year, end_date_obj.month, end_date_obj.day, 23, 59, 59).strftime("%Y%m%dT%H%M%SZ")
                            _rrule = f"RRULE:FREQ=WEEKLY;BYDAY={','.join(byday_codes)};UNTIL={_until}"

                            _days_pretty = "/".join(_code_to_pretty[c] for c in byday_codes)
                            _label = f"{_days_pretty} {_fmt_time(sh, sm)}–{_fmt_time(eh, em)}"

                            _recur_url = (
                                "https://calendar.google.com/calendar/render"
                                f"?action=TEMPLATE"
                                f"&text={_urllib.quote(_summary)}"
                                f"&dates={_start_str}/{_end_str}"
                                f"&details={_urllib.quote(_details)}"
                                f"&location={_urllib.quote('Zoom')}"
                                f"&ctz={_urllib.quote('Africa/Accra')}"
                                f"&recur={_urllib.quote(_rrule)}"
                                f"&sf=true"
                            )
                            _gcal_repeat_links.append((_label, _recur_url))
                except Exception:
                    _gcal_repeat_links = []

                if _gcal_repeat_links:
                    _items = "".join(
                        f"<li style='margin:4px 0;'><a href='{url.replace('&','&amp;')}' target='_blank'>Tap here: {lbl}</a></li>"
                        for (lbl, url) in _gcal_repeat_links
                    )
                    _phone_links_ul = f"<ul style='margin:6px 0 0 18px;padding:0;'>{_items}</ul>"
                else:
                    _phone_links_ul = (
                        "<div style='margin:6px 0 0 2px;color:#444;'>"
                        "No repeating blocks are set yet. Ask the office to add your class times."
                        "</div>"
                    )

                st.markdown(
                    f"""
                    **Computer or iPhone:** Download the **.ics** above and install.
                    - **Computer (Google Calendar web):** calendar.google.com → **Settings** → **Import & export** → **Import**.
                    - **iPhone (Apple Calendar):** Download the `.ics`, open it, choose notifications, then **Done**.

                    **Android (Google Calendar app):** The app **can’t import `.ics`**. Use these links (**with repeat**):
                    {_phone_links_ul}
                    <div style="margin:8px 0 0 2px;"></div>
                    """,
                    unsafe_allow_html=True,
                )

        # ===================== ATTENDANCE =====================
        elif classroom_section == "Attendance":
            with st.container():
                st.markdown(
                    """
                    <div style="
                        padding:10px 12px;
                        background:#f0f9ff;
                        border:1px solid #bae6fd;
                        border-radius:12px;
                        margin: 6px 0 8px 0;
                        display:flex;align-items:center;gap:8px;">
                      <span style="font-size:1.05rem;">📊 <b>Attendance</b></span>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )
                records, sessions_attended, hours_invested = load_attendance_records(
                    student_code, class_name
                )
                cols = st.columns(2)
                cols[0].metric("Attended sessions", sessions_attended)
                cols[1].metric("Invested hours", f"{hours_invested:.1f}")
                if records:
                    df_att = pd.DataFrame(records)
                    df_att["Present"] = df_att.pop("present").map({True: "✓", False: ""})
                    df_att.rename(columns={"session": "Session"}, inplace=True)
                    st.dataframe(
                        df_att,
                        width="stretch",
                        hide_index=True,
                    )
                else:
                    st.info("No attendance records found yet.")

        # ===================== MEMBERS & PROFILE =====================
        elif classroom_section == "Members & Profile":
            # Subtle hint banner
            st.markdown(
                """
                <div style="
                    padding:10px 12px;
                    background:#f0f9ff;
                    border:1px solid #bae6fd;
                    border-radius:12px;
                    margin: 6px 0 8px 0;
                    display:flex;align-items:center;gap:8px;">
                  <span style="font-size:1.05rem;">👥 <b>Class Members</b></span>
                </div>
                """,
                unsafe_allow_html=True,
            )
            with st.container():
                try:
                    df_students = load_student_data()
                except Exception:
                    df_students = pd.DataFrame()
                if df_students is None:
                    df_students = pd.DataFrame()

                for col in (
                    "ClassName",
                    "Name",
                    "Email",
                    "Location",
                    "StudentCode",
                ):
                    if col not in df_students.columns:
                        df_students[col] = ""
                    df_students[col] = (
                        df_students[col].fillna("").astype(str).str.strip()
                    )

                same_class = df_students[df_students["ClassName"] == class_name].copy()
                if not same_class.empty:
                    def _about_for(code: str) -> str:
                        """Fetch the student's bio."""
                        return load_student_profile(code or "")

                    same_class["About"] = same_class["StudentCode"].apply(
                        _about_for
                    )
                _n = len(same_class)
                st.markdown(
                    f"""
                    <div style="display:flex;justify-content:space-between;align-items:center;margin:4px 0 6px 0;">
                      <div style="font-weight:600;color:#0f172a;">{class_name}</div>
                      <span style="background:#0ea5e922;border:1px solid #0ea5e9;color:#0369a1;
                                   padding:3px 8px;border-radius:999px;font-size:.9rem;">
                        {_n} member{'' if _n==1 else 's'}
                      </span>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )

                cols_show = [
                    c for c in ["Name", "Email", "Location", "About"] if c in same_class.columns
                ]
                if not same_class.empty and cols_show:
                    for _, row in same_class[cols_show].reset_index(drop=True).iterrows():
                        name = row.get("Name", "")
                        email = row.get("Email", "")
                        location = row.get("Location", "")
                        about = row.get("About", "")

                        contact = " | ".join(
                            [part for part in [email, location] if part]
                        )

                        st.markdown(
                            f"""
                            <div style="width:100%;padding:8px 0;border-bottom:1px solid #e5e7eb;">
                                <div style="font-weight:600;color:#0f172a;">{name}</div>
                                <div style="font-size:.9rem;color:#475569;">{contact}</div>
                                <div style="margin-top:4px;">{about}</div>
                            </div>
                            """,
                            unsafe_allow_html=True,
                        )
                else:
                    st.info("No members found for this class yet.")
            st.info("Scroll down to update your profile description.")

            # --- Profile section ---
            st.markdown("---")
            with st.container():
                st.markdown(
                    """
                    <div style="
                        padding:10px 12px;
                        background:#f0f9ff;
                        border:1px solid #bae6fd;
                        border-radius:12px;
                        margin: 6px 0 8px 0;
                        display:flex;align-items:center;gap:8px;">
                      <span style="font-size:1.05rem;">👤 <b>Profile</b></span>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )
                student_code = (st.session_state.get("student_code", "") or "").strip()
                loaded_key = "profile_loaded_code"
                about_key = "profile_about"
                edit_key = "profile_editing"
                cancel_profile_key = "profile_cancel"
                if (
                    student_code
                    and (
                        st.session_state.get(loaded_key) != student_code
                        or about_key not in st.session_state
                    )
                ):
                    st.session_state[about_key] = load_student_profile(student_code)
                    st.session_state[loaded_key] = student_code
                if not student_code:
                    st.session_state.setdefault(about_key, "")
                st.session_state.setdefault(edit_key, False)
                st.session_state.setdefault(cancel_profile_key, False)

                ai_flag = f"profile_ai_busy_{student_code}"
                if st.session_state.get(cancel_profile_key):
                    st.session_state[about_key] = (
                        load_student_profile(student_code) if student_code else ""
                    )
                    st.session_state.pop(ai_flag, None)
                    st.session_state[edit_key] = False
                    st.session_state[cancel_profile_key] = False

                editing = st.session_state.get(edit_key, False)
                if editing:
                    if st.session_state.get(ai_flag):
                        with st.spinner("Correcting with AI..."):
                            apply_profile_ai_correction(about_key)
                        st.session_state[ai_flag] = False
                        st.session_state.pop("need_rerun", None)
                    st.text_area("About me", key=about_key, height=300)
                else:
                    st.markdown(st.session_state[about_key])

                if not editing:
                    if st.button("Edit", disabled=not bool(student_code), key=_ukey("edit_profile")):
                        st.session_state[edit_key] = True
                else:
                    col1, col_ai, col2 = st.columns(3)
                    with col1:
                        if st.button("Save", key=_ukey("save_profile")):
                            try:
                                save_student_profile(student_code, st.session_state.get(about_key, ""))
                                st.success("Profile saved.")
                            except Exception:
                                st.error("Failed to save profile.")
                            finally:
                                st.session_state.pop(ai_flag, None)
                                st.session_state[edit_key] = False
                    with col_ai:
                        if st.button(
                            "✨ Correct with AI",
                            key=_ukey("ai_profile"),
                            disabled=st.session_state.get(ai_flag, False),
                        ):
                            st.session_state[ai_flag] = True
                            st.session_state["need_rerun"] = True
                    with col2:
                        if st.button("Cancel", key=_ukey("cancel_profile")):
                            st.session_state[cancel_profile_key] = True
                            st.rerun()

                if not bool(student_code):
                    st.info("Enter your student code to edit your profile.")

        # ===================== JOIN =====================
        elif classroom_section == "Join on Zoom":
            with st.container():
                st.markdown(
                    """
                    <div style="padding: 12px; background: #facc15; color: #000; border-radius: 8px;
                         font-size: 1rem; margin-bottom: 16px; text-align: left; font-weight: 600;">
                      📣 <b>Zoom Classroom (Official)</b><br>
                      This is the <u>official Zoom link</u> for your class. <span style="font-weight:500;">Add the calendar below to get notifications before each class.</span>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )

                ZOOM = {
                    "link": "https://us06web.zoom.us/j/6886900916?pwd=bEdtR3RLQ2dGTytvYzNrMUV3eFJwUT09",
                    "meeting_id": "688 690 0916",
                    "passcode": "german",
                }
                # Allow secrets override
                try:
                    zs = st.secrets.get("zoom", {})
                    if zs.get("link"):       ZOOM["link"]       = zs["link"]
                    if zs.get("meeting_id"): ZOOM["meeting_id"] = zs["meeting_id"]
                    if zs.get("passcode"):   ZOOM["passcode"]   = zs["passcode"]
                except Exception:
                    pass

                # Build iOS/Android deep-link (opens Zoom app directly)
                _mid_digits = ZOOM["meeting_id"].replace(" ", "")
                _pwd_enc = _urllib.quote(ZOOM["passcode"] or "")
                zoom_deeplink = f"zoommtg://zoom.us/join?action=join&confno={_mid_digits}&pwd={_pwd_enc}"

                z1, z2 = st.columns([3, 2])
                with z1:
                    # Primary join button (browser)
                    try:
                        st.link_button("➡️ Join Zoom Meeting (Browser)", ZOOM["link"], key=_ukey("zoom_join_btn"))
                    except Exception:
                        st.markdown(f"[➡️ Join Zoom Meeting (Browser)]({ZOOM['link']})")

                    # Secondary: open in Zoom app (mobile deep link)
                    try:
                        st.link_button("📱 Open in Zoom App", zoom_deeplink, key=_ukey("zoom_app_btn"))
                    except Exception:
                        st.markdown(f"[📱 Open in Zoom App]({zoom_deeplink})")

                    st.write(f"**Meeting ID:** `{ZOOM['meeting_id']}`")
                    st.write(f"**Passcode:** `{ZOOM['passcode']}`")

                    # Copy helpers (mobile-friendly, safe escaping)
                    _link_safe = ZOOM["link"].replace("'", "\\'")
                    _id_safe   = ZOOM["meeting_id"].replace("'", "\\'")
                    _pwd_safe  = ZOOM["passcode"].replace("'", "\\'")
                    if components:
                        components.html(
                            f"""
                            <div style="display:flex;gap:8px;margin-top:8px;">
                              <button id="zCopyLink"
                                      style="padding:6px 10px;border-radius:8px;border:1px solid #cbd5e1;background:#f1f5f9;cursor:pointer;">
                                Copy Link
                              </button>
                              <button id="zCopyId"
                                      style="padding:6px 10px;border-radius:8px;border:1px solid #cbd5e1;background:#f1f5f9;cursor:pointer;">
                                Copy ID
                              </button>
                              <button id="zCopyPwd"
                                      style="padding:6px 10px;border-radius:8px;border:1px solid #cbd5e1;background:#f1f5f9;cursor:pointer;">
                                Copy Passcode
                              </button>
                            </div>
                            <script>
                              (function(){{
                                try {{
                                  var link = '{_link_safe}', mid = '{_id_safe}', pwd = '{_pwd_safe}';
                                  function wire(btnId, txt, label) {{
                                    var b = document.getElementById(btnId);
                                    if (!b) return;
                                    b.addEventListener('click', function(){{
                                      navigator.clipboard.writeText(txt).then(function(){{
                                        b.innerText = '✓ Copied ' + label;
                                        setTimeout(function(){{ b.innerText = 'Copy ' + label; }}, 1500);
                                      }}).catch(function(){{}});
                                    }});
                                  }}
                                  wire('zCopyLink', link, 'Link');
                                  wire('zCopyId',   mid,  'ID');
                                  wire('zCopyPwd',  pwd,  'Passcode');
                                }} catch(e) {{}}
                              }})();
                            </script>
                            """,
                            height=72,
                        )

                with z2:
                    st.info(
                        f"You’re viewing: **{class_name}**  \n\n"
                        "✅ Use the **calendar** tab to receive automatic class reminders.",
                        icon="📅",
                    )


        # ===================== Class Board =====================
        elif classroom_section == "Class Notes & Q&A":
            st.markdown("<div id='classnotes'></div>", unsafe_allow_html=True)
            if st.session_state.pop("__scroll_to_classnotes", False):
                st.markdown(
                    "<script>window.location.hash='classnotes';</script>",
                    unsafe_allow_html=True,
                )
            board_base = (
                db.collection("class_board")
                .document(student_level)
                .collection("classes")
                .document(class_name)
                .collection("posts")
            )

            _new7, _unans, _total = 0, 0, 0
            try:
                _now = _dt.now(_timezone.utc)
                try:
                    from firebase_admin import firestore as fbfs
                    direction_desc = getattr(fbfs.Query, "DESCENDING", "DESCENDING")
                    _qdocs = list(board_base.order_by("created_at", direction=direction_desc).limit(250).stream())
                except Exception:
                    _qdocs = list(board_base.order_by("created_at", direction="DESCENDING").limit(250).stream())

                for _doc in _qdocs:
                    _d = (_doc.to_dict() or {})
                    _total += 1
                    _rc = 0
                    if isinstance(_d.get("answers"), list):
                        _rc = len(_d["answers"])
                    elif isinstance(_d.get("replies"), list):
                        _rc = len(_d["replies"])
                    elif isinstance(_d.get("reply_count"), int):
                        _rc = int(_d["reply_count"])
                    if _rc == 0:
                        _unans += 1
                    _created = _to_datetime_any(_d.get("created_at") or _d.get("ts") or _d.get("timestamp"))
                    if _created and (_now - _created).days <= 7:
                        _new7 += 1
            except Exception:
                pass

            _badges = []
            if _new7 > 0:
                _badges.append(
                    f"<span style='margin-left:8px;background:#16a34a;color:#fff;padding:2px 8px;"
                    f"border-radius:999px;font-size:0.8rem;'>NEW · {_new7}</span>"
                )
            if _unans > 0:
                _badges.append(
                    f"<span style='margin-left:8px;background:#f97316;color:#fff;padding:2px 8px;"
                    f"border-radius:999px;font-size:0.8rem;'>UNANSWERED · {_unans}</span>"
                )
            _badge_html = "".join(_badges)

            st.markdown(
                f'''
                <div style="
                    padding:12px;
                    background:#2e7d32;
                    background-image:radial-gradient(circle, rgba(255,255,255,0.05) 1px, transparent 1px);
                    background-size:6px 6px;
                    color:#f5f5f5;
                    border-radius:8px;
                    margin-bottom:12px;
                    border:2px solid #c8c8c8;
                    box-shadow:inset 0 0 4px rgba(0,0,0,0.6), 0 2px 6px rgba(0,0,0,0.08);
                    font-family:'Chalkboard','Chalkduster','Comic Sans MS',cursive;
                    display:flex;align-items:center;justify-content:space-between;">
                    <div style="font-weight:700;font-size:1.15rem;">💬 Class Board — {class_name} {_badge_html}</div>
                    <div style="font-size:0.92rem;opacity:.9;">Share a post • Comment with {class_name} classmates</div>
                </div>
                ''',
                unsafe_allow_html=True
            )

            def _fmt_ts(ts):
                try:
                    return ts.strftime("%d %b %H:%M:%S")
                except Exception:
                    return ""

            def _normalize_legacy_post_content(text: str) -> str:
                """Clean artifacts from legacy wrapped posts for display."""
                if not text:
                    return ""

                normalized = text.replace("\r\n", "\n").replace("\r", "\n")

                # Remove leading double-spaces introduced by legacy rewrapping.
                normalized = normalized.lstrip()
                normalized = re.sub(r"(?m)(?<=\n)[ \t]{2,}(?=\S)", "", normalized)

                if "·" in normalized:
                    # Convert legacy middot bullets to regular bullets for consistency.
                    normalized = re.sub(
                        r"(?m)(^|\n)\s*·\s*",
                        lambda m: (m.group(1) or "") + "• ",
                        normalized,
                    )
                    normalized = normalized.replace(" · ", " • ")

                return normalized

            def format_post(text: str) -> str:
                """Normalize post text for consistent rendering."""
                cleaned = (text or "")
                cleaned = cleaned.replace("\r\n", "\n").replace("\r", "\n")
                cleaned = cleaned.strip()
                cleaned = re.sub(r"\n\s*\n+", "\n\n", cleaned)
                return cleaned

            def escape_with_linebreaks(text: str) -> str:
                """Escape text and convert newlines into HTML-friendly breaks."""
                if not text:
                    return ""

                normalized = _normalize_legacy_post_content(text)
                normalized = normalized.strip()
                normalized = re.sub(r"\n\s*\n+", "\n\n", normalized)

                paragraphs = [p for p in normalized.split("\n\n") if p != ""]
                if not paragraphs:
                    return ""

                escaped_paragraphs: List[str] = []
                for para in paragraphs:
                    escaped = html.escape(para)
                    escaped_paragraphs.append(escaped.replace("\n", "<br>"))

                return "".join(f"<p>{p}</p>" for p in escaped_paragraphs)

            # ---- Prepare lesson choices ----
            lesson_choices = []
            try:
                schedules = load_level_schedules()
                level_sched = schedules.get(student_level, schedules.get("A1", []))
                for item in level_sched:
                    day = item.get("day")
                    topic = item.get("topic")
                    if day is not None and topic:
                        lesson_choices.append(_schedule.full_lesson_title(item))
            except Exception:
                pass

            st.subheader("➕ Add a new post")
            st.info(
                "Decide whether you're responding to a classmate or posting your own question. "
                "If you're responding, scroll down to find the existing question and reply there. "
                "When you want to post a new question, use the text box below."
            )
            draft_key = "q_text"
            initialize_draft_state(student_code, draft_key)
            if st.session_state.get("__clear_q_form"):
                st.session_state.pop("__clear_q_form", None)
                st.session_state["q_topic"] = ""
                st.session_state["q_link"] = ""
                st.session_state["q_lesson"] = lesson_choices[0] if lesson_choices else ""
                st.session_state.pop("q_ai_suggestion", None)
                st.session_state.pop("q_ai_explanation", None)
                st.session_state.pop("q_ai_diff", None)
                st.session_state["q_forum_timer_minutes"] = 0
                reset_local_draft_state(draft_key)
                _clear_typing_state(
                    level=student_level,
                    class_code=class_name,
                    qid=_NEW_POST_TYPING_ID,
                    draft_key="q_text",
                    student_code=student_code,
                    student_name=student_name,
                )
            lesson = (
                st.selectbox("Lesson", lesson_choices, key="q_lesson")
                if lesson_choices
                else st.text_input("Lesson", key="q_lesson")
            )
            topic = st.text_input("Topic (optional)", key="q_topic")
            link = st.text_input("Link (optional)", key="q_link")
            timer_key = "q_forum_timer_minutes"
            if timer_key not in st.session_state:
                st.session_state[timer_key] = 0
            if IS_ADMIN:
                st.number_input(
                    "Forum timer (minutes)",
                    min_value=0,
                    step=5,
                    key=timer_key,
                    help="Automatically close replies after this many minutes.",
                )

            st.markdown(
                """
                <style>
                textarea[aria-label="Your content"] {
                    background-color: #f1f5f9;
                    color: #0f172a;
                    font-family: 'Chalkboard', 'Chalkduster', 'Comic Sans MS', cursive;
                    font-size: 1.1rem;
                    line-height: 1.4;
                }
                </style>
                """,
                unsafe_allow_html=True,
            )
            ai_flag = "__q_ai_busy"
            if st.session_state.get(ai_flag):
                with st.spinner("Correcting with AI..."):
                    original = st.session_state.get("q_text", "")
                    improved, explanation = apply_status_ai_correction(original)
                    st.session_state["q_ai_suggestion"] = improved
                    st.session_state["q_ai_explanation"] = explanation
                    st.session_state["q_ai_diff"] = diff_with_markers(original, improved)
                st.session_state[ai_flag] = False

            ta_col, ai_col = st.columns([3, 1])
            with ta_col:
                banner = _format_typing_banner(
                    fetch_active_typists(
                        student_level,
                        class_name,
                        _NEW_POST_TYPING_ID,
                    ),
                    student_code,
                )
                if banner:
                    st.caption(banner)
                new_q = st.text_area("Your content", key=draft_key, height=160)
                render_umlaut_pad("q_text", context="classboard_post")
                _update_typing_state(
                    level=student_level,
                    class_code=class_name,
                    qid=_NEW_POST_TYPING_ID,
                    draft_key="q_text",
                    student_code=student_code,
                    student_name=student_name,
                    text=new_q,
                )
                autosave_maybe(
                    student_code,
                    draft_key,
                    st.session_state.get(draft_key, ""),
                    min_secs=2.0,
                    min_delta=12,
                )
            with ai_col:
                if st.button(
                    "✨ Correct with AI",
                    key="qna_ai_correct",
                    disabled=st.session_state.get(ai_flag, False),
                ):
                    st.session_state[ai_flag] = True
                    st.rerun()

            if st.session_state.get("q_ai_diff"):
                st.markdown(st.session_state["q_ai_diff"], unsafe_allow_html=True)
                st.markdown("**Why these changes?**")
                st.markdown(st.session_state.get("q_ai_explanation", ""))
                acc_col, rej_col = st.columns(2)
                with acc_col:
                    if st.button("Accept", key="q_ai_accept"):
                        st.session_state["q_text"] = st.session_state.get("q_ai_suggestion", "")
                        st.session_state.pop("q_ai_suggestion", None)
                        st.session_state.pop("q_ai_explanation", None)
                        st.session_state.pop("q_ai_diff", None)
                        st.session_state["need_rerun"] = True
                with rej_col:
                    if st.button("Reject", key="q_ai_reject"):
                        st.session_state.pop("q_ai_suggestion", None)
                        st.session_state.pop("q_ai_explanation", None)
                        st.session_state.pop("q_ai_diff", None)
                        st.session_state["need_rerun"] = True

            if st.button("Post", key="qna_post_question"):
                formatted_q = format_post(new_q)
                if formatted_q:
                    q_id = str(uuid4())[:8]
                    payload = {
                        "content": formatted_q,
                        "asked_by_name": student_name,
                        "asked_by_code": student_code,
                        "timestamp": _dt.now(UTC),
                        "lesson": lesson,
                        "topic": (topic or "").strip(),
                        "link": (link or "").strip(),
                        "pinned": False,
                    }
                    timer_minutes_val = int(st.session_state.get("q_forum_timer_minutes", 0) or 0)
                    if timer_minutes_val > 0:
                        payload["expires_at"] = _dt.now(UTC) + timedelta(minutes=timer_minutes_val)
                    board_base.document(q_id).set(payload)
                    preview = (formatted_q[:180] + "…") if len(formatted_q) > 180 else formatted_q
                    topic_tag = f" • Topic: {payload['topic']}" if payload["topic"] else ""
                    _notify_slack(
                        f"📝 *New Class Board post* — {class_name}{topic_tag}\n",
                        f"*From:* {student_name} ({student_code})\n",
                        f"*When:* {_dt.now(UTC).strftime('%Y-%m-%d %H:%M')} UTC\n",
                        f"*Content:* {preview}"
                    )
                    _clear_typing_state(
                        level=student_level,
                        class_code=class_name,
                        qid=_NEW_POST_TYPING_ID,
                        draft_key="q_text",
                        student_code=student_code,
                        student_name=student_name,
                    )
                    clear_draft_after_post(student_code, draft_key)
                    st.session_state["__clear_q_form"] = True
                    st.success("Post published!")
                    refresh_with_toast()


            colsa, colsb, colsc = st.columns([2, 1, 1])
            with colsa:
                q_search = st.text_input("Search posts (text or topic)…", key="q_search")
            with colsb:
                show_latest = st.toggle("Newest first", value=True, key="q_show_latest")
            with colsc:
                if st.button("↻ Refresh", key="qna_refresh"):
                    refresh_with_toast()

            try:
                try:
                    from firebase_admin import firestore as fbfs
                    direction_desc = getattr(fbfs.Query, "DESCENDING", "DESCENDING")
                    q_docs = list(board_base.order_by("timestamp", direction=direction_desc).stream())
                except Exception:
                    q_docs = list(board_base.order_by("timestamp", direction="DESCENDING").stream())
                questions = [dict(d.to_dict() or {}, id=d.id) for d in q_docs]
            except Exception:
                q_docs = list(board_base.stream())
                questions = [dict(d.to_dict() or {}, id=d.id) for d in q_docs]
                questions.sort(key=lambda x: x.get("timestamp"), reverse=True)

            if q_search.strip():
                ql = q_search.lower()
                questions = [
                    q for q in questions
                    if ql in str(q.get("content", "")).lower() or ql in str(q.get("topic", "")).lower()
                ]
            if not show_latest:
                questions = list(reversed(questions))

            pinned_qs = [q for q in questions if q.get("pinned")]
            other_qs = [q for q in questions if not q.get("pinned")]
            questions = pinned_qs + other_qs

            def send_comment(
                q_id,
                student_code,
                student_name,
                class_name,
                board_base,
                draft_key,
                last_val_key,
                last_ts_key,
                saved_flag_key,
                saved_at_key,
            ):
                current_text = st.session_state.get(draft_key, "").strip()
                if not current_text:
                    return
                comment_payload = {
                    "content": current_text,
                    "replied_by_name": student_name,
                    "replied_by_code": student_code,
                    "timestamp": _dt.now(_timezone.utc),
                }
                c_ref = board_base.document(q_id).collection("comments")
                c_ref.document(str(uuid4())[:8]).set(comment_payload)
                prev = (
                    comment_payload["content"][:180] + "…"
                ) if len(comment_payload["content"]) > 180 else comment_payload["content"]
                _notify_slack(
                    f"💬 *New Class Board comment* — {class_name}\n",
                    f"*By:* {student_name} ({student_code})  •  *QID:* {q_id}\n",
                    f"*When:* {_dt.now(_timezone.utc).strftime('%Y-%m-%d %H:%M')} UTC\n",
                    f"*Comment:* {prev}",
                )
                save_draft_to_db(student_code, draft_key, "")
                _clear_typing_state(
                    level=student_level,
                    class_code=class_name,
                    qid=q_id,
                    draft_key=draft_key,
                    student_code=student_code,
                    student_name=student_name,
                )
                st.session_state[f"__clear_comment_draft_{q_id}"] = True
                st.session_state[last_val_key] = ""
                st.session_state[last_ts_key] = time.time()
                st.session_state[saved_flag_key] = False
                st.session_state[saved_at_key] = None
                st.success("Comment sent!")
                refresh_with_toast()

            comment_panel_active = (
                st.session_state.get("coursebook_subtab") == "🧑‍🏫 Classroom"
                and st.session_state.get("classroom_page") == "Class Notes & Q&A"
            )
            if comment_panel_active:
                try:
                    from streamlit_autorefresh import st_autorefresh
                except ImportError:
                    pass
                else:
                    st_autorefresh(interval=5000, key="classboard_comment_refresh")

            if not questions:
                st.info("No posts yet.")
            else:
                now_for_timer = _dt.now(UTC)
                for idx, q in enumerate(questions):
                    q_id = q.get("id", "")
                    ts = q.get("timestamp")
                    ts_label = _fmt_ts(ts)
                    pin_html = " 📌" if q.get("pinned") else ""
                    timer_info = build_forum_timer_indicator(q.get("expires_at"), now=now_for_timer)
                    timer_minutes_remaining = int(timer_info.get("minutes") or 0)
                    topic_html = (
                        f"<div style='font-weight:bold;color:#8d4de8;'>{html.escape(str(q.get('topic', '')))}</div>"
                        if q.get("topic")
                        else ""
                    )
                    content_html = escape_with_linebreaks(q.get("content", ""))
                    raw_link_value = q.get("link")
                    raw_link = str(raw_link_value).strip() if raw_link_value else ""
                    safe_link_html = ""
                    if raw_link:
                        parsed_link = urlparse(raw_link)
                        if parsed_link.scheme.lower() in {"http", "https"}:
                            safe_href = html.escape(raw_link, quote=True)
                            safe_label = html.escape(raw_link)
                            safe_link_html = (
                                "<div style='margin-top:4px;'>"
                                f"<a href='{safe_href}' target='_blank' rel='noopener noreferrer'>{safe_label}</a>"
                                "</div>"
                            )
                    link_html = safe_link_html
                    lesson_value = q.get("lesson")
                    lesson = str(lesson_value) if lesson_value else ""
                    if lesson:
                        day_part = lesson.split(":")[0]
                        day = day_part.split()[1] if len(day_part.split()) > 1 else ""
                        course_link = build_course_day_link(day)
                        safe_lesson = html.escape(lesson)
                        link_markup = ""
                        if course_link:
                            safe_course_link = html.escape(course_link, quote=True)
                            link_markup = (
                                f" – <a href='{safe_course_link}' target='_blank' rel='noopener noreferrer'>View page</a>"
                            )
                        lesson_html = (
                            "<div style='font-size:1.1rem;font-weight:600;color:#0f172a;'>"
                            f"📘 {safe_lesson}{link_markup}"
                            "</div>"
                        )
                    else:
                        lesson_html = ""
                    safe_author = html.escape(str(q.get("asked_by_name", "")))
                    safe_timestamp = html.escape(ts_label) if ts_label else ""
                    timestamp_html = (
                        f"<span style='color:#aaa;'> • {safe_timestamp}</span>" if safe_timestamp else ""
                    )
                    timer_html = ""
                    timer_label = timer_info.get("label") or ""
                    if timer_info.get("status") == "open" and timer_label:
                        timer_html = (
                            "<div style='margin-top:4px;font-size:0.95rem;font-weight:600;color:#dc2626;'>"
                            f"{html.escape(str(timer_label))}"
                            "</div>"
                        )
                    elif timer_info.get("status") == "closed" and timer_label:
                        timer_html = (
                            "<div style='margin-top:4px;font-size:0.95rem;font-weight:600;color:#64748b;'>"
                            f"{html.escape(str(timer_label))}"
                            "</div>"
                        )
                    post_html = (
                        "<div style='padding:10px;background:#f8fafc;border:1px solid #ddd;border-radius:6px;margin:6px 0;font-size:1rem;line-height:1.5;'>"
                        f"<b>{safe_author}</b>{pin_html}"
                        f"{timestamp_html}"
                        f"{timer_html}"
                        f"{lesson_html}"
                        f"{topic_html}"
                        f"{content_html}"
                        f"{link_html}"
                        "</div>"
                    )
                    st.markdown(post_html, unsafe_allow_html=True)

                    show_timer_warning = (
                        timer_info.get("status") == "open" and timer_minutes_remaining == 1
                    )

                    clear_q_edit_flag = f"__clear_q_edit_{q_id}"
                    if st.session_state.pop(clear_q_edit_flag, False):
                        for _k in [
                            f"q_edit_text_{q_id}",
                            f"q_edit_topic_{q_id}",
                            f"q_edit_link_{q_id}",
                            f"q_edit_lesson_{q_id}",
                            f"q_edit_text_input_{q_id}",
                            f"q_edit_topic_input_{q_id}",
                            f"q_edit_link_input_{q_id}",
                            f"q_edit_lesson_input_{q_id}",
                            f"q_edit_timer_input_{q_id}",
                        ]:
                            st.session_state.pop(_k, None)
                        _clear_typing_state(
                            level=student_level,
                            class_code=class_name,
                            qid=q_id,
                            draft_key=f"q_edit_text_{q_id}",
                            student_code=student_code,
                            student_name=student_name,
                        )

                    can_modify_q = (q.get("asked_by_code") == student_code) or IS_ADMIN
                    if can_modify_q:
                        qc1, qc2, qc3, _ = st.columns([1, 1, 1, 6])
                        with qc1:
                            if st.button("✏️ Edit", key=f"q_edit_btn_{q_id}"):
                                st.session_state[f"q_editing_{q_id}"] = True
                                st.session_state[f"q_edit_text_{q_id}"] = q.get("content", "")
                                st.session_state[f"q_edit_topic_{q_id}"] = q.get("topic", "")
                                st.session_state[f"q_edit_link_{q_id}"] = q.get("link", "")
                                st.session_state[f"q_edit_lesson_{q_id}"] = q.get("lesson", "")
                                st.session_state[f"q_edit_timer_input_{q_id}"] = timer_minutes_remaining
                        with qc2:
                            if st.button("🗑️ Delete", key=f"q_del_btn_{q_id}"):
                                try:
                                    c_ref = board_base.document(q_id).collection("comments")
                                    for rdoc in c_ref.stream():
                                        rdoc.reference.delete()
                                except Exception:
                                    pass
                                board_base.document(q_id).delete()
                                _notify_slack(
                                    f"🗑️ *Class Board post deleted* — {class_name}\n"
                                    f"*By:* {student_name} ({student_code}) • QID: {q_id}\n"
                                    f"*When:* {_dt.now(UTC).strftime('%Y-%m-%d %H:%M')} UTC"
                                )
                                st.success("Post deleted.")
                                refresh_with_toast()
                        with qc3:
                            pin_label = "📌 Unpin" if q.get("pinned") else "📌 Pin"
                            if st.button(pin_label, key=f"q_pin_btn_{q_id}"):
                                board_base.document(q_id).update({"pinned": not q.get("pinned", False)})
                                refresh_with_toast()

                        if st.session_state.get(f"q_editing_{q_id}", False):
                            with st.form(f"q_edit_form_{q_id}"):
                                if lesson_choices:
                                    current_lesson = st.session_state.get(f"q_edit_lesson_{q_id}", "")
                                    try:
                                        _idx = lesson_choices.index(current_lesson)
                                    except ValueError:
                                        _idx = 0
                                    new_lesson = st.selectbox(
                                        "Edit lesson",
                                        lesson_choices,
                                        index=_idx,
                                        key=f"q_edit_lesson_input_{q_id}"
                                    )
                                else:
                                    new_lesson = st.text_input(
                                        "Edit lesson",
                                        value=st.session_state.get(f"q_edit_lesson_{q_id}", ""),
                                        key=f"q_edit_lesson_input_{q_id}"
                                    )
                                new_topic = st.text_input(
                                    "Edit topic (optional)",
                                    value=st.session_state.get(f"q_edit_topic_{q_id}", ""),
                                    key=f"q_edit_topic_input_{q_id}"
                                )
                                new_link = st.text_input(
                                    "Edit link (optional)",
                                    value=st.session_state.get(f"q_edit_link_{q_id}", ""),
                                    key=f"q_edit_link_input_{q_id}"
                                )
                                if f"q_edit_timer_input_{q_id}" not in st.session_state:
                                    st.session_state[f"q_edit_timer_input_{q_id}"] = timer_minutes_remaining
                                if IS_ADMIN:
                                    st.number_input(
                                        "Forum timer (minutes)",
                                        min_value=0,
                                        step=5,
                                        key=f"q_edit_timer_input_{q_id}",
                                    )
                                banner = _format_typing_banner(
                                    fetch_active_typists(student_level, class_name, q_id),
                                    student_code,
                                )
                                if banner:
                                    st.caption(banner)
                                new_text = st.text_area(
                                    "Edit post",
                                    value=st.session_state.get(f"q_edit_text_{q_id}", ""),
                                    key=f"q_edit_text_input_{q_id}",
                                    height=150
                                )
                                _update_typing_state(
                                    level=student_level,
                                    class_code=class_name,
                                    qid=q_id,
                                    draft_key=f"q_edit_text_{q_id}",
                                    student_code=student_code,
                                    student_name=student_name,
                                    text=new_text,
                                )
                                save_edit = st.form_submit_button("💾 Save")
                                cancel_edit = st.form_submit_button("❌ Cancel")
                            if save_edit:
                                formatted_edit = format_post(new_text)
                                if formatted_edit:
                                    update_payload = {
                                        "content": formatted_edit,
                                        "topic": (new_topic or "").strip(),
                                        "link": (new_link or "").strip(),
                                        "lesson": new_lesson,
                                    }
                                    if IS_ADMIN:
                                        timer_minutes_updated = int(
                                            st.session_state.get(f"q_edit_timer_input_{q_id}", 0) or 0
                                        )
                                        if timer_minutes_updated > 0:
                                            update_payload["expires_at"] = _dt.now(UTC) + timedelta(
                                                minutes=timer_minutes_updated
                                            )
                                        else:
                                            update_payload["expires_at"] = firestore.DELETE_FIELD
                                    board_base.document(q_id).update(update_payload)
                                    _notify_slack(
                                        f"✏️ *Class Board post edited* — {class_name}\n",
                                        f"*By:* {student_name} ({student_code}) • QID: {q_id}\n",
                                        f"*When:* {_dt.now(UTC).strftime('%Y-%m-%d %H:%M')} UTC\n",
                                        f"*New:* {(formatted_edit[:180] + '…') if len(formatted_edit) > 180 else formatted_edit}",
                                    )
                                    _clear_typing_state(
                                        level=student_level,
                                        class_code=class_name,
                                        qid=q_id,
                                        draft_key=f"q_edit_text_{q_id}",
                                        student_code=student_code,
                                        student_name=student_name,
                                    )
                                    st.session_state[f"q_editing_{q_id}"] = False
                                    st.session_state[f"__clear_q_edit_{q_id}"] = True
                                    st.success("Post updated.")
                                    refresh_with_toast()
                            if cancel_edit:
                                _clear_typing_state(
                                    level=student_level,
                                    class_code=class_name,
                                    qid=q_id,
                                    draft_key=f"q_edit_text_{q_id}",
                                    student_code=student_code,
                                    student_name=student_name,
                                )
                                st.session_state[f"q_editing_{q_id}"] = False
                                st.session_state[f"__clear_q_edit_{q_id}"] = True
                                refresh_with_toast()

                    c_ref = board_base.document(q_id).collection("comments")
                    try:
                        comments_docs = list(c_ref.order_by("timestamp").stream())
                    except Exception:
                        comments_docs = list(c_ref.stream())
                        comments_docs.sort(key=lambda c: (c.to_dict() or {}).get("timestamp"))

                    if comments_docs:
                        for c in comments_docs:
                            cid = c.id
                            c_data = c.to_dict() or {}
                            c_label = _fmt_ts(c_data.get("timestamp"))
                            commenter = c_data.get("replied_by_name", "")
                            role = "user" if c_data.get("replied_by_code") == student_code else "assistant"
                            message = st.chat_message(role)
                            header_html = f"**{commenter}**"
                            if c_label:
                                header_html = (
                                    f"{header_html} <span style='color:#94a3b8;font-size:0.85rem;'>{c_label}</span>"
                                )
                            message.markdown(header_html, unsafe_allow_html=True)
                            message.markdown(c_data.get("content", ""))

                            clear_c_edit_flag = f"__clear_c_edit_{q_id}_{cid}"
                            if st.session_state.pop(clear_c_edit_flag, False):
                                for _k in [
                                    f"c_edit_text_{q_id}_{cid}",
                                    f"c_edit_text_input_{q_id}_{cid}",
                                ]:
                                    st.session_state.pop(_k, None)
                                _clear_typing_state(
                                    level=student_level,
                                    class_code=class_name,
                                    qid=q_id,
                                    draft_key=f"c_edit_text_{q_id}_{cid}",
                                    student_code=student_code,
                                    student_name=student_name,
                                )

                            can_modify_c = (c_data.get("replied_by_code") == student_code) or IS_ADMIN
                            if can_modify_c:
                                controls = message.container()
                                rc1, rc2, _ = controls.columns([1, 1, 6])
                                with rc1:
                                    if st.button("✏️ Edit", key=f"c_edit_btn_{q_id}_{cid}"):
                                        st.session_state[f"c_editing_{q_id}_{cid}"] = True
                                        st.session_state[f"c_edit_text_{q_id}_{cid}"] = c_data.get("content", "")
                                with rc2:
                                    if st.button("🗑️ Delete", key=f"c_del_btn_{q_id}_{cid}"):
                                        c.reference.delete()
                                        _notify_slack(
                                            f"🗑️ *Class Board comment deleted* — {class_name}\n"
                                            f"*By:* {student_name} ({student_code}) • QID: {q_id}\n"
                                            f"*When:* {_dt.now(_timezone.utc).strftime('%Y-%m-%d %H:%M')} UTC"
                                        )
                                        st.success("Comment deleted.")
                                        refresh_with_toast()

                                if st.session_state.get(f"c_editing_{q_id}_{cid}", False):
                                    edit_container = controls.container()
                                    with edit_container.form(f"c_edit_form_{q_id}_{cid}"):
                                        banner = _format_typing_banner(
                                            fetch_active_typists(
                                                student_level,
                                                class_name,
                                                q_id,
                                            ),
                                            student_code,
                                        )
                                        if banner:
                                            st.caption(banner)
                                        new_rtext = st.text_area(
                                            "Edit comment",
                                            value=st.session_state.get(f"c_edit_text_{q_id}_{cid}", ""),
                                            key=f"c_edit_text_input_{q_id}_{cid}",
                                            height=80
                                        )
                                        _update_typing_state(
                                            level=student_level,
                                            class_code=class_name,
                                            qid=q_id,
                                            draft_key=f"c_edit_text_{q_id}_{cid}",
                                            student_code=student_code,
                                            student_name=student_name,
                                            text=new_rtext,
                                        )
                                        csave = st.form_submit_button("💾 Save")
                                        ccancel = st.form_submit_button("❌ Cancel")
                                    if csave and new_rtext.strip():
                                        c.reference.update({
                                            "content": new_rtext.strip(),
                                            "edited_at": _dt.now(_timezone.utc),
                                        })
                                        _notify_slack(
                                            f"✏️ *Class Board comment edited* — {class_name}\n"
                                            f"*By:* {student_name} ({student_code}) • QID: {q_id}\n"
                                            f"*When:* {_dt.now(_timezone.utc).strftime('%Y-%m-%d %H:%M')} UTC\n"
                                            f"*New:* {(new_rtext[:180] + '…') if len(new_rtext) > 180 else new_rtext}"
                                        )
                                        _clear_typing_state(
                                            level=student_level,
                                            class_code=class_name,
                                            qid=q_id,
                                            draft_key=f"c_edit_text_{q_id}_{cid}",
                                            student_code=student_code,
                                            student_name=student_name,
                                        )
                                        st.session_state[f"c_editing_{q_id}_{cid}"] = False
                                        st.session_state[f"__clear_c_edit_{q_id}_{cid}"] = True
                                        st.success("Comment updated.")
                                        refresh_with_toast()
                                    if ccancel:
                                        _clear_typing_state(
                                            level=student_level,
                                            class_code=class_name,
                                            qid=q_id,
                                            draft_key=f"c_edit_text_{q_id}_{cid}",
                                            student_code=student_code,
                                            student_name=student_name,
                                        )
                                        st.session_state[f"c_editing_{q_id}_{cid}"] = False
                                        st.session_state[f"__clear_c_edit_{q_id}_{cid}"] = True
                                        refresh_with_toast()

                    draft_key = f"classroom_comment_draft_{q_id}"
                    last_val_key, last_ts_key, saved_flag_key, saved_at_key = _draft_state_keys(draft_key)
                    if draft_key not in st.session_state:
                        txt, ts = load_draft_meta_from_db(student_code, draft_key)
                        st.session_state[draft_key] = txt or ""
                        st.session_state[last_val_key] = st.session_state[draft_key]
                        st.session_state[last_ts_key] = time.time()
                        st.session_state[saved_flag_key] = bool(txt)
                        st.session_state[saved_at_key] = ts
                    clear_flag = f"__clear_comment_draft_{q_id}"
                    if st.session_state.pop(clear_flag, False):
                        st.session_state[draft_key] = ""
                        _clear_typing_state(
                            level=student_level,
                            class_code=class_name,
                            qid=q_id,
                            draft_key=draft_key,
                            student_code=student_code,
                            student_name=student_name,
                        )
                    def apply_ai_correction(q_id: str, draft_key: str, current_text: str) -> None:
                        if not current_text.strip():
                            return
                        try:
                            resp = client.chat.completions.create(
                                model="gpt-4o-mini",
                                messages=[
                                    {
                                        "role": "system",
                                        "content": (
                                            "You are a helpful assistant that corrects German replies. "
                                            "Return only the corrected reply."
                                        ),
                                    },
                                    {
                                        "role": "user",
                                        "content": f"Question: {q.get('content','')}\nReply: {current_text}",
                                    },
                                ],
                                temperature=0,
                                max_tokens=300,
                            )
                            ai_text = (resp.choices[0].message.content or "").strip()
                            flagged = resp.choices[0].finish_reason == "content_filter"
                        except Exception:
                            ai_text = ""
                            flagged = False
                        if ai_text:
                            st.session_state[draft_key] = ai_text
                            save_ai_response(q_id, ai_text, flagged)

                    ai_flag = f"q_ai_busy_{q_id}"

                    current_text = st.session_state.get(draft_key, "")
                    if not isinstance(current_text, str):
                        current_text = ""
                    if st.session_state.get(ai_flag):
                        with st.spinner("Correcting with AI..."):
                            apply_ai_correction(q_id, draft_key, current_text)
                        st.session_state[ai_flag] = False
                        current_text = st.session_state.get(draft_key, "")
                        if not isinstance(current_text, str):
                            current_text = ""

                    banner = _format_typing_banner(
                        fetch_active_typists(student_level, class_name, q_id),
                        student_code,
                    )
                    if banner:
                        st.caption(banner)
                    reply_timer_label = build_forum_reply_indicator_text(timer_info)
                    if reply_timer_label:
                        reply_color = "#dc2626" if timer_info.get("status") == "open" else "#64748b"
                        st.markdown(
                            "<div style='font-size:0.95rem;font-weight:600;color:%s;margin:6px 0 -4px;'>%s</div>"
                            % (
                                reply_color,
                                html.escape(str(reply_timer_label)),
                            ),
                            unsafe_allow_html=True,
                        )
                    if show_timer_warning:
                        st.info(
                            "⏳ Time up soon—replies close in under a minute.",
                            icon="⏳",
                        )
                    st.text_area(
                        "Reply to this thread…",
                        key=draft_key,
                        placeholder="Reply to this thread…",
                        label_visibility="collapsed",
                    )
                    render_umlaut_pad(
                        draft_key,
                        context=f"classboard_reply_{q_id}",
                    )

                    current_text = st.session_state.get(draft_key, "")
                    if not isinstance(current_text, str):
                        current_text = ""
                    _update_typing_state(
                        level=student_level,
                        class_code=class_name,
                        qid=q_id,
                        draft_key=draft_key,
                        student_code=student_code,
                        student_name=student_name,
                        text=current_text,
                    )
                    autosave_maybe(student_code, draft_key, current_text, min_secs=2.0, min_delta=12)

                    send_col, ai_col = st.columns([1, 1])

                    with send_col:
                        if st.button(
                            "Send Reply",
                            key=f"q_send_comment_{q_id}",
                            type="primary",
                            width="stretch",
                        ):
                            if not current_text.strip():
                                st.warning("Type a reply first.")
                            else:
                                save_now(draft_key, student_code, show_toast=False)
                                send_comment(
                                    q_id,
                                    student_code,
                                    student_name,
                                    class_name,
                                    board_base,
                                    draft_key,
                                    last_val_key,
                                    last_ts_key,
                                    saved_flag_key,
                                    saved_at_key,
                                )
                                st.rerun()

                    with ai_col:
                        if st.button(
                            "✨ Correct with AI",
                            key=f"q_ai_btn_{q_id}",
                            disabled=st.session_state.get(ai_flag, False),
                            width="stretch",
                        ):
                            st.session_state[ai_flag] = True
                            st.rerun()

                    if idx < len(questions) - 1:
                        st.divider()


    # === LEARNING NOTES SUBTAB ===
    elif cb_subtab == "📒 Learning Notes":
        st.markdown("""
            <div style="padding: 14px; background: #8d4de8; color: #fff; border-radius: 8px; 
            text-align:center; font-size:1.5rem; font-weight:700; margin-bottom:16px; letter-spacing:.5px;">
            📒 My Learning Notes
            </div>
        """, unsafe_allow_html=True)

        student_code = st.session_state.get("student_code", "") or ""
        if not student_code:
            st.error("Student code is required to view notes.")
            st.stop()
        key_notes = f"notes_{student_code}"

        if key_notes not in st.session_state:
            st.session_state[key_notes] = load_notes_from_db(student_code)
        notes = st.session_state[key_notes]

        if st.session_state.get("switch_to_edit_note"):
            st.session_state["course_notes_radio"] = "➕ Add/Edit Note"
            del st.session_state["switch_to_edit_note"]
        elif st.session_state.get("switch_to_library"):
            st.session_state["course_notes_radio"] = "📚 My Notes Library"
            del st.session_state["switch_to_library"]

        notes_subtab = st.radio(
            "Notebook",
            ["➕ Add/Edit Note", "📚 My Notes Library"],
            horizontal=True,
            key="course_notes_radio"
        )

        if notes_subtab == "➕ Add/Edit Note":
            # >>>> New helper message for pre-filled note context <<<<
            editing = st.session_state.get("edit_note_idx", None) is not None
            if editing:
                idx = st.session_state["edit_note_idx"]
                title = st.session_state.get("edit_note_title", "")
                tag = st.session_state.get("edit_note_tag", "")
                text = st.session_state.get("edit_note_text", "")
            else:
                title, tag, text = "", "", ""

            student_level = st.session_state.get("student_level", "A1")
            lesson_choices = []
            try:
                schedules = load_level_schedules()
                level_sched = schedules.get(student_level, schedules.get("A1", []))
                for item in level_sched:
                    day = item.get("day")
                    topic = item.get("topic")
                    if day is not None and topic:
                        lesson_choices.append(
                            f"Day {day} - {topic} (Chapter {item.get('chapter', '?')})"
                        )
            except Exception:
                pass

            if title and tag:
                st.info(f"You're adding a note for **{title}** ({tag}).")

            st.markdown("#### ✍️ Create a new note or update an old one")
            if st.session_state.pop("reset_note_form", False):
                for k in [
                    "learning_note_title",
                    "learning_note_tag",
                    "learning_note_draft",
                    "learning_note_last_saved",
                    "learning_note_lesson",
                ]:
                    st.session_state.pop(k, None)

            with st.form("note_form", clear_on_submit=True):
                st.session_state.setdefault("learning_note_title", title)
                st.session_state.setdefault("learning_note_tag", tag)
                st.session_state.setdefault("learning_note_draft", text)
                st.session_state.setdefault("learning_note_last_saved", None)
                st.session_state.setdefault(
                    "learning_note_lesson",
                    st.session_state.get(
                        "edit_note_lesson",
                        lesson_choices[0] if lesson_choices else "",
                    ),
                )

                st.selectbox("Lesson", lesson_choices, key="learning_note_lesson")

                st.text_input(
                    "Note Title",
                    max_chars=50,
                    key="learning_note_title",
                )
                st.text_input(
                    "Category/Tag (optional)",
                    max_chars=20,
                    key="learning_note_tag",
                )
                ta_col, ai_col = st.columns([3, 1])
                with ta_col:
                    st.text_area(
                        "Your Note",
                        height=200,
                        max_chars=3000,
                        key="learning_note_draft",
                    )
                with ai_col:
                    ai_btn = st.form_submit_button("✨ Correct with AI")

                col1, col2 = st.columns(2)
                save_btn = col1.form_submit_button("Save")
                cancel_btn = editing and col2.form_submit_button("❌ Cancel Edit")
                if save_btn:
                    autosave_learning_note(student_code, key_notes)
                    if not editing:
                        st.session_state["reset_note_form"] = True
                        st.session_state["need_rerun"] = True
                if st.session_state.get("learning_note_last_saved"):
                    st.caption(
                        f"Last saved {st.session_state['learning_note_last_saved']} UTC"
                    )

            if ai_btn:
                with st.spinner("Correcting with AI..."):
                    original = st.session_state.get("learning_note_draft", "")
                    improved, explanation = apply_note_ai_correction(original)
                    st.session_state["note_ai_suggestion"] = improved
                    st.session_state["note_ai_explanation"] = explanation
                    st.session_state["note_ai_diff"] = diff_with_markers(original, improved)
                st.session_state["need_rerun"] = True

            if st.session_state.get("note_ai_diff"):
                st.markdown(st.session_state["note_ai_diff"], unsafe_allow_html=True)
                st.markdown("**Why these changes?**")
                st.markdown(st.session_state.get("note_ai_explanation", ""))
                acc_col, rej_col = st.columns(2)
                with acc_col:
                    if st.button("Accept", key="note_ai_accept"):
                        st.session_state["learning_note_draft"] = st.session_state.get(
                            "note_ai_suggestion", ""
                        )
                        st.session_state.pop("note_ai_suggestion", None)
                        st.session_state.pop("note_ai_explanation", None)
                        st.session_state.pop("note_ai_diff", None)
                        st.session_state["need_rerun"] = True
                with rej_col:
                    if st.button("Reject", key="note_ai_reject"):
                        st.session_state.pop("note_ai_suggestion", None)
                        st.session_state.pop("note_ai_explanation", None)
                        st.session_state.pop("note_ai_diff", None)
                        st.session_state["need_rerun"] = True

            if cancel_btn:

                for k in [
                    "edit_note_idx",
                    "edit_note_title",
                    "edit_note_text",
                    "edit_note_tag",
                    "edit_note_lesson",
                    "learning_note_title",
                    "learning_note_tag",
                    "learning_note_draft",
                    "learning_note_lesson",
                    "learning_note_last_saved",
                ]:
                    if k in st.session_state:
                        del st.session_state[k]

                st.session_state["switch_to_library"] = True
                refresh_with_toast()

        elif notes_subtab == "📚 My Notes Library":
            st.markdown("#### 📚 All My Notes")

            if not notes:
                st.info("No notes yet. Add your first note in the ➕ tab!")
            else:
                search_term = st.text_input("🔎 Search your notes…", "")
                if search_term.strip():
                    filtered = []
                    st.markdown('<div style="height:10px"></div>', unsafe_allow_html=True)
                    for n in notes:
                        if (
                            search_term.lower() in n.get("title", "").lower()
                            or search_term.lower() in n.get("tag", "").lower()
                            or search_term.lower() in n.get("text", "").lower()
                            or search_term.lower() in n.get("lesson", "").lower()
                        ):
                            filtered.append(n)
                    notes_to_show = filtered
                    if not filtered:
                        st.warning("No matching notes found!")
                else:
                    notes_to_show = notes

                # --- Download Buttons (TXT, PDF, DOCX) FOR ALL NOTES ---
                all_notes = []
                for n in notes_to_show:
                    note_text = f"Title: {n.get('title','')}\n"
                    if n.get('tag'):
                        note_text += f"Tag: {n['tag']}\n"
                    if n.get("lesson"):
                        note_text += f"Lesson: {n['lesson']}\n"
                    note_text += n.get('text','') + "\n"
                    note_text += f"Date: {n.get('updated', n.get('created',''))}\n"
                    note_text += "-"*32 + "\n"
                    all_notes.append(note_text)
                txt_data = "\n".join(all_notes)

                st.download_button(
                    label="⬇️ Download All Notes (TXT)",
                    data=txt_data.encode("utf-8"),
                    file_name=f"{student_code}_notes.txt",
                    mime="text/plain"
                )

                # --- PDF Download (all notes, Unicode/emoji ready!) ---

                pdf_bytes = generate_notes_pdf(notes_to_show)

                st.download_button(
                    label="⬇️ Download All Notes (PDF)",
                    data=pdf_bytes,
                    file_name=f"{student_code}_notes.pdf",
                    mime="application/pdf",
                )

                # --- DOCX Download (all notes) ---
                def export_notes_to_docx(notes, student_code="student"):
                    doc = Document()
                    doc.add_heading("My Learning Notes", 0)
                    doc.add_heading("Table of Contents", level=1)
                    for idx, note in enumerate(notes):
                        doc.add_paragraph(f"{idx+1}. {note.get('title', '(No Title)')} - {note.get('created', note.get('updated',''))}")
                    doc.add_page_break()
                    for note in notes:
                        doc.add_heading(note.get('title','(No Title)'), level=1)
                        if note.get("tag"):
                            doc.add_paragraph(f"Tag: {note.get('tag','')}")
                        if note.get("lesson"):
                            doc.add_paragraph(f"Lesson: {note.get('lesson','')}")
                        doc.add_paragraph(note.get('text', ''))
                        doc.add_paragraph(f"Date: {note.get('created', note.get('updated',''))}")
                        doc.add_paragraph('-' * 40)
                        doc.add_paragraph("")
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as f:
                        doc.save(f.name)
                        return f.name
                docx_path = export_notes_to_docx(notes_to_show, student_code)
                with open(docx_path, "rb") as f:
                    st.download_button(
                        label="⬇️ Download All Notes (DOCX)",
                        data=f.read(),
                        file_name=f"{student_code}_notes.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                os.remove(docx_path)

                st.markdown("---")
                pinned_notes = [n for n in notes_to_show if n.get("pinned")]
                other_notes = [n for n in notes_to_show if not n.get("pinned")]
                show_list = pinned_notes + other_notes
                for i, note in enumerate(show_list):
                    st.markdown(
                        f"<div style='padding:12px 0 6px 0; font-weight:600; color:#7c3aed; font-size:1.18rem;'>"
                        f"{'📌 ' if note.get('pinned') else ''}{note.get('title','(No Title)')}"
                        f"</div>", unsafe_allow_html=True)
                    if note.get("tag"):
                        st.caption(f"🏷️ Tag: {note['tag']}")
                    if note.get("lesson"):
                        st.caption(f"📘 Lesson: {note['lesson']}")
                    st.markdown(
                        f"<div style='margin-top:-5px; margin-bottom:6px; font-size:1.08rem; line-height:1.7;'>{note['text'].replace('\n', '<br>')}</div>",
                        unsafe_allow_html=True)
                    st.caption(f"🕒 {note.get('updated',note.get('created',''))}")

                    # --- Per-Note Download Buttons (TXT, PDF, DOCX) ---
                    download_cols = st.columns([1,1,1])
                    with download_cols[0]:
                        # TXT per note
                        txt_note = f"Title: {note.get('title','')}\n"
                        if note.get('tag'):
                            txt_note += f"Tag: {note['tag']}\n"
                        if note.get("lesson"):
                            txt_note += f"Lesson: {note['lesson']}\n"
                        txt_note += note.get('text', '') + "\n"
                        txt_note += f"Date: {note.get('updated', note.get('created',''))}\n"
                        st.download_button(
                            label="⬇️ TXT",
                            data=txt_note.encode("utf-8"),
                            file_name=f"{student_code}_{note.get('title','note').replace(' ','_')}.txt",
                            mime="text/plain",
                            key=f"download_txt_{i}"
                        )
                    with download_cols[1]:
                        # PDF per note (Unicode/emoji ready!)
                        pdf_bytes_single = generate_single_note_pdf(note)
                        st.download_button(
                            label="⬇️ PDF",
                            data=pdf_bytes_single,
                            file_name=f"{student_code}_{note.get('title','note').replace(' ','_')}.pdf",
                            mime="application/pdf",
                            key=f"download_pdf_{i}",
                        )
                    with download_cols[2]:
                        # DOCX per note
                        doc_single = Document()
                        doc_single.add_heading(note.get('title','(No Title)'), level=1)
                        if note.get("tag"):
                            doc_single.add_paragraph(f"Tag: {note.get('tag','')}")
                        if note.get("lesson"):
                            doc_single.add_paragraph(f"Lesson: {note.get('lesson','')}")
                        doc_single.add_paragraph(note.get('text', ''))
                        doc_single.add_paragraph(f"Date: {note.get('updated', note.get('created',''))}")
                        single_docx_io = io.BytesIO()
                        doc_single.save(single_docx_io)
                        st.download_button(
                            label="⬇️ DOCX",
                            data=single_docx_io.getvalue(),
                            file_name=f"{student_code}_{note.get('title','note').replace(' ','_')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"download_docx_{i}"
                        )

                    cols = st.columns([1,1,1,1])
                    with cols[0]:
                        if st.button("✏️ Edit", key=f"edit_{i}"):
                            st.session_state["edit_note_idx"] = i
                            st.session_state["edit_note_title"] = note["title"]
                            st.session_state["edit_note_text"] = note["text"]
                            st.session_state["edit_note_tag"] = note.get("tag", "")
                            st.session_state["edit_note_lesson"] = note.get("lesson", "")
                            st.session_state["switch_to_edit_note"] = True
                            refresh_with_toast()
                    with cols[1]:
                        if st.button("🗑️ Delete", key=f"del_{i}"):
                            notes.remove(note)
                            st.session_state[key_notes] = notes
                            save_notes_to_db(student_code, notes)
                            st.success("Note deleted.")
                            refresh_with_toast()
                    with cols[2]:
                        if note.get("pinned"):
                            if st.button("📌 Unpin", key=f"unpin_{i}"):
                                note["pinned"] = False
                                st.session_state[key_notes] = notes
                                save_notes_to_db(student_code, notes)
                                refresh_with_toast()
                        else:
                            if st.button("📍 Pin", key=f"pin_{i}"):
                                note["pinned"] = True
                                st.session_state[key_notes] = notes
                                save_notes_to_db(student_code, notes)
                                refresh_with_toast()
                    with cols[3]:
                        st.caption("")



# =========================== MY RESULTS & RESOURCES ===========================
if tab == "My Results and Resources":
    render_results_and_resources_tab()


@st.cache_data
def build_dict_df(levels):
    rows = []
    sentence_map = {}

    # Collect sentences for each token in the sentence bank
    for lvl in levels:
        for item in SENTENCE_BANK.get(lvl, []):
            sentence = item.get("target_de", "")
            for tok in item.get("tokens", []):
                t = str(tok).strip()
                if not t or t in [",", ".", "!", "?", ":", ";"]:
                    continue
                sentence_map.setdefault((lvl, t), sentence)

    # Build initial rows from the vocab lists
    for lvl in levels:
        for entry in VOCAB_LISTS.get(lvl, []):
            de = entry[0]
            en = entry[1]
            pron = entry[2] if len(entry) > 2 else ""
            sent = sentence_map.get((lvl, de), "")
            rows.append(
                {
                    "Level": lvl,
                    "German": de,
                    "English": en,
                    "Pronunciation": pron,
                    "Sentence": sent,
                }
            )

    df = (
        pd.DataFrame(rows)
        if rows
        else pd.DataFrame(
            columns=["Level", "German", "English", "Pronunciation", "Sentence"]
        )
    )

    # Add extra tokens that appear in the sentence bank but not in the vocab list
    extra = []
    for (lvl, t), sent in sentence_map.items():
        if not ((df["German"] == t) & (df["Level"] == lvl)).any():
            extra.append(
                {
                    "Level": lvl,
                    "German": t,
                    "English": "",
                    "Pronunciation": "",
                    "Sentence": sent,
                }
            )
    if extra:
        df = pd.concat([df, pd.DataFrame(extra)], ignore_index=True)

    if not df.empty:
        df = df.drop_duplicates(subset=["Level", "German"]).reset_index(drop=True)
    return df


def render_vocab_trainer_section() -> None:
    # --- Who is this? ---
    student_code = st.session_state.get("student_code", "") or ""
    if not student_code:
        st.error("Student code is required to access the vocab trainer.")
        return

    # --- Lock the level from your Sheet/profile ---
    student_level_locked = (
        get_student_level(student_code, default=None)
        or st.session_state.get("student_level")
        or "A1"
    )
    # Header
    st.markdown(
        """
        <div style="
            padding:8px 12px; background:#6f42c1; color:#fff;
            border-radius:6px; text-align:center; margin-bottom:8px;
            font-size:1.3rem;">
        📚 Vocab Trainer
        </div>
        """,
        unsafe_allow_html=True
    )
    st.markdown(f"**Practicing Level:** `{student_level_locked}` (from your profile)")
    st.caption("Your level is loaded automatically from the school list. Ask your tutor if this looks wrong.")
    st.divider()

    subtab = st.radio(
        "Choose practice:",
        ["Sentence Builder", "Vocab Practice", "Dictionary"],
        horizontal=True,
        key="vocab_practice_subtab"
    )

    # ===========================
    # SUBTAB: Sentence Builder  (unchanged logic, audio not needed here)
    # ===========================
    if subtab == "Sentence Builder":
        render_sentence_builder(student_code, student_level_locked)

    # ===========================
    # SUBTAB: Vocab Practice  (download-only audio)
    # ===========================
    elif subtab == "Vocab Practice":
        defaults = {
            "vt_history": [], "vt_list": [], "vt_index": 0,
            "vt_score": 0, "vt_total": None, "vt_saved": False, "vt_session_id": None,
            "vt_mode": "Only new words",
        }
        for k, v in defaults.items():
            st.session_state.setdefault(k, v)

        # Stats
        stats = render_vocab_stats(student_code)

        # Level lock
        level = student_level_locked
        items = VOCAB_LISTS.get(level, [])
        completed = set(stats["completed_words"])
        not_done = [p for p in items if p[0] not in completed]
        st.info(f"{len(not_done)} words NOT yet done at {level}.")

        if st.button("🔁 Start New Practice", key="vt_reset"):
            for k in defaults:
                st.session_state[k] = defaults[k]
            refresh_with_toast()

        if st.session_state.vt_total is None:
            with st.form("vt_setup"):
                st.subheader("Daily Practice Setup")
                mode = st.radio(
                    "Select words:",
                    ["Only new words", "All words"],
                    horizontal=True,
                    key="vt_mode",
                )
                session_vocab = (not_done if mode == "Only new words" else items).copy()
                maxc = len(session_vocab)
                if maxc == 0:
                    st.success("🎉 All done! Switch to 'All words' to repeat.")
                    return
                count = st.number_input(
                    "How many today?", 1, maxc, min(7, maxc), key="vt_count"
                )
                submitted = st.form_submit_button("Start")
            if submitted:
                import random
                from uuid import uuid4
                random.shuffle(session_vocab)
                st.session_state.vt_list = session_vocab[:count]
                st.session_state.vt_total = count
                st.session_state.vt_index = 0
                st.session_state.vt_score = 0
                st.session_state.vt_history = [
                    ("assistant", f"Hallo! Ich bin Herr Felix. Let's do {count} words!")
                ]
                st.session_state.vt_saved = False
                st.session_state.vt_session_id = str(uuid4())
                refresh_with_toast()
        else:
            st.markdown("### Daily Practice Setup")
            st.info(
                f"{st.session_state.vt_total} words · {st.session_state.get('vt_mode')}"
            )
            if st.button("Change goal", key="vt_change_goal"):
                st.session_state.vt_total = None
                refresh_with_toast()

        tot = st.session_state.vt_total
        idx = st.session_state.vt_index
        score = st.session_state.vt_score

        if st.session_state.vt_history:
            if isinstance(tot, int) and tot:
                remaining = tot - idx
                c1, c2 = st.columns(2)
                with c1:
                    st.metric("Words", f"{idx}/{tot}", f"{remaining} left")
                    st.progress(idx / tot)
                with c2:
                    st.metric("Score", score)

            st.markdown("### 🗨️ Practice Chat")
            for who, msg in st.session_state.vt_history:
                render_message(who, msg)

        if isinstance(tot, int) and idx < tot:
            current = st.session_state.vt_list[idx]
            word = current[0]
            answer = current[1]

            # ---- AUDIO (download-only: prefer sheet link; fallback to gTTS bytes) ----
            audio_url = get_audio_url(level, word)
            if audio_url:
                st.markdown(f"[⬇️ Download / Open MP3]({audio_url})")
            else:
                audio_bytes = _dict_tts_bytes_de(word)  # fallback generation
                if audio_bytes:
                    st.download_button(
                        "⬇️ Download MP3",
                        data=audio_bytes,
                        file_name=f"{word}.mp3",
                        mime="audio/mpeg",
                        key=f"dl_{idx}"
                    )
                else:
                    st.caption("Audio not available yet.")

            # nicer input styling
            st.markdown(
                """
                <style>
                div[data-baseweb="input"] input { font-size: 18px !important; font-weight: 600 !important; color: black !important; }
                </style>
                """,
                unsafe_allow_html=True
            )

            usr = st.text_input(
                f"{word} = ?",
                key=f"vt_input_{idx}",
                placeholder="Type your answer here...",
            )
            render_umlaut_pad(
                f"vt_input_{idx}",
                context=f"vocab_practice_{student_code}",
            )
            if usr and st.button("Check", key=f"vt_check_{idx}"):
                st.session_state.vt_history.append(("user", usr))
                if is_correct_answer(usr, answer):
                    st.session_state.vt_score += 1
                    fb = f"✅ Correct! '{word}' = '{answer}'"
                else:
                    fb = f"❌ Nope. '{word}' = '{answer}'"
                st.session_state.vt_history.append(("assistant", fb))
                st.session_state.vt_index += 1
                refresh_with_toast()

        if isinstance(tot, int) and idx >= tot:
            score = st.session_state.vt_score
            words = [item[0] for item in (st.session_state.vt_list or [])]
            st.markdown(f"### 🏁 Done! You scored {score}/{tot}.")
            if not st.session_state.get("vt_saved", False):
                if not st.session_state.get("vt_session_id"):
                    from uuid import uuid4
                    st.session_state.vt_session_id = str(uuid4())
                if not vocab_attempt_exists(student_code, st.session_state.vt_session_id):
                    save_vocab_attempt(
                        student_code=student_code,
                        level=level,
                        total=tot,
                        correct=score,
                        practiced_words=words,
                        session_id=st.session_state.vt_session_id
                    )
                st.session_state.vt_saved = True
                refresh_with_toast()
            if st.button("Practice Again", key="vt_again"):
                for k in defaults:
                    st.session_state[k] = defaults[k]
                refresh_with_toast()

    # ===========================
    # SUBTAB: Dictionary  (download-only audio)
    # ===========================
    elif subtab == "Dictionary":
        import io
        import json
        import difflib
        import pandas as pd

        # functions used here
        _map = {"ä":"ae","ö":"oe","ü":"ue","ß":"ss"}
        def _norm(s: str) -> str:
            s = (s or "").strip().lower()
            for k,v in _map.items(): s = s.replace(k, v)
            return "".join(ch for ch in s if ch.isalnum() or ch.isspace())

        # Build data from selected levels
        available_levels = sorted(VOCAB_LISTS.keys())
        has_unknown = False
        if "nan" in available_levels:
            available_levels = [lvl for lvl in available_levels if lvl != "nan"]
            available_levels.append("Unknown level")
            has_unknown = True
        if has_unknown:
            st.info("Words without a level are listed under 'Unknown level'.")
        default_levels = [student_level_locked] if student_level_locked in available_levels else []
        levels_display = st.multiselect(
            "Select level(s)",
            available_levels,
            default=default_levels,
            key="dict_levels",
        )
        levels = ["nan" if lvl == "Unknown level" else lvl for lvl in levels_display]
        df_dict = build_dict_df(levels)
        for c in ["Level","German","English","Pronunciation"]:
            if c not in df_dict.columns: df_dict[c] = ""
        df_dict["g_norm"] = df_dict["German"].astype(str).map(_norm)
        df_dict["e_norm"] = df_dict["English"].astype(str).map(_norm)
        df_dict = df_dict.sort_values(["German"]).reset_index(drop=True)

        # Sticky search UI
        st.markdown(
            """
            <style>
              .sticky-search { position: sticky; top: 0; z-index: 999; background: white; padding: 8px 0 10px 0; }
              input[type="text"] { font-size: 18px !important; }
              .chip { display:inline-block; padding:6px 10px; border-radius:999px; border:1px solid #e5e7eb; margin-right:6px; margin-bottom:6px; }
            </style>
            """,
            unsafe_allow_html=True
        )
        with st.container():
            st.markdown('<div class="sticky-search">', unsafe_allow_html=True)
            cols = st.columns([6, 3, 3])
            with cols[0]:
                pending_dict_q = st.session_state.pop("dict_q_pending", None)
                if pending_dict_q is not None:
                    st.session_state["dict_q"] = pending_dict_q
                q = st.text_input("🔎 Search (German or English)", key="dict_q", placeholder="e.g., Wochenende, bakery, spielen")
            with cols[1]:
                search_in = st.selectbox("Field", ["Both", "German", "English"], 0, key="dict_field")
            with cols[2]:
                match_mode = st.selectbox("Match", ["Contains", "Starts with", "Exact"], 0, key="dict_mode")
            st.markdown('</div>', unsafe_allow_html=True)

        # Filter + choose top row
        df_view = df_dict.copy()
        suggestions = []
        top_row = None

        if q:
            qn = _norm(q)
            g_contains = df_view["g_norm"].str.contains(qn, na=False) if search_in in ("Both","German") else pd.Series([False]*len(df_view))
            g_starts   = df_view["g_norm"].str.startswith(qn, na=False) if search_in in ("Both","German") else pd.Series([False]*len(df_view))
            g_exact    = df_view["g_norm"].eq(qn) if search_in in ("Both","German") else pd.Series([False]*len(df_view))
            e_contains = df_view["e_norm"].str.contains(qn, na=False) if search_in in ("Both","English") else pd.Series([False]*len(df_view))
            e_starts   = df_view["e_norm"].str.startswith(qn, na=False) if search_in in ("Both","English") else pd.Series([False]*len(df_view))
            e_exact    = df_view["e_norm"].eq(qn) if search_in in ("Both","English") else pd.Series([False]*len(df_view))

            mask = (g_contains | e_contains) if match_mode=="Contains" else (g_starts | e_starts) if match_mode=="Starts with" else (g_exact | e_exact)
            if mask.any():
                exact_mask = (g_exact | e_exact) & mask
                starts_mask = (g_starts | e_starts) & mask
                df_view = df_view[mask].reset_index(drop=True)
                exact_mask = exact_mask[mask].reset_index(drop=True)
                starts_mask = starts_mask[mask].reset_index(drop=True)
                if not df_view.empty:
                    top_row = df_view[exact_mask].iloc[0] if exact_mask.any() else df_view[starts_mask].iloc[0] if starts_mask.any() else df_view.iloc[0]
            else:
                vocab_all = df_view["German"].astype(str).unique().tolist()
                suggestions = difflib.get_close_matches(q, vocab_all, n=5, cutoff=0.72)
                if not suggestions:
                    st.info("No matches found.")
                # Still show a card for the query itself
                dummy = {"Level": student_level_locked, "German": q, "English": "", "Pronunciation": "", "g_norm": qn, "e_norm": ""}
                df_view = pd.concat([df_view, pd.DataFrame([dummy])], ignore_index=True)
                top_row = pd.Series(dummy)
        else:
            if not df_view.empty: top_row = df_view.iloc[0]

        # Details panel (download-only audio)
        if top_row is not None and len(top_row) > 0:
            de  = str(top_row["German"])
            en  = str(top_row.get("English", "") or "")
            lvl = str(top_row.get("Level", student_level_locked))

            st.markdown(f"### {de}")
            if en: st.markdown(f"**Meaning:** {en}")
            pron = str(top_row.get("Pronunciation", "") or "").strip()
            if pron:
                st.markdown(f"**Pronunciation:** {pron}")

            # Show first example sentence containing the word
            example_sentence = ""
            for item in SENTENCE_BANK.get(lvl, []):
                tokens = [str(tok).strip().lower() for tok in item.get("tokens", [])]
                if de.lower() in tokens:
                    example_sentence = item.get("target_de") or " ".join(item.get("tokens", []))
                    break
            if example_sentence:
                st.markdown(example_sentence)

            sheet_audio = get_audio_url(lvl, de)
            sheet_audio = prepare_audio_url(sheet_audio) if sheet_audio else None
            if sheet_audio:
                render_audio_player(sheet_audio, verified=True)
                st.markdown(f"[⬇️ Download / Open MP3]({sheet_audio})")
            else:
                audio_bytes = _dict_tts_bytes_de(de)
                if audio_bytes:
                    render_audio_player(audio_bytes)
                    st.download_button(
                        "⬇️ Download MP3",
                        data=audio_bytes,
                        file_name=f"{de}.mp3",
                        mime="audio/mpeg",
                        key=f"dl_{de}_{lvl}",
                    )
                else:
                    st.caption("Audio not available yet.")

        if q and suggestions:
            st.markdown("**Did you mean:**")
            bcols = st.columns(min(5, len(suggestions)))
            for i, s in enumerate(suggestions[:5]):
                with bcols[i]:
                    if st.button(s, key=f"sugg_{i}"):
                        st.session_state["dict_q_pending"] = s
                        refresh_with_toast()

        levels_label = ", ".join(levels) if levels else "none"
        with st.expander(
            f"Browse all words for levels: {levels_label}", expanded=False
        ):
            df_show = df_view[["German", "English"]].copy()
            st.dataframe(df_show, width="stretch", height=420)







#Maincode for me

if tab == "Chat • Grammar • Exams":
    st.markdown("## 🗣️ Chat • Grammar • Exams")
    st.caption("Simple & clear: last 3 messages shown; input stays below. 3 keywords • 6 questions.")

    # ---- Links
    RECORDER_URL = "https://script.google.com/macros/s/AKfycbzMIhHuWKqM2ODaOCgtS7uZCikiZJRBhpqv2p6OyBmK1yAVba8HlmVC1zgTcGWSTfrsHA/exec"
    PRACTICE_URL = "https://script.google.com/macros/s/AKfycbyJ5lTeXUgaGw-rejDuh_2ex7El_28JgKLurOOsO1c8LWfVE-Em2-vuWuMn1hC5-_IN/exec"

    # ---------- Styles (bubbles, chips, sticky input, contrast buttons, cards) ----------
    st.markdown("""
    <style>
      .bubble-wrap{ display:flex; gap:8px; margin:8px 0; align-items:flex-start; }
      .bubble-a{ background:#fffbe6; border:1px solid #fde68a; padding:12px 14px; border-radius:14px; line-height:1.55; max-width:92%; }
      .bubble-u{ background:#eef2ff; border:1px solid #c7d2fe; padding:12px 14px; border-radius:14px; line-height:1.55; margin-left:auto; max-width:92%; }
      .lbl-a{ font-size:.8rem; color:#7c2d12; font-weight:800; margin-bottom:4px; }
      .lbl-u{ font-size:.8rem; color:#1e40af; font-weight:800; text-align:right; margin:0 4px 4px 0; }
      .kw-title{ font-weight:900; font-size:1.05rem; margin:2px 0 6px 0; }
      .kw-chip{ display:inline-block; margin:0 6px 6px 0; padding:6px 12px; border-radius:999px; border:1px solid #065f46; background:#d1fae5; font-weight:800; }
      .sticky-input{ position:sticky; bottom:0; background:#f8fafc; padding:10px 8px;
                     border-top:1px solid #e5e7eb; box-shadow:0 -6px 20px rgba(2,6,23,.06); border-radius:12px; }
      .btn-row{ display:flex; flex-wrap:wrap; gap:10px; }
      .btn {
        display:inline-block; padding:12px 16px; border-radius:12px; text-decoration:none; font-weight:900;
        background:#111827; color:#ffffff; border:1px solid #0b1220; box-shadow:0 4px 10px rgba(0,0,0,.25);
      }
      .btn.secondary{ background:#065f46; border-color:#064e3b; }
      .rec-banner{
        border:1px solid #a7f3d0; background:#ecfdf5; color:#064e3b;
        padding:10px 12px; border-radius:10px; margin:6px 0 10px 0; font-weight:700;
      }
      .panel { border:1px solid rgba(148,163,184,.35); border-radius:12px; padding:12px 14px; background:#ffffff; }
      .pres-card{
        border:2px solid #c7d2fe; background:#eef2ff; border-radius:12px; padding:12px 14px; margin-top:10px;
      }
      .pres-head{ display:flex; justify-content:space-between; align-items:center; gap:10px; }
      .pres-copy{ padding:6px 10px; border-radius:8px; border:1px solid #1f2937; background:#111827; color:#fff; font-weight:800; }
      .meta{ color:#64748b; font-size:.9rem; }
      .typing-notice{
        display:flex; align-items:center; gap:8px; margin-bottom:8px;
        background:#e0f2fe; border:1px solid #bae6fd; color:#0f172a;
        font-weight:700; font-size:.9rem; padding:8px 12px; border-radius:12px;
      }
      .typing-notice .typing span{ background:#0284c7; }
      @media (max-width: 640px){
        .block-container{ padding-bottom: 4.5rem; }
        .btn{ width:100%; text-align:center; }
        button,[role="button"]{ min-height:48px; }
      }
      .typing span{ display:inline-block; width:6px; height:6px; margin:0 2px; border-radius:50%;
                    background:#94a3b8; opacity:.2; animation: t 1s infinite; }
      .typing span:nth-child(2){ animation-delay:.15s; } .typing span:nth-child(3){ animation-delay:.3s; }
      @keyframes t{ 0%{opacity:.2; transform:translateY(0)} 50%{opacity:1; transform:translateY(-2px)} 100%{opacity:.2; transform:translateY(0)} }
    </style>
    """, unsafe_allow_html=True)

    student_code_tc = (st.session_state.get("student_code") or "").strip()
    student_row_chat = st.session_state.get("student_row") or {}
    student_display_name = (
        _safe_str(student_row_chat.get("Name"))
        or _safe_str(st.session_state.get("student_name"))
        or "Student"
    )
    student_label_html = html.escape(student_display_name)

    def _resolve_topic_coach_db():
        """Return a Firestore client for Topic Coach persistence if available."""

        global db  # type: ignore  # Streamlit runtime assigns this at module scope
        existing = globals().get("db")
        if existing is None:
            existing = getattr(_falowen_sessions, "db", None) or getattr(
                _falowen_sessions, "_db_client", None
            )
        if existing is not None:
            return existing

        getter = getattr(_falowen_sessions, "get_db", None)
        if callable(getter):
            try:
                existing = getter()
            except Exception as exc:
                logging.debug("Topic Coach Firestore unavailable: %s", exc)
                return None
            if existing is not None:
                try:
                    _falowen_sessions.db = existing
                except Exception:
                    pass
                if hasattr(_falowen_sessions, "_db_client"):
                    try:
                        _falowen_sessions._db_client = existing
                    except Exception:
                        pass
                db = existing
                return existing
        return None

    topic_db = _resolve_topic_coach_db()
    topic_doc_ref = None
    topic_messages: List[Dict[str, Any]] = []
    topic_meta: Dict[str, Any] = {}
    if student_code_tc:
        topic_doc_ref, topic_messages, topic_meta = load_topic_coach_state(
            topic_db, student_code_tc
        )

    def _infer_topic_qcount(messages: List[Dict[str, Any]]) -> int:
        count = 0
        assistant_seen = False
        for msg in messages:
            if not isinstance(msg, dict):
                continue
            role = msg.get("role")
            if role == "assistant":
                assistant_seen = True
            elif role == "user" and assistant_seen:
                count = min(6, count + 1)
        return count

    loaded_qcount_raw = topic_meta.get("qcount") if isinstance(topic_meta, dict) else None
    try:
        loaded_qcount = int(loaded_qcount_raw) if loaded_qcount_raw is not None else None
    except Exception:
        loaded_qcount = None
    if loaded_qcount is None:
        loaded_qcount = _infer_topic_qcount(topic_messages)
    loaded_qcount = max(0, loaded_qcount)

    loaded_finalized = (
        bool(topic_meta.get("finalized")) if isinstance(topic_meta, dict) else False
    )

    # ---------- Widget keys (make them UNIQUE across app) ----------
    POST_TOAST_FLAG    = "_cchat_show_post_toast"
    KEY_LEVEL_SLIDER   = "cchat_w_level"
    KEY_FORCE_DE_TOG   = "cchat_w_force_de"
    KEY_MAX_WORDS_NUM  = "cchat_w_max_words"
    KEY_NEWCHAT_BTN    = "cchat_w_btn_new_bottom"
    KEY_CHAT_INPUT     = "cchat_w_chat_input"
    # Grammar keys: one set shared across the grammar widgets
    KEY_GRAM_TEXT      = "cchat_w_gram_text"
    KEY_GRAM_LEVEL     = "cchat_w_gram_level"
    KEY_GRAM_ASK_BTN   = "cchat_w_gram_go"
    KEY_CONN_MODE      = "cchat_w_conn_mode"
    KEY_CONN_TEXT      = "cchat_w_conn_text"
    KEY_CONN_SCENARIO  = "cchat_w_conn_scenario"
    KEY_CONN_SUGGEST   = "cchat_w_conn_suggest"
    KEY_CONN_SESSION   = "cchat_w_conn_session"
    KEY_CONN_COACH     = "cchat_w_conn_coach"
    KEY_CONN_RESPONSE  = "cchat_w_conn_response"
    KEY_CONN_CLEAR     = "cchat_w_conn_clear"
    # Also make Regen button unique
    KEY_REGEN_BTN      = "cchat_w_btn_regen_v2"

    level_options = ["A1", "A2", "B1", "B2"]
    ensure_student_level()
    roster_level = _safe_upper(st.session_state.get("student_level"), "")
    match = re.search("|".join(level_options), roster_level) if roster_level else None
    default_level = match.group(0) if match else "A2"

    active_level = sync_level_state(
        st,
        student_code=student_code_tc,
        default_level=default_level,
        level_options=level_options,
        slider_key=KEY_LEVEL_SLIDER,
        grammar_key=KEY_GRAM_LEVEL,
    )

    focus_meta = topic_meta.get("focus_tips") if isinstance(topic_meta, dict) else None
    if isinstance(focus_meta, list):
        initial_focus = [
            str(item).strip()
            for item in focus_meta
            if str(item).strip() and str(item).strip().lower() not in {"none", "nan"}
        ]
    else:
        initial_focus = None

    (
        chat_data_key,
        qcount_data_key,
        finalized_data_key,
        focus_data_key,
    ) = _initialise_topic_coach_session_state(
        st.session_state,
        student_code=student_code_tc,
        level=active_level,
        messages=topic_messages,
        qcount=loaded_qcount,
        finalized=loaded_finalized,
        focus_tips=initial_focus,
    )

    def _save_topic_coach_transcript(
        focus_override: Optional[Iterable[str]] = None,
    ) -> None:
        if not student_code_tc:
            return
        doc_ref = topic_doc_ref or get_topic_coach_doc(topic_db, student_code_tc)
        if doc_ref is None:
            return
        focus_payload = list(focus_override or st.session_state.get(focus_data_key, []) or [])
        persist_topic_coach_state(
            doc_ref,
            messages=list(st.session_state.get(chat_data_key, [])),
            qcount=st.session_state.get(qcount_data_key, 0),
            finalized=st.session_state.get(finalized_data_key, False),
            focus_tips=focus_payload,
        )

    # ---------- Subtabs ----------
    tab_labels = ["🧑‍🏫 Topic Coach", "🛠️ Grammar", "📝 Exams"]
    base_tab_labels = tab_labels[:]
    focus_tab = st.session_state.get("_chat_focus_tab")
    if focus_tab not in base_tab_labels:
        focus_tab = base_tab_labels[0]
        st.session_state["_chat_focus_tab"] = focus_tab

    st.markdown(
        """
        <style>
          .chat-grammar-tabs [data-baseweb="tab-list"]{
            flex-wrap:wrap;
            row-gap:0.35rem;
            column-gap:0.35rem;
            justify-content:center;
          }
          .chat-grammar-tabs [data-baseweb="tab"]{
            flex:0 0 auto;
            border-radius:999px !important;
            padding:4px 14px !important;
            white-space:nowrap;
          }
          @media (max-width: 640px){
            .chat-tab-selector{ display:block; margin-bottom:0.5rem; }
            .chat-tab-selector label{ font-weight:700; color:#0f172a; }
          }
          @media (min-width: 641px){
            .chat-tab-selector{ display:none; }
          }
        </style>
        """,
        unsafe_allow_html=True,
    )

    selector_key = "chat_tab_selector"

    def _sync_chat_tab_focus() -> None:
        chosen = st.session_state.get(selector_key)
        if chosen and chosen in base_tab_labels:
            if st.session_state.get("_chat_focus_tab") != chosen:
                st.session_state["_chat_focus_tab"] = chosen
                st.session_state["need_rerun"] = True

    with st.container():
        st.markdown('<div class="chat-tab-selector">', unsafe_allow_html=True)
        st.selectbox(
            "Select a tool",
            base_tab_labels,
            index=base_tab_labels.index(focus_tab),
            key=selector_key,
            on_change=_sync_chat_tab_focus,
            help="Use this menu on phones to switch between Chat • Grammar • Exams tools.",
        )
        st.markdown("</div>", unsafe_allow_html=True)

    if st.session_state.get(selector_key) not in base_tab_labels:
        st.session_state[selector_key] = focus_tab

    if focus_tab in tab_labels:
        idx = tab_labels.index(focus_tab)
        tab_labels = tab_labels[idx:] + tab_labels[:idx]

    st.markdown('<div class="chat-grammar-tabs">', unsafe_allow_html=True)
    tab_contexts = st.tabs(tab_labels)
    st.markdown('</div>', unsafe_allow_html=True)
    tab_lookup = dict(zip(tab_labels, tab_contexts))
    tab_tc = tab_lookup["🧑‍🏫 Topic Coach"]
    tab_gram = tab_lookup["🛠️ Grammar"]
    tab_exam = tab_lookup["📝 Exams"]

    # ===================== Topic Coach (intro, feedback, finalize) =====================
    with tab_tc:
        if st.session_state.pop(POST_TOAST_FLAG, False):
            toast_ok("Answer sent!")

        st.info(
            "Run a 6-question speaking session with Herr Felix. You'll get corrections,"
            " ideas, and a final summary plus ~60-word presentation when you finish."
        )

        focus_tips_display = []
        for tip in st.session_state.get(focus_data_key, []) or []:
            text_tip = str(tip).strip()
            if text_tip and text_tip.lower() not in {"none", "nan"}:
                focus_tips_display.append(text_tip)
        if focus_tips_display:
            st.markdown("#### 🎯 Focus Tips from your last session")
            st.markdown("\n".join(f"- {tip}" for tip in focus_tips_display))
            st.caption("Start with these corrections in mind before you answer the first question.")

        # Recorder reminder banner + button
        st.markdown(
            f"""
            <div class="rec-banner">
              🎙️ You can also <a class="btn secondary" href="{RECORDER_URL}" target="_blank" rel="noopener">Open Recorder</a>
              to record your answers while you practice.
            </div>
            """,
            unsafe_allow_html=True
        )

        # Controls
        colA, colB, colC = st.columns([1,1,1.2])
        with colA:
            cur_level = st.session_state.get(KEY_LEVEL_SLIDER, default_level)
            if cur_level not in level_options:
                cur_level = default_level
            level = st.select_slider("Level (CEFR)", level_options, value=cur_level, key=KEY_LEVEL_SLIDER)
        with colB:
            force_de = st.toggle("Force German replies 🇩🇪", key=KEY_FORCE_DE_TOG, value=st.session_state.get(KEY_FORCE_DE_TOG, False))
        with colC:
            max_words = st.number_input(
                "Max words per reply",
                min_value=60, max_value=400,
                value=int(st.session_state.get(KEY_MAX_WORDS_NUM, 140)),
                step=10, key=KEY_MAX_WORDS_NUM
            )

        # Progress
        q_done = int(st.session_state[qcount_data_key] or 0)
        st.markdown(f"**Progress:** ✅ {q_done}/6 answered")
        st.progress(q_done / 6.0)

        if st.session_state[finalized_data_key]:
            st.success("🎉 Session complete — summary & ~60-word presentation generated. You can regenerate if you like.")
            if st.button("🔁 Regenerate presentation", key=KEY_REGEN_BTN):
                convo = [{"role": "system", "content": "You are Herr Felix. FINALIZE NOW: The student has answered 6 questions. Do not ask more questions. Output two parts: 1) An English summary (strengths, mistakes, improvements). 2) A ~60-word presentation using their own words (add a few if needed). Keep it clear and usable for class. No extra chit-chat."}]
                for m in st.session_state[chat_data_key]:
                    convo.append({"role": m["role"], "content": m["content"]})
                placeholder = st.empty()
                placeholder.markdown("<div class='bubble-a'><div class='typing'><span></span><span></span><span></span></div></div>", unsafe_allow_html=True)
                time.sleep(random.uniform(0.8, 1.2))
                try:
                    resp = client.chat.completions.create(
                        model="gpt-4o-mini",
                        messages=convo,
                        temperature=0.2,
                        max_tokens=800,
                    )
                    reply_raw = (resp.choices[0].message.content or "").strip()
                except Exception as e:
                    reply_raw = f"(Error) {e}"
                placeholder.empty()
                st.session_state[chat_data_key].append({
                    "role": "assistant",
                    "content": reply_raw,
                    "ts": datetime.now(UTC).isoformat()
                })
                _save_topic_coach_transcript()
                st.rerun()

        st.divider()

        # History: older collapsed; last 3 messages shown old→new
        history = st.session_state[chat_data_key] or []
        older = history[:-3] if len(history) > 3 else []
        latest = history[-3:] if len(history) > 3 else history

        typing_notice_placeholder: Optional[Any] = None

        if older:
            with st.expander(f"Show earlier ({len(older)})"):
                for m in older:
                    if m["role"] == "user":
                        st.markdown(
                            f"<div class='bubble-wrap'><div class='lbl-u'>{student_label_html}</div></div>",
                            unsafe_allow_html=True,
                        )
                        st.markdown(f"<div class='bubble-u'>{m['content']}</div>", unsafe_allow_html=True)
                    else:
                        st.markdown(f"<div class='bubble-wrap'><div class='lbl-a'>Herr Felix</div></div>", unsafe_allow_html=True)
                        st.markdown(f"<div class='bubble-a'>{m['content']}</div>", unsafe_allow_html=True)

        for m in latest:
            if m["role"] == "user":
                st.markdown(
                    f"<div class='bubble-wrap'><div class='lbl-u'>{student_label_html}</div></div>",
                    unsafe_allow_html=True,
                )
                st.markdown(f"<div class='bubble-u'>{m['content']}</div>", unsafe_allow_html=True)
            else:
                st.markdown(f"<div class='bubble-wrap'><div class='lbl-a'>Herr Felix</div></div>", unsafe_allow_html=True)
                st.markdown(f"<div class='bubble-a'>{m['content']}</div>", unsafe_allow_html=True)

        typing_notice_placeholder = st.empty()


        # ---- coaching system prompt (intro + feedback + expand + keywords) ----
        correction_lang = "English"
        system_text = (
            f"You are Herr Felix, a supportive and innovative German teacher.\n"
            f"GOALS:\n"
            f"- Run a 6-question speaking practice session guided by exactly 3 keywords.\n"
            f"- After each student answer: (1) correct errors, (2) explain briefly in {correction_lang}, "
            f"(3) add richer ideas/templates/examples to help them speak more, then (4) ask ONE next German question.\n"
            f"- Keep questions level-appropriate, short, and creative; never ask >2 questions per keyword.\n\n"
            f"FIRST REPLY (MANDATORY):\n"
            f"- Your FIRST paragraph must be in English: congratulate the topic, explain how the session works (6 Qs, feedback each time), "
            f"invite questions/translation requests, and set expectations.\n"
            f"- Then show a **KEYWORDS** line with exactly 3 bold chips (e.g., **Supermarkt**, **Preise**, **Einkaufszettel**).\n"
            f"- After the English intro, ask exactly ONE short German question to begin.\n"
            f"- Never start your first reply with German.\n\n"
            f"- The overall 6 questions should train student ability to talk in the past,present and future and also good use of connectors\n\n"
            f"- When student input looks like a letter question, stop the chat and refer them to go to ideas geenrator at the schreiben trainer as you are only trained to coach on sprechen presentation"
            f"FIRST TOPIC BEHAVIOR:\n"
            f"- If the student's FIRST message looks like a topic (1–3 words, e.g., 'Einkaufen', 'Reisen', 'Umwelt'), "
            f"start with a short English topic intro: what it covers, typical situations, 3–5 helpful ideas, "
            f"and a micro-vocabulary set (3–6 important words with brief glosses). Then begin Question 1 in German.\n\n"
            f"FEEDBACK AFTER EACH ANSWER:\n"
            f"- 1) 'Corrections': fix errors; give 1–2 short notes in {correction_lang}; show a corrected German sentence.\n"
            f"- 2) 'Idea boost': give 1–2 extra angles students can add; include a tiny template + one filled example in German.\n"
            f"- 3) 'Next question' (German only): exactly one question.\n\n"
            f"FINISHING:\n"
            f"- Stop after 6 questions (≤2 per keyword). Then output:\n"
            f"  (A) English summary: strengths, mistakes, what to improve.\n"
            f"  (B) A ~60-word class presentation using the student's own words (add a few if needed).\n\n"
            f"KEYWORDS & STYLE:\n"
            f"- Choose exactly 3 useful keywords for the topic. At the top of each reply, show a **KEYWORDS** header "
            f"with three bold chips (e.g., **Supermarkt**, **Preise**, **Einkaufszettel**).\n"
            f"- Short paragraphs, simple bullets. Keep each reply under {max_words} words.\n"
        )
        system_text += f" CEFR level: {level}."
        if force_de:
            system_text += " After your first English intro, ask questions in German; explanations/feedback stay in English."

        # ---- build conversation & enforce first-reply English intro ----
        convo = [{"role": "system", "content": system_text}]
        for m in st.session_state[chat_data_key]:
            convo.append({"role": m["role"], "content": m["content"]})

        no_assistant_yet = not any(m["role"] == "assistant" for m in st.session_state[chat_data_key])
        if no_assistant_yet:
            convo.append({
                "role": "system",
                "content": (
                    "FIRST_REPLY_MODE: Start with a short English intro (overview, session plan, encouragement). "
                    "Then present a **KEYWORDS** line with exactly 3 bold chips. "
                    "Finally ask ONE short German question to begin. "
                    "Do not skip the English intro even if the user's first message is long or detailed."
                )
            })
        else:
            convo.append({
                "role": "system",
                "content": (
                    "FEEDBACK_MODE: For the user's last message, first give 'Corrections' with a brief explanation in English "
                    "and a corrected German sentence; then 'Idea boost' with a tiny template and one filled German example; "
                    "THEN ask exactly ONE next German question (unless already finished)."
                )
            })


        # ---- sticky input (recorder reminder + new chat + input) ----
        with st.container():
            st.markdown("<div class='sticky-input'>", unsafe_allow_html=True)
            st.markdown(
                f"<div class='meta'>Tip: You can also <a href='{RECORDER_URL}' target='_blank' rel='noopener'>record your answer</a> while you practice.</div>",
                unsafe_allow_html=True
            )
            col_left, col_right = st.columns([1, 10])
            with col_left:
                if st.button("🧹 New chat", key=KEY_NEWCHAT_BTN, width="stretch"):
                    st.session_state[chat_data_key] = []
                    st.session_state[qcount_data_key] = 0
                    st.session_state[finalized_data_key] = False
                    _save_topic_coach_transcript()
                    st.toast("Cleared")
                    st.rerun()
            with col_right:
                user_msg = st.chat_input(
                    "Hallo! 👋 What would you like to talk about? Type here so we can chat",
                    key=KEY_CHAT_INPUT
                )
                render_umlaut_pad(KEY_CHAT_INPUT, context="chat_tab", disabled=True)
            st.markdown("</div>", unsafe_allow_html=True)

        # ---- SEND ----
        if user_msg:
            # Store user message
            st.session_state[chat_data_key].append({
                "role": "user", "content": user_msg, "ts": datetime.now(UTC).isoformat()
            })

            # Count rule: don't count first topic; only increment if an assistant turn exists before this user turn
            has_assistant_turn = any(m["role"] == "assistant" for m in st.session_state[chat_data_key][:-1])
            if has_assistant_turn and not st.session_state[finalized_data_key]:
                st.session_state[qcount_data_key] = min(6, int(st.session_state[qcount_data_key] or 0) + 1)

            _save_topic_coach_transcript()
            st.session_state[POST_TOAST_FLAG] = True

            if typing_notice_placeholder is not None:
                typing_notice_placeholder.markdown(
                    HERR_FELIX_TYPING_HTML,
                    unsafe_allow_html=True,
                )

            # Build conversation
            convo = [{"role": "system", "content": system_text}]
            for m in st.session_state[chat_data_key]:
                convo.append({"role": m["role"], "content": m["content"]})

            # Detect first 'topic-like' input (1–3 words, no punctuation) → INTRO_MODE
            def _looks_like_topic(s: str) -> bool:
                s = (s or "").strip()
                if not s: return False
                if any(ch in s for ch in ".,;:!?/\\()[]{}\"'`~@#$%^&*_+=|<>"):
                    return False
                words = s.split()
                return 1 <= len(words) <= 3

            no_assistant_yet = not any(m["role"] == "assistant" for m in st.session_state[chat_data_key])
            if no_assistant_yet and _looks_like_topic(user_msg):
                convo.append({
                    "role": "system",
                    "content": (
                        "INTRO_MODE: The first user input is a topic. Start with a short English topic overview, "
                        "3–5 concrete ideas, and a 3–6 word mini-vocabulary list (German + tiny gloss). "
                        "Then ask the first question in German (ONE question only)."
                    )
                })

            # Finalization at 6 questions
            finalize_now = (int(st.session_state[qcount_data_key]) >= 6) and (not st.session_state[finalized_data_key])
            if finalize_now:
                convo.append({
                    "role": "system",
                    "content": (
                        "FINALIZE NOW: The student has answered 6 questions. "
                        "Do not ask more questions. Output two parts only:\n"
                        "1) English summary (strengths, mistakes, what to improve).\n"
                        "2) A ~60-word presentation using their own words (add a few if needed) for class speaking.\n"
                        "Keep it clear and usable. No extra chit-chat."
                    )
                })
            else:
                # Always ensure corrections + idea boost + next Q on normal turns
                convo.append({
                    "role": "system",
                    "content": (
                        "FEEDBACK_MODE: For the user's last message, first give 'Corrections' with a brief explanation in English "
                        "and a corrected German sentence; then 'Idea boost' with a tiny template and one filled German example; "
                        "THEN ask exactly ONE next German question (unless already finished)."
                    )
                })

            # Typing pulse
            placeholder = st.empty()
            placeholder.markdown(
                "<div class='bubble-a'><div class='typing'><span></span><span></span><span></span></div></div>",
                unsafe_allow_html=True,
            )
            time.sleep(random.uniform(0.8, 1.2))

            # Call model
            try:
                resp = client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=convo,
                    temperature=0.2,
                    max_tokens=800 if finalize_now else 600,
                )
                reply_raw = (resp.choices[0].message.content or "").strip()
            except Exception as e:
                reply_raw = f"(Error) {e}"

            placeholder.empty()

            if typing_notice_placeholder is not None:
                typing_notice_placeholder.empty()

            # Convert "Keywords:" line → bold chips (skip for final message usually)
            chips_html = ""
            if not finalize_now:
                kw_match = re.search(r"(?:^|\n)\s*(?:\*\*?)?(KEYWORDS?|Keywords?)\b(?:\*\*?)?\s*:\s*(.+)", reply_raw)
                if kw_match:
                    raw = re.split(r"[•,|/]", kw_match.group(2))
                    kws = [k.strip(" .*-_") for k in raw if k.strip()]
                    if kws:
                        chips_html = "<div class='kw-title'>**KEYWORDS**</div>" + "".join(
                            f"<span class='kw-chip'><b>{k}</b></span>" for k in kws[:3]
                        )
                        reply_raw = re.sub(r"(?:^|\n)\s*(?:\*\*?)?(KEYWORDS?|Keywords?)\b(?:\*\*?)?\s*:.*", "", reply_raw, count=1)

            # Store assistant message
            bubble = (chips_html + ("<div style='height:6px'></div>" if chips_html else "")) + reply_raw
            st.session_state[chat_data_key].append({
                "role": "assistant", "content": bubble, "ts": datetime.now(UTC).isoformat()
            })

            # If final, mark and try to render presentation card
            if finalize_now:
                st.session_state[finalized_data_key] = True

                # Extract a ~60-word presentation paragraph
                def _extract_presentation(text: str) -> str:
                    paras = [p.strip() for p in re.split(r"\n{2,}|\r{2,}", text) if p.strip()]
                    best = ""
                    best_len = 0
                    for p in paras:
                        wc = len(re.findall(r"\b\w+\b", p))
                        if 45 <= wc <= 90 and wc > best_len:
                            best, best_len = p, wc
                    if not best:
                        best = paras[-1] if paras else text
                    return best

                pres = _extract_presentation(reply_raw)
                wc = len(re.findall(r"\b\w+\b", pres))
                html_id = f"pres_{uuid4().hex}"
                components.html(f"""
                    <div class="pres-card" id="{html_id}">
                      <div class="pres-head">
                        <div><b>🎤 60-word class presentation</b> <span class="meta">(~{wc} words)</span></div>
                        <button class="pres-copy" onclick="
                          const txt = document.querySelector('#{html_id} .pres-body').innerText;
                          navigator.clipboard.writeText(txt);
                          this.innerText='Copied!';
                          setTimeout(()=>this.innerText='Copy',1200);
                        ">Copy</button>
                      </div>
                      <div class="pres-body" style="margin-top:8px; white-space:pre-wrap;">{html.escape(pres)}</div>
                    </div>
                """, height=170)

                new_focus = _extract_focus_tips_from_history(
                    st.session_state.get(chat_data_key, [])
                )
                st.session_state[focus_data_key] = new_focus

            _save_topic_coach_transcript(
                focus_override=st.session_state.get(focus_data_key, [])
            )
            st.rerun()

    # ===================== Grammar (simple, one-box) =====================
    with tab_gram:
        level_options = ["A1", "A2", "B1", "B2"]
        default_gram = st.session_state.get("_cchat_last_profile_level") or level_options[0]
        if default_gram not in level_options:
            default_gram = level_options[0]
        cur_level_g = st.session_state.get(KEY_GRAM_LEVEL, default_gram)
        if cur_level_g not in level_options:
            cur_level_g = default_gram
        st.session_state.setdefault(KEY_GRAM_LEVEL, cur_level_g)

        tab_quick, tab_connectors = st.tabs(["Quick Grammar Help", "Connector Trainer"])

        with tab_quick:
            st.info(
                "Paste a sentence or grammar question here to get quick corrections,"
                " short explanations in English, and German example sentences for your level."
            )
            gcol1, gcol2 = st.columns([3, 1])
            gram_typing_notice = None
            with gcol1:
                gram_typing_notice = st.empty()
                gram_q = st.text_area(
                    "Type your grammar question or paste text",
                    height=160,
                    key=KEY_GRAM_TEXT,
                    placeholder="z.B. Ist es 'wegen dem' oder 'wegen des'? Oder: Ich bin gestern in den Park gegangen…",
                )
            with gcol2:
                gram_level = st.select_slider(
                    "Level",
                    level_options,
                    value=st.session_state.get(KEY_GRAM_LEVEL, cur_level_g),
                    key=KEY_GRAM_LEVEL,
                )
                ask = st.button("Ask", type="primary", width="stretch", key=KEY_GRAM_ASK_BTN)

            if ask and (gram_q or "").strip():
                sys = (
                    "You are a German grammar helper. "
                    "All EXPLANATIONS must be in English ONLY. "
                    "The CORRECTED TEXT and EXAMPLE SENTENCES must be in German ONLY. "
                    "Match the CEFR level and be concise. "
                    "If text was pasted, first give a short corrected German version of the user's text, "
                    "then provide exactly 3 concise English bullet points explaining the key grammar points and corrections, "
                    "and finally give 1–2 short German example sentences that illustrate the rule. "
                    "If it's only a question (no text to correct), give the English explanation and German examples only. "
                    "Keep the whole answer compact and classroom-friendly."
                )
                if gram_typing_notice is not None:
                    gram_typing_notice.markdown(
                        HERR_FELIX_TYPING_HTML,
                        unsafe_allow_html=True,
                    )

                placeholder = st.empty()
                placeholder.markdown(
                    "<div class='bubble-a'><div class='typing'><span></span><span></span><span></span></div></div>",
                    unsafe_allow_html=True,
                )
                time.sleep(random.uniform(0.8, 1.2))
                if topic_db is None:
                    logging.debug("Grammar Firestore logging skipped: no client available")
                else:
                    try:
                        topic_db.collection("falowen_grammar_questions").add(
                            {
                                "student_code": st.session_state.get("student_code"),
                                "level": gram_level,
                                "question": gram_q,
                                "created_at": firestore.SERVER_TIMESTAMP,
                            }
                        )
                    except Exception:
                        logging.warning("Failed to log grammar question", exc_info=True)
                try:
                    resp = client.chat.completions.create(
                        model="gpt-4o-mini",
                        messages=[
                            {"role": "system", "content": sys + f" CEFR level: {gram_level}."},
                            {"role": "user", "content": gram_q},
                        ],
                        temperature=0.1,
                        max_tokens=700,
                    )
                    out = (resp.choices[0].message.content or "").strip()
                except Exception as e:
                    out = f"(Error) {e}"
                placeholder.empty()
                if gram_typing_notice is not None:
                    gram_typing_notice.empty()
                st.markdown(
                    "<div class='bubble-wrap'><div class='lbl-a'>Herr Felix</div></div>",
                    unsafe_allow_html=True,
                )
                st.markdown(f"<div class='bubble-a'>{out}</div>", unsafe_allow_html=True)

        with tab_connectors:
            gram_level = st.session_state.get(KEY_GRAM_LEVEL, cur_level_g)
            st.info(
                "Build smoother sentences with German connectors. Choose your own or have Herr Felix suggest a set,"
                " then get coaching with examples matched to your level."
            )

            connector_mode = st.radio(
                "How should we choose connectors?",
                ("I'll choose connectors", "Suggest for me"),
                key=KEY_CONN_MODE,
                horizontal=True,
            )

            connectors_text = ""
            suggestion_placeholder = st.empty()
            if connector_mode == "I'll choose connectors":
                connectors_text = st.text_area(
                    "Connectors or topics to focus on",
                    key=KEY_CONN_TEXT,
                    height=120,
                    placeholder="z.B. weil, obwohl, trotzdem – oder Themen wie: Mein Wochenende, Arbeit vs. Freizeit",
                )
                suggestion_placeholder.empty()
            else:
                st.caption(
                    "Let Herr Felix suggest 3–5 connectors for your current CEFR level "
                    f"({gram_level})."
                )
                if st.button("Suggest connectors", key=KEY_CONN_SUGGEST):
                    suggestion_placeholder.markdown(
                        "<div class='bubble-a'><div class='typing'><span></span><span></span><span></span></div></div>",
                        unsafe_allow_html=True,
                    )
                    try:
                        resp = client.chat.completions.create(
                            model="gpt-4o-mini",
                            messages=[
                                {
                                    "role": "system",
                                    "content": (
                                        "You are a German writing coach. Suggest 3-5 connectors appropriate for the "
                                        "given CEFR level. Provide them as a simple numbered list with each connector "
                                        "in German and a short English hint about when to use it."
                                    ),
                                },
                                {
                                    "role": "user",
                                    "content": f"CEFR level: {gram_level}. Suggest connectors for practice.",
                                },
                            ],
                            temperature=0.3,
                            max_tokens=300,
                        )
                        suggestions = (resp.choices[0].message.content or "").strip()
                    except Exception as exc:
                        suggestions = f"(Error) {exc}"
                    suggestion_placeholder.empty()
                    st.session_state[KEY_CONN_SESSION] = suggestions
                suggestions = st.session_state.get(KEY_CONN_SESSION, "")
                if suggestions:
                    suggestion_placeholder.markdown(
                        (
                            "<div class='bubble-wrap'><div class='lbl-a'>Herr Felix</div></div>"
                            f"<div class='bubble-a'>{suggestions}</div>"
                        ),
                        unsafe_allow_html=True,
                    )
                else:
                    suggestion_placeholder.empty()
                connectors_text = suggestions

            scenario_text = st.text_area(
                "Sentence or scenario to practice",
                key=KEY_CONN_SCENARIO,
                height=150,
                placeholder="Beschreibe kurz, was du sagen möchtest. z.B. Ich erzähle über meinen letzten Urlaub und möchte Kontraste zeigen.",
            )

            coach_response_placeholder = st.empty()
            existing_coach_reply = st.session_state.get(KEY_CONN_RESPONSE, "")
            if existing_coach_reply:
                coach_response_placeholder.markdown(
                    (
                        "<div class='bubble-wrap'><div class='lbl-a'>Herr Felix</div></div>"
                        f"<div class='bubble-a'>{existing_coach_reply}</div>"
                    ),
                    unsafe_allow_html=True,
                )

            col_coach, col_clear = st.columns([1, 1])
            coach_btn = col_coach.button("Coach me", type="primary", key=KEY_CONN_COACH)
            clear_btn = col_clear.button("Clear session", type="secondary", key=KEY_CONN_CLEAR)

            if clear_btn:
                st.session_state.pop(KEY_CONN_SESSION, None)
                st.session_state.pop(KEY_CONN_TEXT, None)
                st.session_state.pop(KEY_CONN_SCENARIO, None)
                st.session_state.pop(KEY_CONN_RESPONSE, None)
                st.session_state.pop(KEY_CONN_MODE, None)
                suggestion_placeholder.empty()
                coach_response_placeholder.empty()
                rerun_fn = getattr(st, "experimental_rerun", None) or getattr(st, "rerun", None)
                if callable(rerun_fn):
                    rerun_fn()

            if coach_btn:
                gram_level = st.session_state.get(KEY_GRAM_LEVEL, cur_level_g)
                connector_payload = connectors_text.strip()
                scenario_payload = scenario_text.strip()
                if not connector_payload and connector_mode == "I'll choose connectors":
                    st.warning("Bitte gib mindestens einen Connector oder ein Thema an – oder nutze die Vorschlagsfunktion.")
                else:
                    coach_response_placeholder.markdown(
                        "<div class='bubble-a'><div class='typing'><span></span><span></span><span></span></div></div>",
                        unsafe_allow_html=True,
                    )
                    time.sleep(random.uniform(0.8, 1.2))
                    sys_msg = (
                        "You are Herr Felix, a German connector coach. Review any provided connectors (or suggest a short "
                        "set if none were given), explain how to use them in English, and give short German example "
                        "sentences that match the student's CEFR level. Keep the tone encouraging and classroom-friendly."
                    )
                    user_parts = [f"CEFR level: {gram_level}"]
                    if connector_payload:
                        user_parts.append(f"Connectors to cover: {connector_payload}")
                    else:
                        user_parts.append("No connectors provided – please suggest a helpful set.")
                    if scenario_payload:
                        user_parts.append(f"Practice scenario: {scenario_payload}")
                    user_msg = "\n".join(user_parts)

                    try:
                        resp = client.chat.completions.create(
                            model="gpt-4o-mini",
                            messages=[
                                {"role": "system", "content": sys_msg},
                                {"role": "user", "content": user_msg},
                            ],
                            temperature=0.4,
                            max_tokens=700,
                        )
                        coaching_reply = (resp.choices[0].message.content or "").strip()
                    except Exception as exc:
                        coaching_reply = f"(Error) {exc}"
                    st.session_state[KEY_CONN_RESPONSE] = coaching_reply
                    coach_response_placeholder.markdown(
                        (
                            "<div class='bubble-wrap'><div class='lbl-a'>Herr Felix</div></div>"
                            f"<div class='bubble-a'>{coaching_reply}</div>"
                        ),
                        unsafe_allow_html=True,
                    )

                    if topic_db is None:
                        logging.debug("Connector Firestore logging skipped: no client available")
                    else:
                        try:
                            topic_db.collection("falowen_connector_sessions").add(
                                {
                                    "student_code": st.session_state.get("student_code"),
                                    "level": gram_level,
                                    "mode": connector_mode,
                                    "connectors": connector_payload,
                                    "scenario": scenario_payload,
                                    "response": coaching_reply,
                                    "suggestions": st.session_state.get(KEY_CONN_SESSION, ""),
                                    "created_at": firestore.SERVER_TIMESTAMP,
                                }
                            )
                        except Exception:
                            logging.warning("Failed to log connector session", exc_info=True)


    # ===================== Exams (Speaking • Lesen • Hören) =====================
    with tab_exam:
        # Level-aware Goethe links (Lesen & Hören)
        lesen_links = {
            "A1": [("Goethe A1 Lesen (Lesen & Hören page)", "https://www.goethe.de/ins/mm/en/spr/prf/gzsd1/ueb.html")],
            "A2": [("Goethe A2 Lesen (Lesen & Hören page)", "https://www.goethe.de/ins/mm/en/spr/prf/gzsd2/ueb.html")],
            "B1": [("Goethe B1 Lesen (Lesen & Hören page)", "https://www.goethe.de/ins/mm/en/spr/prf/gzb1/ueb.html")],
            "B2": [("Goethe B2 Lesen (Lesen & Hören page)", "https://www.goethe.de/ins/mm/en/spr/prf/gzb2/ue9.html")],
            "C1": [("Goethe C1 Lesen (Lesen & Hören page)", "https://www.goethe.de/ins/be/en/spr/prf/gzc1/u24.html")],
        }
        hoeren_links = {
            "A1": [("Goethe A1 Hören (Lesen & Hören page)", "https://www.goethe.de/ins/mm/en/spr/prf/gzsd1/ueb.html")],
            "A2": [("Goethe A2 Hören (Lesen & Hören page)", "https://www.goethe.de/ins/mm/en/spr/prf/gzsd2/ueb.html")],
            "B1": [("Goethe B1 Hören (Lesen & Hören page)", "https://www.goethe.de/ins/mm/en/spr/prf/gzb1/ueb.html")],
            "B2": [("Goethe B2 Hören (Lesen & Hören page)", "https://www.goethe.de/ins/mm/en/spr/prf/gzb2/ue9.html")],
            "C1": [("Goethe C1 Hören (Lesen & Hören page)", "https://www.goethe.de/ins/be/en/spr/prf/gzc1/u24.html")],
        }

        st.info(
            "Use these subtabs to get exam-ready: Speaking shows your practice attempts and scores,"
            " while Lesen and Hören open official Goethe online exercises for your chosen level."
        )

        level_for_exams = st.session_state.get("exam_lesen_level", default_level)
        if level_for_exams not in lesen_links or level_for_exams not in hoeren_links:
            level_for_exams = st.session_state.get("exam_hoeren_level", level_for_exams)
        if level_for_exams not in lesen_links or level_for_exams not in hoeren_links:
            if default_level in lesen_links and default_level in hoeren_links:
                level_for_exams = default_level
            else:
                common_levels = [lvl for lvl in lesen_links if lvl in hoeren_links]
                if common_levels:
                    level_for_exams = common_levels[0]
                else:
                    level_for_exams = next(
                        iter(lesen_links.keys()),
                        next(iter(hoeren_links.keys()), default_level),
                    )

        sub_speak, sub_lesen, sub_hoeren = st.tabs(["🗣️ Speaking", "📖 Lesen", "🎧 Hören"])

        def _link_buttons(links: Iterable[Tuple[str, str]]) -> None:
            """Render a set of practice links as buttons in the Exams tab."""

            items = list(links or [])
            if not items:
                st.caption("No practice links available for this level yet.")
                return

            for idx, (label, url) in enumerate(items):
                if not label or not url:
                    continue

                st.link_button(label, url, width="stretch")


    with sub_speak:
        st.markdown(
            """
            <div class="panel">
              <b>Speaking practice</b><br>
              • Open the practice page and follow the prompts to record and submit.<br>
              • Keep sentences short and clear — you’ve got this! 💪
            </div>
            """,
            unsafe_allow_html=True,
        )

        st.markdown(
            '<a class="btn" href="https://blog.falowen.app/blog/a1-speaking-exam-guide/" target="_blank" rel="noopener">📘 Read the A1 Speaking Exam Guide</a>',
            unsafe_allow_html=True,
        )

        # Button to your speaking practice page (PRACTICE_URL)
        st.markdown(
            f'<a class="btn secondary" href="{PRACTICE_URL}" target="_blank" rel="noopener">📝 Open Speaking Practice</a>',
            unsafe_allow_html=True,
        )

        st.info(
            "Track your latest speaking scores and audio directly on the speaking practice page. "
            "Once you've submitted new attempts, reopen the practice tool to review your feedback."
        )

        # ---------- Lesen ----------
        with sub_lesen:
            st.markdown(
                """
                <div class="panel">
                  <b>Reading (Lesen)</b><br>
                  • Choose your level to open official Goethe online practice.
                </div>
                """,
                unsafe_allow_html=True,
            )
            lv = st.select_slider(
                "Level",
                options=list(lesen_links.keys()),
                value=level_for_exams,
                key="exam_lesen_level",
            )
            _link_buttons(lesen_links.get(lv, []))

        # ---------- Hören ----------
        with sub_hoeren:
            st.markdown(
                """
                <div class="panel">
                  <b>Listening (Hören)</b><br>
                  • Choose your level to open official Goethe online practice.
                </div>
                """,
                unsafe_allow_html=True,
            )
            lv_h = st.select_slider(
                "Level",
                options=list(hoeren_links.keys()),
                value=level_for_exams,
                key="exam_hoeren_level",
            )
            _link_buttons(hoeren_links.get(lv_h, []))

    st.divider()
    render_app_footer(FOOTER_LINKS)



# =========================================
# Vocab
# =========================================


# ================================
# CONFIG: Sheet for Vocab + Audio
# ================================
def is_correct_answer(user_input: str, answer: str) -> bool:
    """Return True if the user's input matches the expected answer.

    Comparison ignores leading/trailing whitespace and letter casing. In
    addition, leading English articles ("the", "a", "an") are removed and a
    fuzzy similarity check is performed using :class:`difflib.SequenceMatcher`.
    """

    from difflib import SequenceMatcher

    normalized_user = user_input.strip().lower()
    normalized_answer = answer.strip().lower()

    # Fast path for exact matches
    if normalized_user == normalized_answer:
        return True

    def _strip_article(s: str) -> str:
        for article in ("the ", "a ", "an "):
            if s.startswith(article):
                return s[len(article):].lstrip()
        return s

    normalized_user = _strip_article(normalized_user)
    normalized_answer = _strip_article(normalized_answer)

    similarity = SequenceMatcher(None, normalized_user, normalized_answer).ratio()
    return similarity >= 0.85

# ================================
# TAB: Vocab Trainer (locked by Level)
# ================================
if tab == "Schreiben Trainer":
    st.markdown(
        '''
        <div style="
            padding: 8px 12px;
            background: #d63384;
            color: #fff;
            border-radius: 6px;
            text-align: center;
            margin-bottom: 8px;
            font-size: 1.3rem;">
            ✍️ Schreiben Trainer (Writing Practice)
        </div>
        ''',
        unsafe_allow_html=True
    )

    st.info(
        """
        ✍️ **This section is for Writing (Schreiben) only.**
        Practice German letters, emails, and essays for A1–C1 exams—now with automatic level detection.

        Want to practice presentations or focus on Speaking, Reading, or Listening?
        👉 Switch to **Exam Mode & Chat • Grammar • Exams** (tab above)!

        Your writing will be assessed and scored out of 25 marks, just like in the real exam.
        """,
        icon="✉️"
    )

    st.divider()

    # --- Writing stats summary with Firestore ---
    student_code = st.session_state.get("student_code", "demo")
    stats = render_schreiben_stats(student_code)
    student_name = _safe_str(st.session_state.get("student_name"), "Student") or "Student"

    # --- Update session states for new student (preserves drafts, etc) ---
    prev_student_code = st.session_state.get("prev_student_code", None)
    if student_code != prev_student_code:
        stats = stats or get_schreiben_stats(student_code)
        st.session_state[f"{student_code}_last_feedback"] = None
        st.session_state[f"{student_code}_last_user_letter"] = None
        st.session_state[f"{student_code}_delta_compare_feedback"] = None
        st.session_state[f"{student_code}_final_improved_letter"] = ""
        st.session_state[f"{student_code}_awaiting_correction"] = False
        st.session_state[f"{student_code}_improved_letter"] = ""
        st.session_state["prev_student_code"] = student_code

    pending_subtab = st.session_state.pop("schreiben_pending_subtab", None)
    if pending_subtab:
        st.session_state[f"schreiben_sub_tab_{student_code}"] = pending_subtab

    # --- Sub-tabs for the Trainer ---
    st.markdown(
        """
        <style>
        div[role="radiogroup"][aria-label="Choose Mode"],
        div[data-testid="stHorizontalBlock"] [aria-label="Choose Mode"] {
            display: flex;
            flex-wrap: wrap;
            gap: 0.5rem;
        }
        div[role="radiogroup"][aria-label="Choose Mode"] > label,
        div[data-testid="stHorizontalBlock"] [aria-label="Choose Mode"] > label {
            flex: 1 1 180px;
            min-width: 160px;
        }
        @media (max-width: 640px) {
            div[role="radiogroup"][aria-label="Choose Mode"] > label,
            div[data-testid="stHorizontalBlock"] [aria-label="Choose Mode"] > label {
                flex-basis: 100%;
            }
        }
        </style>
        """,
        unsafe_allow_html=True,
    )
    sub_tab = st.radio(
        "Choose Mode",
        [
            "Practice Letters",
            "Mark My Letter",
            "Ideas Generator (Letter Coach)",
            "Vocab Trainer",

        ],
        horizontal=True,
        key=f"schreiben_sub_tab_{student_code}"
    )

        # --- Level picker: Auto-detect from student code (manual override removed) ---
    if student_code:
        detected_level = get_level_from_code(student_code)
        # Only apply detected level when first seeing this student code
        if st.session_state.get("prev_student_code_for_level") != student_code:
            st.session_state["schreiben_level"] = detected_level
            st.session_state["prev_student_code_for_level"] = student_code
    else:
        detected_level = "A1"
        if "schreiben_level" not in st.session_state:
            st.session_state["schreiben_level"] = detected_level

    # Ensure current writing level variable reflects auto-detected one
    schreiben_level = st.session_state.get("schreiben_level", "A1")

    st.markdown(
        f"<span style='color:gray;font-size:0.97em;'>Auto-detected level from your code: <b>{detected_level}</b></span>",
        unsafe_allow_html=True
    )


    st.divider()

    # ----------- PRACTICE LETTERS -----------
    if sub_tab == "Practice Letters":
        try:
            from src.schreiben_prompts_module import get_prompts_for_level
        except Exception:  # pragma: no cover - fallback if module missing
            def get_prompts_for_level(_level):
                return []

        prompts = get_prompts_for_level(schreiben_level)
        if prompts:
            options = [p["Thema"] for p in prompts]
            selected_theme = st.selectbox(
                "Choose a prompt",
                options,
                key=f"practice_prompt_{student_code}",
            )
            st.markdown(
                "**[German Writing Rules](https://drive.google.com/file/d/1o7_ez3WSNgpgxU_nEtp6EO1PXDyi3K3b/view?usp=sharing)**",
            )
            prompt = next((p for p in prompts if p["Thema"] == selected_theme), None)
            if prompt:
                st.markdown(f"### ✉️ {prompt['Thema']}")
                st.markdown("\n".join(f"- {p}" for p in prompt['Punkte']))
        else:
            st.selectbox(
                "Choose a prompt",
                ["(no prompts available)"],
                key=f"practice_prompt_{student_code}",
            )
        st.info(
            "Use \u201cMark My Letter\u201d to submit your response for evaluation and \u201cIdeas Generator (Letter Coach)\u201d when you need inspiration.",
        )

    if sub_tab == "Vocab Trainer":
        render_vocab_trainer_section()

    # ----------- 1. MARK MY LETTER -----------
    if sub_tab == "Mark My Letter":
        daily_so_far = get_schreiben_usage(student_code)
        st.markdown(f"**Daily usage:** {daily_so_far} / {SCHREIBEN_DAILY_LIMIT}")

        st.markdown(
            "**[German Writing Rules](https://drive.google.com/file/d/1o7_ez3WSNgpgxU_nEtp6EO1PXDyi3K3b/view?usp=sharing)**",
        )

        try:
            _ = _wkey
        except NameError:
            import hashlib

            def _wkey(base: str) -> str:
                sc = str(st.session_state.get("student_code", "anon"))
                return f"{base}_{hashlib.md5(f'{base}|{sc}'.encode()).hexdigest()[:8]}"

        draft_key = _wkey("schreiben_letter")
        existing_draft = load_draft_from_db(student_code, draft_key)
        existing_feedback, existing_letter = load_schreiben_feedback(student_code)
        if existing_feedback or existing_letter:
            st.session_state[f"{student_code}_last_feedback"] = existing_feedback
            st.session_state[f"{student_code}_last_user_letter"] = existing_letter

        letter_disabled = daily_so_far >= SCHREIBEN_DAILY_LIMIT
        if letter_disabled:
            st.warning(
                "You've reached today's limit for letter corrections. Please come back"
                " tomorrow after the daily quota resets to submit a new letter. While"
                " you wait, feel free to explore the other tabs for prompts, ideas, or"
                " practice activities."
            )

        user_letter = st.text_area(
            "Paste or type your German letter/essay here.",
            key=draft_key,
            value=existing_draft,
            on_change=lambda: save_now(draft_key, student_code),
            height=400,
            placeholder="Write your German letter here...",
            disabled=letter_disabled,
        )

        render_umlaut_pad(
            draft_key,
            context=f"schreiben_mark_{student_code}",
            disabled=letter_disabled,
        )

        autosave_maybe(student_code, draft_key, user_letter, min_secs=2.0, min_delta=20)

        if st.button("\U0001f4be Save Draft", key=f"save_draft_btn_{student_code}"):
            save_now(draft_key, student_code)
            toast_once("Draft saved!", "✅")
        st.caption("Auto-saves every few seconds or click 'Save Draft' to save now.")

        def clear_feedback_and_start_new():
            for k in [
                "last_feedback",
                "last_user_letter",
                "delta_compare_feedback",
                "improved_letter",
                "final_improved_letter",
            ]:
                st.session_state.pop(f"{student_code}_{k}", None)
            st.session_state[f"{student_code}_awaiting_correction"] = False
            st.session_state.pop(draft_key, None)
            save_now(draft_key, student_code)
            lv, lt, sf, sa = _draft_state_keys(draft_key)
            for key in (lv, lt, sf, sa):
                st.session_state.pop(key, None)
            delete_schreiben_feedback(student_code)
            st.session_state["need_rerun"] = True

        if st.session_state.get(f"{student_code}_last_feedback"):
            st.info(
                "Draft auto-save is paused while feedback is visible. "
                "Clear feedback to resume saving."
            )

        # --- Word count and Goethe exam rules ---
        import re
        def get_level_requirements(level):
            reqs = {
                "A1": {"min": 25, "max": 40, "desc": "A1 formal/informal letters should be 25–40 words. Cover all bullet points."},
                "A2": {"min": 30, "max": 40, "desc": "A2 formal/informal letters should be 30–40 words. Cover all bullet points."},
                "B1": {"min": 80, "max": 150, "desc": "B1 letters/essays should be about 80–150 words, with all points covered and clear structure."},
                "B2": {"min": 150, "max": 250, "desc": "B2 essays are 180–220 words, opinion essays or reports, with good structure and connectors."},
                "C1": {"min": 230, "max": 350, "desc": "C1 essays are 230–250+ words. Use advanced structures and express opinions clearly."}
            }
            return reqs.get(level.upper(), reqs["A1"])

        def count_words(text):
            return len(re.findall(r'\b\w+\b', text))

        if user_letter.strip():
            words = re.findall(r'\b\w+\b', user_letter)
            chars = len(user_letter)
            st.info(f"**Word count:** {len(words)} &nbsp;|&nbsp; **Character count:** {chars}")

            # -- Apply Goethe writing rules here --
            requirements = get_level_requirements(detected_level)  # << USE AUTO-DETECTED LEVEL
            word_count = count_words(user_letter)
            min_wc = requirements["min"]
            max_wc = requirements["max"]

            if detected_level in ("A1", "A2"):
                if word_count < min_wc:
                    st.error(f"⚠️ Your letter is too short for {detected_level} ({word_count} words). {requirements['desc']}")
                    st.stop()
                elif word_count > max_wc:
                    st.warning(f"ℹ️ Your letter is a bit long for {detected_level} ({word_count} words). The exam expects {min_wc}-{max_wc} words.")
            else:
                if word_count < min_wc:
                    st.error(f"⚠️ Your essay is too short for {detected_level} ({word_count} words). {requirements['desc']}")
                    st.stop()
                elif word_count > max_wc + 40 and detected_level in ("B1", "B2"):
                    st.warning(f"ℹ️ Your essay is longer than the usual limit for {detected_level} ({word_count} words). Try to stay within the guidelines.")

        # --------- Reset correction states (do not indent inside above ifs)
        for k, v in [
            ("last_feedback", None),
            ("last_user_letter", None),
            ("delta_compare_feedback", None),
            ("improved_letter", ""),
            ("awaiting_correction", False),
            ("final_improved_letter", "")
        ]:
            session_key = f"{student_code}_{k}"
            if session_key not in st.session_state:
                st.session_state[session_key] = v

        # Namespaced correction state per student (reset on session)
        for k, v in [
            ("last_feedback", None),
            ("last_user_letter", None),
            ("delta_compare_feedback", None),
            ("improved_letter", ""),
            ("awaiting_correction", False),
            ("final_improved_letter", "")
        ]:
            session_key = f"{student_code}_{k}"
            if session_key not in st.session_state:
                st.session_state[session_key] = v

        submit_disabled = (not user_letter.strip()) or (daily_so_far >= SCHREIBEN_DAILY_LIMIT)
        feedback_btn = st.button(
            "Get Feedback",
            type="primary",
            disabled=submit_disabled,
            key=f"feedback_btn_{student_code}"
        )

        if feedback_btn:
            st.session_state[f"{student_code}_awaiting_correction"] = True
            ai_prompt = (
                f"You are Herr Felix, a supportive and innovative German letter writing trainer.\n"
                f"You help students prepare for A1, A2, B1, B2, and C1 German exam letters or essays.\n"
                f"The student has submitted a {schreiben_level} German letter or essay.\n"
                f"Your job is to mark, score, and explain feedback in a kind, step-by-step way.\n"
                f"Always answer in English.\n"
                f"Begin with a warm greeting that uses the student's name ({student_name}) and refer to them by name throughout your feedback.\n"
                f"1. Give a quick summary (one line) of how well the student did overall.\n"
                f"2. Then show a detailed breakdown of strengths and weaknesses in 4 areas:\n"
                f"   Grammar, Vocabulary, Spelling, Structure.\n"
                f"3. For each area, say what was good and what should improve.\n"
                f"4. Highlight every mistake with [wrong]...[/wrong] and every good example with [correct]...[/correct].\n"
                f"5. Give 2-3 improvement tips in bullet points.\n"
                f"6. At the end, give a realistic score out of 25 in the format: Score: X/25.\n"
                f"7. For A1 and A2, be strict about connectors, basic word order, modal verbs, and correct formal/informal greeting.\n"
                f"8. For B1+, mention exam criteria and what examiner wants.\n"
                f"9. Never write a new letter for the student, only mark what they submit.\n"
                f"10. When possible, point out specific lines or examples from their letter in your feedback.\n"
                f"11. When student score is 18 or above then they have passed. When score is less than 18, is a fail and they must try again before submitting to prevent low marks.\n"
                f"12. After completion, remind them to only copy their improved letter without your feedback, go to 'my course' on the app and submit together with their lesen and horen answers. They only share the letter and feedback with their teacher for evaluation only when they preparing for the exams\n"
                
            )

            with st.spinner("🧑‍🏫 Herr Felix is typing..."):
                try:
                    completion = client.chat.completions.create(
                        model="gpt-4o",
                        messages=[
                            {"role": "system", "content": ai_prompt},
                            {"role": "user", "content": user_letter},
                        ],
                        temperature=0.6,
                    )
                    feedback = completion.choices[0].message.content
                    st.session_state[f"{student_code}_last_feedback"] = feedback
                    st.session_state[f"{student_code}_last_user_letter"] = user_letter
                    st.session_state[f"{student_code}_delta_compare_feedback"] = None
                except Exception:
                    st.error("AI feedback failed. Please check your OpenAI setup.")
                    feedback = None

            if feedback:
                st.markdown("[⬇️ Jump to feedback](#feedback-reference)")
                st.session_state[f"{student_code}_awaiting_correction"] = True

                save_schreiben_feedback(student_code, feedback, user_letter)

                # --- Save to Firestore ---
                score_match = re.search(r"Score[: ]+(\d+)", feedback)
                score = int(score_match.group(1)) if score_match else 0
                passed = score >= 18
                save_submission(
                    student_code=student_code,
                    score=score,
                    passed=passed,
                    timestamp=None,  # Not needed
                    level=schreiben_level,
                    letter=user_letter
                )
                update_schreiben_stats(student_code)
                inc_schreiben_usage(student_code)
                save_draft_to_db(student_code, draft_key, "")
                st.session_state.pop(draft_key, None)

        elif (
            st.session_state.get(f"{student_code}_last_feedback")
            and st.session_state.get(f"{student_code}_last_user_letter")
        ):
            st.markdown("[⬇️ Jump to feedback](#feedback-reference)")
            
        # --- Improvement section: Compare, download, WhatsApp ---
        if st.session_state.get(f"{student_code}_last_feedback") and st.session_state.get(f"{student_code}_last_user_letter"):
            st.markdown("---")
            st.markdown('<div id="feedback-reference"></div>', unsafe_allow_html=True)
            st.markdown("#### 📝 Feedback from Herr Felix (Reference)")
            st.markdown(
                highlight_feedback(st.session_state[f"{student_code}_last_feedback"]),
                unsafe_allow_html=True
            )
            clear_feedback_reference = st.button(
                "🗑️ Clear feedback and start a new letter",
                key=f"clear_feedback_{student_code}_reference",
            )
            if clear_feedback_reference:
                clear_feedback_and_start_new()
            st.markdown(
                """
                <div style="background:#e3f7da; border-left:7px solid #44c767;
                color:#295327; padding:1.15em; margin-top:1em; border-radius:10px; font-size:1.09em;">
                    🔁 <b>Try to improve your letter!</b><br>
                    Paste your improved version below and click <b>Compare My Improvement</b>.<br>
                    The AI will highlight what’s better, what’s still not fixed, and give extra tips.<br>
                    <b>You can download or share the improved version & new feedback below.</b>
                </div>
                """, unsafe_allow_html=True
            )
            improved_letter = st.text_area(
                "Your improved version (try to fix the mistakes Herr Felix mentioned):",
                key=f"{student_code}_improved_letter",
                height=400,
                placeholder="Paste your improved letter here..."
            )
            render_umlaut_pad(
                f"{student_code}_improved_letter",
                context=f"schreiben_improve_{student_code}",
            )
            compare_clicked = st.button("Compare My Improvement", key=f"compare_btn_{student_code}")

            if compare_clicked and improved_letter.strip():
                ai_compare_prompt = (
                    "You are Herr Felix, a supportive German writing coach. "
                    "A student first submitted this letter:\n\n"
                    f"{st.session_state[f'{student_code}_last_user_letter']}\n\n"
                    "Your feedback was:\n"
                    f"{st.session_state[f'{student_code}_last_feedback']}\n\n"
                    "Now the student has submitted an improved version below.\n"
                    f"The student's name is {student_name}. Begin by greeting {student_name} and continue to use their name warmly throughout your response.\n"
                    "Compare both versions and:\n"
                    "- Tell the student exactly what they improved, and which mistakes were fixed.\n"
                    "- Point out if there are still errors left, with new tips for further improvement.\n"
                    "- Encourage the student. If the improvement is significant, say so.\n"
                    "1. If student dont improve after the third try, end the chat politely and tell the student to try again tomorrow. Dont continue to give the feedback after third try.\n"
                    "2. Always explain your feeback in English for them to understand. You can still highlight their german phrases. But your correction should be english\n"
                    "3. For A1 and A2 students, make sure a sentence is not more than 7 words."
                    "4. For A1 and A2 students, break their phrases down for them when they use relative clauses."
                    "5. For A1 and A2 students, only recommend connectors such as deshalb, weil, ich mochte wissen,und,oder."
                    "- Give a revised score out of 25 (Score: X/25)."
                )
                with st.spinner("👨‍🏫 Herr Felix is comparing your improvement..."):
                    try:
                        result = client.chat.completions.create(
                            model="gpt-4o",
                            messages=[
                                {"role": "system", "content": ai_compare_prompt},
                                {"role": "user", "content": improved_letter}
                            ],
                            temperature=0.5,
                        )
                        compare_feedback = result.choices[0].message.content
                        st.session_state[f"{student_code}_delta_compare_feedback"] = compare_feedback
                        st.session_state[f"{student_code}_final_improved_letter"] = improved_letter
                    except Exception as e:
                        st.session_state[f"{student_code}_delta_compare_feedback"] = f"Sorry, there was an error comparing your letters: {e}"

            if st.session_state.get(f"{student_code}_delta_compare_feedback"):
                st.markdown("---")
                st.markdown("### 📝 Improvement Feedback from Herr Felix")
                st.markdown(highlight_feedback(st.session_state[f"{student_code}_delta_compare_feedback"]), unsafe_allow_html=True)

                # PDF & WhatsApp buttons
                from fpdf import FPDF
                import urllib.parse
                import os

                def sanitize_text(text):
                    return text

                # PDF
                pdf = FPDF()
                pdf.add_font("DejaVu", "", "font/DejaVuSans.ttf", uni=True)
                pdf.add_page()
                pdf.set_font("DejaVu", size=12)
                improved_letter = st.session_state.get(f"{student_code}_final_improved_letter", "")
                improved_feedback = st.session_state[f"{student_code}_delta_compare_feedback"]
                pdf.multi_cell(0, 10, f"Your Improved Letter:\n\n{sanitize_text(improved_letter)}\n\nFeedback from Herr Felix:\n\n{sanitize_text(improved_feedback)}")
                pdf_output = f"Feedback_{student_code}_{schreiben_level}_improved.pdf"
                pdf.output(pdf_output)
                with open(pdf_output, "rb") as f:
                    pdf_bytes = f.read()
                st.download_button(
                    "⬇️ Download Improved Version + Feedback (PDF)",
                    pdf_bytes,
                    file_name=pdf_output,
                    mime="application/pdf"
                )
                os.remove(pdf_output)

                # WhatsApp share
                wa_message = (
                    f"Hi, here is my IMPROVED German letter and AI feedback:\n\n"
                    f"{improved_letter}\n\n"
                    f"Feedback:\n{st.session_state[f'{student_code}_delta_compare_feedback']}"
                )
                wa_url = (
                    "https://api.whatsapp.com/send"
                    "?phone=233205706589"
                    f"&text={urllib.parse.quote(wa_message)}"
                )
                st.markdown(
                    f"[📲 Send Improved Letter & Feedback to Tutor on WhatsApp]({wa_url})",
                    unsafe_allow_html=True,
                )
                clear_feedback = st.button(
                    "🗑️ Clear feedback and start a new letter",
                    key=f"clear_feedback_{student_code}",
                )
                if clear_feedback:
                    for k in [
                        "last_feedback",
                        "last_user_letter",
                        "delta_compare_feedback",
                        "improved_letter",
                        "final_improved_letter",
                    ]:
                        st.session_state.pop(f"{student_code}_{k}", None)
                    st.session_state[f"{student_code}_awaiting_correction"] = False
                    st.session_state.pop(draft_key, None)
                    save_now(draft_key, student_code)
                    lv, lt, sf, sa = _draft_state_keys(draft_key)
                    for key in (lv, lt, sf, sa):
                        st.session_state.pop(key, None)
                    st.session_state["need_rerun"] = True
    if sub_tab == "Ideas Generator (Letter Coach)":
        import io

        # === NAMESPACED SESSION KEYS (per student) ===
        student_code = st.session_state.get("student_code", "demo")
        student_name = _safe_str(st.session_state.get("student_name"), "Student") or "Student"
        ns_prefix = f"{student_code}_letter_coach_"
        def ns(key): return ns_prefix + key

        prompt_draft_key = ns("prompt_draft")
        chat_draft_key = ns("chat_draft")
        drafts_hydrated_key = ns("drafts_hydrated")
        draft_sync_key = ns("draft_sync_cache")

        # --- Reset per-student Letter Coach state on student change ---
        prev_letter_coach_code = st.session_state.get("prev_letter_coach_code", None)
        if student_code != prev_letter_coach_code:
            last_prompt, last_chat = load_letter_coach_progress(student_code)
            st.session_state[ns("prompt")] = last_prompt or ""
            st.session_state[ns("chat")] = last_chat or []
            st.session_state[ns("stage")] = 1 if last_chat else 0
            st.session_state["prev_letter_coach_code"] = student_code
            st.session_state[drafts_hydrated_key] = False
            st.session_state[draft_sync_key] = ("", "")

        # --- Set per-student defaults if missing ---
        for k, default in [
            ("prompt", ""),
            ("chat", []),
            ("stage", 0),
            ("clear_prompt", False),
            ("clear_chat", False),
            ("clear_chat_draft", False),
            ("drafts_hydrated", False),
            ("draft_sync_cache", ("", "")),
        ]:
            if ns(k) not in st.session_state:
                st.session_state[ns(k)] = default

        if not st.session_state.get(drafts_hydrated_key):
            prompt_draft, chat_draft, _, updated_at = load_letter_coach_draft(student_code)
            reset_local_draft_state(
                prompt_draft_key,
                prompt_draft or "",
                saved=bool(prompt_draft),
                saved_at=updated_at,
            )
            reset_local_draft_state(
                chat_draft_key,
                chat_draft or "",
                saved=bool(chat_draft),
                saved_at=updated_at,
            )
            st.session_state[draft_sync_key] = (
                prompt_draft or "",
                chat_draft or "",
            )
            st.session_state[drafts_hydrated_key] = True

        def sync_letter_coach_draft_state() -> None:
            if not student_code:
                return
            prompt_text = st.session_state.get(prompt_draft_key, "") or ""
            chat_text = st.session_state.get(chat_draft_key, "") or ""
            payload = (prompt_text, chat_text)
            if st.session_state.get(draft_sync_key) == payload:
                return
            save_letter_coach_draft(student_code, prompt_text, chat_text)
            st.session_state[draft_sync_key] = payload

        def clear_letter_coach_draft_state() -> None:
            if not student_code:
                return
            clear_letter_coach_draft(student_code)
            st.session_state[draft_sync_key] = ("", "")

        def save_prompt_draft_now(*, show_toast: bool = True) -> None:
            save_now(prompt_draft_key, student_code, show_toast=show_toast)
            sync_letter_coach_draft_state()

        def save_chat_draft_now(*, show_toast: bool = True) -> None:
            save_now(chat_draft_key, student_code, show_toast=show_toast)
            sync_letter_coach_draft_state()

        if st.session_state.get(ns("reset_coach")):
            st.session_state[ns("prompt")] = ""
            st.session_state[ns("chat")] = []
            st.session_state[ns("stage")] = 0
            reset_local_draft_state(prompt_draft_key, "")
            reset_local_draft_state(chat_draft_key, "")
            save_prompt_draft_now()
            save_chat_draft_now()
            clear_letter_coach_draft_state()
            st.session_state.pop(ns("reset_coach"))

        st.markdown(
            "**[German Writing Rules](https://drive.google.com/file/d/1o7_ez3WSNgpgxU_nEtp6EO1PXDyi3K3b/view?usp=sharing)**",
        )
        

# Patch: enforce '…möchte' in the reason line + add conclusion 'Ich freue mich im Voraus auf Ihre/deine Antwort'
        LETTER_COACH_PROMPTS = {
            "A1": (
                "You are Herr Felix, a creative, supportive German letter-writing coach for A1 students. "
                "Your mission: idea generator + step-by-step coach. "
                "Always reply in English. You may show short German fragments (2–3 words), but never full sentences. "
                "Classify each student message as: NEW PROMPT, CONTINUATION, or QUESTION. "
                "• If QUESTION: answer simply, encourage progress, then prompt the next step only. "
                "• If CONTINUATION: give brief corrections and guide the next step only. "
                "    1) Give short ideas, structure, tips, and 2–3 word German fragments. Don’t overfeed; let them think. "
                "    2) Allowed connectors only: 'und', 'aber', 'weil', 'deshalb', 'ich möchte wissen, ob', 'ich möchte wissen, wann'. "
                "       Do NOT suggest 'da', 'dass', relative clauses, or advanced tenses (keep present + basic modals). "
                "    3) Requests: teach 'Könnten Sie … [Infinitiv am Ende]?' and show the main verb at the end. "
                "    4) Greeting + Introduction (fragments; enforce comma/space + register consistency): "
                "       • Formal: 'Ich hoffe, es geht Ihnen gut. Ich schreibe Ihnen, weil ich … möchte.' "
                "       • Informal: 'Wie geht es dir? Ich hoffe, es geht dir gut. Ich schreibe dir, weil ich … möchte.' "
                "       The reason line must END with 'möchte' to keep it simple and consistent at A1. "
                "    5) Closing/conclusion templates (teach explicitly): "
                "       • Formal: 'Ich freue mich im Voraus auf Ihre Antwort.'  +  'Mit freundlichen Grüßen,' + [Name] "
                "       • Informal: 'Ich freue mich im Voraus auf deine Antwort.'  +  'Viele Grüße,' + [Name] "
                "    6) Word-order guard rails: "
                "       • 'weil'-clause → verb at the end (warn if not). "
                "       • Requests with 'Könnten Sie' → infinitive at the end (warn if not). "
                "    7) If a line uses ≥2 conjunctions or is long/complex, warn and split. "
                "    8) If a line exceeds ~7–8 words, break into short statements with full stops. "
                "    9) Final letter length target: 25–35 words. "
                "    10) Scenario mini-banks (fragments only): "
                "        • Cancellation: Wetter/Gesundheit + 'Termin absagen' "
                "        • Enquiry/registration: 'Anfrage stellen'; add 'Wie viel kostet …?' "
                "        • Registration/course: 'anfangen'/'beginnen' "
                "        • Appointment: 'neuen Termin vereinbaren' "
                "        • Apology: 'Es tut mir leid.' "
                "    11) Never write full sentences. Provide only fragments/keywords; the student writes the sentences. "
                "    12) Remind students to type their own words (no translator); you will correct them. "
                "If NEW PROMPT: give a 5-part overview (greeting, introduction, reason, request, closing) with micro-examples (fragments only). "
                "Always end with: 'Your next recommended step:' and request exactly one part at a time—first greeting (wait), then introduction (wait), then reason, then request, then closing. "
                "After each reply: correct briefly, give one tip, then again: 'Your next recommended step:' for the next single part. "
                "Session pacing: aim to finish in ~10 student replies. If not done, say: 'Most letters can be completed in about 10 steps. Please try to finish soon.' "
                "At 14 replies without completion: 'We have reached the end of this coaching session. Please copy your letter below so far and paste it into the “Mark My Letter” tool for full AI feedback and a score.' "
            ),
            "A2": (
                "You are Herr Felix, a creative, supportive German letter-writing coach for A2 students. "
                "Role: idea generator + step-by-step coach. "
                "Always reply in English. You may show short German fragments (2–3 words), never full sentences. "
                "Classify each message: NEW PROMPT, CONTINUATION, or QUESTION. "
                "• If QUESTION: answer simply, encourage progress, then prompt the next step only. "
                "• If CONTINUATION: correct briefly and guide the next step only. "
                "    1) Require sequencing: 'Zuerst' (first idea), 'Dann' or 'Außerdem' (next idea), 'Zum Schluss' (final/closing bridge). Prefer 'Zuerst' over 'Erstens'. "
                "    2) Connectors: 'und', 'aber', 'weil', 'denn', 'deshalb', 'ich möchte wissen, ob/wann/wo'. Recommend one at a time; if ≥2 used in a short line, simplify to one. "
                "    3) Greeting + Introduction templates (teach explicitly; split into fragments if needed): "
                "       • Formal: 'Ich hoffe, es geht Ihnen gut. Ich schreibe Ihnen, weil …' "
                "       • Informal: 'Wie geht es dir? Ich hoffe, es geht dir gut. Ich schreibe dir, weil …' "
                "       Enforce comma after 'Ihnen/dir,' in 'Ich schreibe Ihnen/dir, weil …'. "
                "    4) After every reply, give one tip or one phrase fragment—never full sentences. "
                "    5) Keep lines short: ~7–8 words; split long lines. "
                "    6) Letter length target: 30–40 words. "
                "    7) Scenarios: cancellations (health/weather; 'absagen'), enquiries/registrations ('Anfrage stellen'; include 'Wie viel kostet …?'), appointments ('neuen Termin vereinbaren'). "
                "    8) Apologies: 'Es tut mir leid.' "
                "    9) Always correct grammar and suggest improved fragments when needed. "
                "Steps: greeting → introduction → 'Zuerst' idea → 'Außerdem' (or 'Dann') → 'Zum Schluss' → polite closing cue ('Ich freue mich …'). "
                "Always end with: 'Your next recommended step:' and ask for exactly one section at a time. "
                "Do not write the full letter; guide only. Remind students to type their own words; you will correct them. "
                "Session pacing: finish in ~10 replies; if not, remind to finish soon. End at 14 with: copy/paste into 'Mark My Letter' for feedback. "
            ),
            "B1": (
                "You are Herr Felix, a supportive German letter/essay coach for B1 students—idea generator + step-by-step coach. "
                "Always reply in English; show only short German fragments (2–3 words), never full sentences. "
                "Detect type: formal letter, informal letter, or opinion essay. If unclear, ask which type. "
                "    1) Give short ideas, structure, tips, and 2–3 word German fragments. Don’t overfeed. "
                "    2) Enforce paragraph logic with clear starters/sequence. "
                "    3) After each student line, add 1–2 ideas if helpful (fragments only). "
                "    4) Length targets: formal letter 40–50 words; informal letter & opinion essay 80–90 words (intro, body, conclusion). "
                "    5) Provide fragments only; the student completes each sentence. "
                "    6) Remind them to type their own words; you will correct mistakes. "
                "    7) Never write full sentences for them. "
                "Greeting options for forum/opinion posts (teach explicitly; choose one): "
                "    • 'Hallo zusammen,'  • 'Liebe Forenmitglieder,'  • 'Liebes Forum,'  • 'Liebe Community,' "
                "    Avoid: 'Lieber Forummitglieder' (wrong: gender/number and compound). "
                "Opinion essay template (fragments only): "
                "    • 'Heutzutage ist das Thema' + [Thema] + 'ein wichtiges Thema in unserem Leben.' "
                "    • 'Ich bin der Meinung, dass' + [Info] + ', weil' + [Info] + '.' "
                "    • 'Einerseits gibt es viele Vorteile.'  'Zum Beispiel' + [Verb/Modal] + [Info] + '.' "
                "    • 'Andererseits gibt es auch Nachteile.'  'Ein Beispiel dafür ist' + [Nomen] + '. ' + 'Kleine Info.' "
                "    • 'Ich glaube, dass' + [eigene Meinung] + '.' "
                "    • 'Zusammenfassend lässt sich sagen, dass' + [Thema] + '… positiv/negativ … beeinflussen kann.' "
                "Process: ask one section at a time with 'Your next recommended step:' (intro → pros → cons → opinion → conclusion). "
                "Session pacing: ~10 replies; end at 14 with 'Mark My Letter'. "
            ),
            "B2": (
                "You are Herr Felix, a supportive German writing coach for B2—idea generator + step-by-step coach. "
                "Always reply in English; you may show short German fragments (2–3 words), never full sentences. "
                "Detect type: formal letter, informal letter, or opinion/argumentative essay. If unclear, ask which type. "
                "    1) Give short ideas, structure, tips, and 2–3 word German fragments. Don’t overfeed. "
                "    2) Enforce paragraph logic with clear sequence and topic focus. "
                "    3) Add 1–2 ideas after each submission (fragments only) if helpful. "
                "    4) Length targets: formal letter 100–150 words; opinion/argumentative essay 150–170 words. "
                "Forum/opinion post greetings (pick one): "
                "    • 'Hallo zusammen,'  • 'Liebe Forenmitglieder,'  • 'Liebes Forum,'  • 'Liebe Community,' "
                "    Avoid: 'Lieber Forummitglieder' (wrong form). "
                "Opinion/argumentative scaffold (fragments only): thesis → arguments (pro) → counter (contra) → conclusion. "
                "Always end with 'Your next recommended step:' and request exactly one section at a time. "
                "Session pacing: finish in ~10 replies; end at 14 with 'Mark My Letter' paste instruction. "
            ),


            "B2": (
                "You are Herr Felix, a supportive German writing coach for B2—idea generator + step-by-step coach. "
                "Always reply in English; you may show short German fragments (2–3 words), never full sentences. "
                "Detect type: formal letter, informal letter, or opinion/argumentative essay. If unclear, ask which type. "
                "    1) Give short ideas, structure, tips, and 2–3 word German fragments. Don’t overfeed. "
                "    2) Enforce paragraph logic with clear sequence and topic focus. "
                "    3) Add 1–2 ideas after each student submission (fragments only) if helpful. "
                "    4) Length targets: formal letter 100–150 words; opinion/argumentative essay 150–170 words. "
                "    5) Always correct grammar and suggest stronger phrasing (fragments). "
                "Greeting + Introduction templates (teach explicitly; split into fragments if needed): "
                "    • Formal: 'Ich hoffe, es geht Ihnen gut. Ich schreibe Ihnen, weil …' "
                "    • Informal: 'Wie geht es dir? Ich hoffe, es geht dir gut. Ich schreibe dir, weil …' "
                "    Enforce comma after 'Ihnen/dir,' in 'Ich schreibe Ihnen/dir, weil …'. "
                "Formal: greeting → intro → clear argument/reason → supporting details → closing. "
                "Informal: greeting → personal intro → main point/reason → brief examples → closing. "
                "Opinion/argumentative: intro with thesis → arguments with examples → counterargument(s) → conclusion. "
                "Always end with: 'Your next recommended step:' and ask for exactly one section at a time. "
                "After each reply, give feedback, then 'Your next recommended step:' again. "
                "Model connectors: 'denn', 'dennoch', 'außerdem', 'jedoch', 'zum Beispiel', 'einerseits … andererseits'. "
                "Session pacing: finish in ~10 replies; end at 14 with 'Mark My Letter' paste instruction. "
            ),
            "C1": (
                "You are Herr Felix, an advanced, supportive German writing coach for C1—idea generator + step-by-step coach. "
                "Primarily reply in English; you may include German where useful and then explain it clearly. "
                "Detect type: formal letter, informal letter, or academic/opinion essay. If unclear, ask which type. "
                "    1) Give short ideas, structure, tips, and 2–3 word German fragments. Don’t overfeed. "
                "    2) Enforce paragraph logic with sequence, cohesion, and topic sentences. "
                "    3) Add 1–2 ideas after each submission (fragments) if helpful. "
                "    4) Length targets: formal letter 120–150 words; opinion/academic essay 230–250 words. "
                "    5) Correct grammar and suggest precise, higher-register phrasing (explain briefly if advanced). "
                "Formal: greeting → sophisticated introduction → detailed argument → evidence/examples → closing. "
                "Informal: greeting → nuanced intro → main point/reason → personal stance → polished closing. "
                "Academic/opinion: intro with thesis & context → structured arguments → counterpoints → conclusion. "
                "Always end with: 'Your next recommended step:' and ask for exactly one section at a time. "
                "After each answer, provide feedback, then 'Your next recommended step:' again. "
                "Model advanced connectors: 'nicht nur … sondern auch', 'obwohl', 'dennoch', 'folglich', 'somit'. "
                "Session pacing: finish in ~10 replies; at 14, end and ask the student to paste into 'Mark My Letter' for scoring. "
            ),
        }



        def reset_letter_coach():
            for k in [
                "letter_coach_stage", "letter_coach_chat", "letter_coach_prompt",
                "letter_coach_type", "selected_letter_lines", "letter_coach_uploaded"
            ]:
                st.session_state[k] = 0 if k == "letter_coach_stage" else []
            st.session_state["letter_coach_uploaded"] = False

        def bubble(role, text):
            if role == "assistant":
                return f"""<div style='background: #f4eafd; color: #7b2ff2; border-radius: 16px 16px 16px 3px; margin-bottom: 8px; margin-right: 80px; box-shadow: 0 2px 8px rgba(123,47,242,0.08); padding: 13px 18px; text-align: left; max-width: 88vw; font-size: 1.12rem;'><b>👨‍🏫 Herr Felix:</b><br>{text}</div>"""
            return f"""<div style='background: #eaf4ff; color: #1a237e; border-radius: 16px 16px 3px 16px; margin-bottom: 8px; margin-left: 80px; box-shadow: 0 2px 8px rgba(26,35,126,0.07); padding: 13px 18px; text-align: right; max-width: 88vw; font-size: 1.12rem;'><b>🙋 You:</b><br>{text}</div>"""

        # --- General Instructions for Students (Minimal Welcome + Subline) ---
        st.markdown(
            """
            <div style="
                background: linear-gradient(97deg, #f4eafd 75%, #ffe0f5 100%);
                border-radius: 12px;
                border: 1px solid #e6d3fa;
                box-shadow: 0 2px 8px #e5e1fa22;
                padding: 0.75em 1em 0.72em 1em;
                margin-bottom: 1.1em;
                margin-top: 0.1em;
                color: #4b2976;
                font-size: 1.03rem;
                font-family: 'Segoe UI', 'Roboto', 'Arial', sans-serif;
                text-align: center;
                ">
                <span style="font-size:1.19em; vertical-align:middle;">✉️</span>
                <span style="font-size:1.05em; font-weight: 500; margin-left:0.24em;">
                    Welcome to <span style="color:#7b2ff2;">Letter Coach</span>
                </span>
                <div style="color:#b48be6; font-size:0.97em; margin-top:0.35em;">
                    Get started below 👇
                </div>
            </div>
            """,
            unsafe_allow_html=True
        )

        # --- Stage 0: Prompt input ---
        if st.session_state[ns("stage")] == 0:
            if st.button("Start new write-up"):
                st.session_state[ns("reset_coach")] = True
                st.session_state["need_rerun"] = True
            st.markdown("### ✏️ Enter your exam prompt or draft to start coaching")
            draft_key = prompt_draft_key
            initialize_draft_state(student_code, draft_key)


            if st.session_state.pop(ns("clear_prompt"), False):
                reset_local_draft_state(draft_key, "")
                save_prompt_draft_now()

            prompt = st.text_area(
                "Exam prompt",
                key=draft_key,
                height=120,
                placeholder="e.g., Schreiben Sie eine formelle E-Mail an Ihre Nachbarin ...",
                label_visibility="collapsed",
                on_change=save_prompt_draft_now,
            )

            render_umlaut_pad(
                draft_key,
                context=f"letter_coach_prompt_{student_code}",
                disabled=False,
            )

            autosave_maybe(
                student_code,
                draft_key,
                st.session_state.get(draft_key, ""),
                min_secs=2.0,
                min_delta=12,
            )
            sync_letter_coach_draft_state()

            st.caption("Draft auto-saves every few seconds.")

            saved_at = st.session_state.get(f"{draft_key}_saved_at")
            if saved_at:
                st.caption(f"Last saved at {saved_at.strftime('%H:%M:%S')}")

            prompt = st.session_state.get(draft_key, "")

            if prompt:
                word_count = len(prompt.split())
                char_count = len(prompt)
                st.markdown(
                    (
                        "<div style='color:#7b2ff2; font-size:0.97em; margin-bottom:0.18em;'>"
                        f"Words: <b>{word_count}</b> &nbsp;|&nbsp; Characters: <b>{char_count}</b>"
                        "</div>"
                    ),
                    unsafe_allow_html=True,
                )

            if st.button("✉️ Start Letter Coach"):
                save_prompt_draft_now()

                prompt = st.session_state.get(draft_key, "")
                if prompt:
                    st.session_state[ns("prompt")] = prompt
                    student_level = st.session_state.get("schreiben_level", "A1")
                    system_prompt = LETTER_COACH_PROMPTS[student_level].format(prompt=prompt)
                    system_prompt += (
                        f" The student's name is {student_name}. Always greet {student_name} and weave their name into your encouragement and guidance."
                    )
                    chat_history = [
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": prompt}
                    ]
                    try:
                        resp = client.chat.completions.create(
                            model="gpt-4o",
                            messages=chat_history,
                            temperature=0.22,
                            max_tokens=380
                        )
                        ai_reply = resp.choices[0].message.content
                    except Exception:
                        ai_reply = "Sorry, there was an error generating a response. Please try again."
                    chat_history.append({"role": "assistant", "content": ai_reply})

                    st.session_state[ns("chat")] = chat_history
                    st.session_state[ns("stage")] = 1
                    save_letter_coach_progress(
                        student_code,
                        student_level,
                        st.session_state[ns("prompt")],
                        st.session_state[ns("chat")],
                    )
                    st.session_state[ns("clear_prompt")] = True
                    st.session_state["need_rerun"] = True
                    
            if prompt:
                st.markdown("---")
                st.markdown(f"📝 **Letter/Essay Prompt or Draft:**\n\n{prompt}")

        # --- Stage 1: Coaching Chat ---
        elif st.session_state[ns("stage")] == 1:
            st.markdown("---")
            st.markdown(f"📝 **Letter/Essay Prompt:**\n\n{st.session_state[ns('prompt')]}")
            chat_history = st.session_state[ns("chat")]
            for msg in chat_history[1:]:
                st.markdown(bubble(msg["role"], msg["content"]), unsafe_allow_html=True)
            num_student_turns = sum(1 for msg in chat_history[1:] if msg["role"] == "user")
            if num_student_turns == 10:
                st.info("🔔 You have written 10 steps. Most students finish in 7–10 turns. Try to complete your letter soon!")
            elif num_student_turns == 12:
                st.warning(
                    "⏰ You have reached 12 writing turns. "
                    "Usually, your letter should be complete by now. "
                    "If you want feedback, click **END SUMMARY** or download your letter as TXT. "
                    "You can always start a new session for more practice."
                )
            elif num_student_turns > 12:
                st.warning(
                    f"🚦 You are now at {num_student_turns} turns. "
                    "Long letters are okay, but usually a good letter is finished in 7–12 turns. "
                    "Try to wrap up, click **END SUMMARY** or download your letter as TXT."
                )


            draft_key = chat_draft_key
            initialize_draft_state(student_code, draft_key)


            if st.session_state.pop(ns("clear_chat_draft"), False):
                reset_local_draft_state(draft_key, "")
                save_chat_draft_now(show_toast=False)

            if st.session_state.pop(ns("clear_chat"), False):
                reset_local_draft_state(draft_key, "")
                save_chat_draft_now(show_toast=False)

            st.text_area(
                "Chat input",
                key=draft_key,
                height=400,
                placeholder="Type your reply, ask about a section, or paste your draft here...",
                label_visibility="collapsed",

            )

            render_umlaut_pad(
                draft_key,
                context=f"letter_coach_chat_{student_code}",
                disabled=False,
            )
            
            autosave_maybe(
                student_code,
                draft_key,
                st.session_state.get(draft_key, ""),
                min_secs=0.2,
                min_delta=1,
            )
            sync_letter_coach_draft_state()
           
            saved_at = st.session_state.get(f"{draft_key}_saved_at")
            if saved_at:
                st.caption(f"Last saved at {saved_at.strftime('%H:%M:%S')}")

            letter_draft_key = ns("letter_draft_saved")
            initialize_draft_state(student_code, letter_draft_key)
            letter_draft = st.session_state.get(letter_draft_key, "")

            def _reset_letter_coach_session() -> None:
                """Clear the current coaching session and return to the prompt stage."""

                st.session_state[ns("clear_chat_draft")] = True
                st.session_state[ns("chat")] = []
                st.session_state[ns("prompt")] = ""
                st.session_state[ns("selected_letter_lines")] = []
                st.session_state[ns("reset_coach")] = True
                reset_local_draft_state(prompt_draft_key, "")
                save_prompt_draft_now(show_toast=False)
                clear_letter_coach_draft_state()
                st.session_state["need_rerun"] = True

            send_col, reset_col = st.columns([2, 3])
            with send_col:
                send = st.button(
                    "Send",
                    key=ns("send_button"),
                    type="primary",
                    use_container_width=True,
                )
            with reset_col:
                start_new_top = st.button(
                    "Start New Letter Coach",
                    key=ns("start_new_letter_top"),
                    help="Clear this chat and begin a fresh coaching session.",
                    use_container_width=True,
                )

            if start_new_top:
                _reset_letter_coach_session()

            if send:
                user_input = st.session_state[draft_key].strip()
                save_chat_draft_now()

            else:
                user_input = ""

            if user_input:
                chat_history.append({"role": "user", "content": user_input})
                student_level = st.session_state.get("schreiben_level", "A1")
                system_prompt = LETTER_COACH_PROMPTS[student_level].format(prompt=st.session_state[ns("prompt")])
                system_prompt += (
                    f" The student's name is {student_name}. Always greet {student_name} warmly and include their name in your feedback."
                )
                with st.spinner("👨‍🏫 Herr Felix is typing..."):
                    resp = client.chat.completions.create(
                        model="gpt-4o",
                        messages=[{"role": "system", "content": system_prompt}] + chat_history[1:],
                        temperature=0.22,
                        max_tokens=380
                    )
                    ai_reply = resp.choices[0].message.content
                chat_history.append({"role": "assistant", "content": ai_reply})
                st.session_state[ns("chat")] = chat_history
                save_letter_coach_progress(
                    student_code,
                    student_level,
                    st.session_state[ns("prompt")],
                    st.session_state[ns("chat")],
                )
                st.session_state[ns("clear_chat")] = True
                st.session_state["need_rerun"] = True

            # ----- LIVE AUTO-UPDATING LETTER DRAFT, Download + Copy -----
            import streamlit.components.v1 as components

            user_msgs = [
                msg["content"]
                for msg in st.session_state[ns("chat")][1:]
                if msg.get("role") == "user"
            ]

            st.markdown("""
                **📝 Your Letter Draft**
                - Tick the lines you want to include in your letter draft.
                - You can untick any part you want to leave out.
                - Only ticked lines will appear in your downloadable draft below.
            """)

            saved_letter_draft = st.session_state.get(letter_draft_key, "")
            saved_at_key = f"{letter_draft_key}_saved_at"
            saved_at = st.session_state.get(saved_at_key)
            selection_key = ns("selected_letter_lines")

            def _restore_selected_lines() -> List[bool]:
                if not user_msgs:
                    return []
                if not saved_at:
                    return [True] * len(user_msgs)
                saved_lines = saved_letter_draft.splitlines()
                if not saved_lines:
                    return [False] * len(user_msgs)
                remaining = Counter(saved_lines)
                restored: List[bool] = []
                for msg in user_msgs:
                    if remaining.get(msg, 0):
                        restored.append(True)
                        remaining[msg] -= 1
                    else:
                        restored.append(False)
                return restored

            if selection_key not in st.session_state:
                st.session_state[selection_key] = _restore_selected_lines()
            else:
                current_selection = st.session_state[selection_key]
                if len(current_selection) < len(user_msgs):
                    current_selection.extend([True] * (len(user_msgs) - len(current_selection)))
                elif len(current_selection) > len(user_msgs):
                    st.session_state[selection_key] = current_selection[:len(user_msgs)]

            selected_lines: List[str] = []
            previous_letter_draft = saved_letter_draft
            for i, line in enumerate(user_msgs):
                selected = st.checkbox(
                    line,
                    value=st.session_state[selection_key][i],
                    key=ns(f"letter_line_{i}")
                )
                st.session_state[selection_key][i] = selected
                if selected:
                    selected_lines.append(line)

            letter_draft = "\n".join(selected_lines) if selected_lines else ""
            if letter_draft != previous_letter_draft:
                st.session_state[letter_draft_key] = letter_draft
                autosave_maybe(
                    student_code,
                    letter_draft_key,
                    letter_draft,
                    min_secs=0.0,
                    min_delta=0,
                )
                saved_at = st.session_state.get(saved_at_key)

            letter_draft = st.session_state.get(letter_draft_key, "")

            saved_at = st.session_state.get(saved_at_key)
            if saved_at:
                st.caption(f"Last saved at {saved_at.strftime('%H:%M:%S')}")

            # --- Live word/character count for the letter draft ---
            draft_word_count = len(letter_draft.split())
            draft_char_count = len(letter_draft)
            st.markdown(
                f"<div style='color:#7b2ff2; font-size:0.97em; margin-bottom:0.18em;'>"
                f"Words: <b>{draft_word_count}</b> &nbsp;|&nbsp; Characters: <b>{draft_char_count}</b>"
                "</div>",
                unsafe_allow_html=True
            )

            # --- Modern, soft header (copy/download) ---
            st.markdown(
                """
                <div style="
                    background:#23272b;
                    color:#eee;
                    border-radius:10px;
                    padding:0.72em 1.04em;
                    margin-bottom:0.4em;
                    font-size:1.07em;
                    font-weight:400;
                    border:1px solid #343a40;
                    box-shadow:0 2px 10px #0002;
                    text-align:left;
                ">
                    <span style="font-size:1.12em; color:#ffe082;">📝 Your Letter So Far</span><br>
                    <span style="font-size:1.00em; color:#b0b0b0;">copy often or download below to prevent data loss</span>
                </div>
                """,
                unsafe_allow_html=True
            )

            # --- Mobile-friendly copy/download box ---
            components.html(f"""
                <textarea id="letterBox_{student_code}" readonly rows="6" style="
                    width: 100%;
                    border-radius: 12px;
                    background: #f9fbe7;
                    border: 1.7px solid #ffe082;
                    color: #222;
                    font-size: 1.12em;
                    font-family: 'Fira Mono', 'Consolas', monospace;
                    padding: 1em 0.7em;
                    box-shadow: 0 2px 8px #ffe08266;
                    margin-bottom: 0.5em;
                    resize: none;
                    overflow:auto;
                " onclick="this.select()">{letter_draft}</textarea>
                <button onclick="navigator.clipboard.writeText(document.getElementById('letterBox_{student_code}').value)" 
                    style="
                        background:#ffc107;
                        color:#3e2723;
                        font-size:1.08em;
                        font-weight:bold;
                        padding:0.48em 1.12em;
                        margin-top:0.4em;
                        border:none;
                        border-radius:7px;
                        cursor:pointer;
                        box-shadow:0 2px 8px #ffe08255;
                        width:100%;
                        max-width:320px;
                        display:block;
                        margin-left:auto;
                        margin-right:auto;
                    ">
                    📋 Copy Text
                </button>
                <style>
                    @media (max-width: 480px) {{
                        #letterBox_{student_code} {{
                            font-size: 1.16em !important;
                            min-width: 93vw !important;
                        }}
                    }}
                </style>
            """, height=175)

            st.markdown("""
                <div style="
                    background:#ffe082;
                    padding:0.9em 1.2em;
                    border-radius:10px;
                    margin:0.4em 0 1.2em 0;
                    color:#543c0b;
                    font-weight:600;
                    border-left:6px solid #ffc107;
                    font-size:1.08em;">
                    📋 <span>On phone, tap in the box above to select all for copy.<br>
                    Or just tap <b>Copy Text</b>.<br>
                    To download, use the button below.</span>
                </div>
            """, unsafe_allow_html=True)

            st.download_button(
                "⬇️ Download Letter as TXT",
                letter_draft.encode("utf-8"),
                file_name="my_letter.txt"
            )

            if st.button(
                "Start New Letter Coach",
                key=ns("start_new_letter_bottom"),
                help="Clear this chat and return to the prompt input screen.",
            ):
                _reset_letter_coach_session()


















































if st.session_state.pop("need_rerun", False):
    # Mark done so we don't schedule again
    st.session_state["post_login_rerun"] = True
    st.rerun()

